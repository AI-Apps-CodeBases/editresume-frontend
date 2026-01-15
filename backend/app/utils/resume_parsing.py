"""Resume parsing utilities - extracted from legacy_app.py"""

from __future__ import annotations

import logging
import re

logger = logging.getLogger(__name__)


def extract_pdf_text(file_content: bytes) -> tuple[str, list[str]]:
    """Extract text from PDF using multiple methods"""
    text = ""
    methods = []

    # Method 1: PyMuPDF (fitz) - Most reliable
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_content, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        methods.append("PyMuPDF")
        logger.info(f"PyMuPDF extracted {len(text)} characters")
    except Exception as e:
        logger.warning(f"PyMuPDF failed: {e}")

    # Method 2: pdfplumber - Good for complex layouts
    if len(text.strip()) < 100:
        try:
            import io

            import pdfplumber

            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            methods.append("pdfplumber")
            logger.info(f"pdfplumber extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

    # Method 3: PyPDF2 - Fallback
    if len(text.strip()) < 100:
        try:
            import io

            import PyPDF2

            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            methods.append("PyPDF2")
            logger.info(f"PyPDF2 extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")

    return text, methods


def extract_docx_text(file_content: bytes) -> tuple[str, list[str]]:
    """Extract text from DOCX using multiple methods"""
    text = ""
    methods = []

    # Method 1: python-docx - Standard method
    try:
        import io

        from docx import Document

        doc = Document(io.BytesIO(file_content))

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

        methods.append("python-docx")
        logger.info(f"python-docx extracted {len(text)} characters")
    except Exception as e:
        logger.warning(f"python-docx failed: {e}")

    # Method 2: docx2txt - Alternative method
    if len(text.strip()) < 100:
        try:
            import io

            import docx2txt

            text = docx2txt.process(io.BytesIO(file_content))
            methods.append("docx2txt")
            logger.info(f"docx2txt extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"docx2txt failed: {e}")

    return text, methods


def extract_doc_text(file_content: bytes) -> tuple[str, list[str]]:
    """Extract text from DOC files"""
    text = ""
    methods = []

    # Method 1: antiword (if available)
    try:
        import os
        import subprocess
        import tempfile

        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name

        try:
            result = subprocess.run(
                ["antiword", temp_file_path],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0:
                text = result.stdout
                methods.append("antiword")
                logger.info(f"antiword extracted {len(text)} characters")
        finally:
            os.unlink(temp_file_path)
    except Exception as e:
        logger.warning(f"antiword failed: {e}")

    # Method 2: Try reading as text (fallback)
    if len(text.strip()) < 100:
        try:
            text = file_content.decode("utf-8", errors="ignore")
            methods.append("UTF-8 decode")
            logger.info(f"UTF-8 decode extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"UTF-8 decode failed: {e}")

    return text, methods


def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""

    # Remove excessive whitespace
    text = re.sub(r"\n\s*\n\s*\n", "\n\n", text)  # Max 2 newlines
    text = re.sub(r"[ \t]+", " ", text)  # Multiple spaces to single

    # Remove common PDF artifacts
    text = re.sub(r"\f", "\n", text)  # Form feeds to newlines

    # Clean up bullet points
    text = re.sub(r"[•·▪▫‣⁃]", "•", text)  # Normalize bullets

    return text.strip()


def _collapse_spaced_letters(line: str) -> str:
    """Collapse spaced letters like 'M E R T' -> 'MERT' within a line."""
    def collapse_chunk(chunk: str) -> str:
        return re.sub(
            r"(?:\b[A-Za-z]\s){2,}[A-Za-z]\b",
            lambda m: m.group(0).replace(" ", ""),
            chunk,
        )

    # Split on large gaps to avoid merging words from different columns
    parts = re.split(r" {2,}|\t+", line)
    collapsed = [collapse_chunk(part) for part in parts]
    return " ".join([p for p in collapsed if p])


def _repair_hyphenation(text: str) -> str:
    """Join words split by hyphenated line breaks."""
    return re.sub(r"(\w)-\n(\w)", r"\1\2", text)


def normalize_extracted_text(text: str) -> str:
    """Normalize extracted text for downstream parsing."""
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _repair_hyphenation(text)
    text = re.sub(r"^--- Page \d+ ---\s*$", "", text, flags=re.MULTILINE)

    normalized_lines = [_collapse_spaced_letters(line) for line in text.split("\n")]
    normalized = "\n".join(normalized_lines)
    normalized = clean_extracted_text(normalized)
    return normalized


def parse_resume_with_regex(text: str) -> dict:
    """Simple, reliable resume parser that actually works"""
    logger.info("Using simple regex-based resume parsing")

    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # Extract basic info
    name = ""
    title = ""
    email = ""
    phone = ""
    location = ""

    # Find name (first line that looks like a name)
    for line in lines[:3]:
        if (
            len(line) > 2
            and not any(char.isdigit() for char in line)
            and "@" not in line
            and "•" not in line
        ):
            name = line
            break

    # Find title (line after name)
    for i, line in enumerate(lines[1:4]):
        if (
            len(line) > 2
            and not any(char.isdigit() for char in line)
            and "@" not in line
            and "•" not in line
            and line != name
        ):
            title = line
            break

    # Find email
    email_match = re.search(
        r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", text
    )
    if email_match:
        email = email_match.group()

    # Find phone
    phone_match = re.search(
        r"(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})", text
    )
    if phone_match:
        phone = phone_match.group().strip()

    # Find location
    location_match = re.search(r"([A-Za-z\s]+),\s*([A-Z]{2})", text)
    if location_match:
        location = location_match.group().strip()

    # Extract sections
    sections = []
    section_dict = {}  # Track sections by normalized title to merge duplicates
    current_section = None
    current_bullets = []

    section_keywords = [
        "work experience",
        "professional experience",
        "employment",
        "experience",
        "projects",
        "education",
        "skills",
        "certifications",
    ]

    def normalize_section_title(title: str) -> str:
        """Normalize section title for comparison (case-insensitive, trimmed, semantic grouping)"""
        normalized = title.lower().strip()

        # Map semantic variations to canonical names
        semantic_map = {
            "work experience": "work experience",
            "professional experience": "work experience",
            "employment": "work experience",
            "employment history": "work experience",
            "career history": "work experience",
            "professional history": "work experience",
            "work history": "work experience",
            "experience": "work experience",  # If "experience" appears alone, treat as work experience
            "academic projects": "projects",
            "project experience": "projects",
            "project": "projects",
            "technical skills": "skills",
            "core competencies": "skills",
            "competencies": "skills",
            "expertise": "skills",
            "skill": "skills",
            "education & training": "education",
            "academic background": "education",
            "educational background": "education",
            "certification": "certifications",
            "certificate": "certifications",
            "award": "awards",
            "honor": "awards",
            "honors": "awards",
        }

        # Check if any semantic mapping applies
        for variant, canonical in semantic_map.items():
            if variant in normalized:
                return canonical

        # Return normalized title if no semantic mapping found
        return normalized

    def save_section(section_title: str, bullets: list):
        """Save or merge section with existing one if duplicate"""
        if not section_title:
            return

        # Filter out empty bullets
        filtered_bullets = [b for b in bullets if b and str(b).strip()]
        if not filtered_bullets:
            return

        normalized_title = normalize_section_title(section_title)

        # Check if section with same normalized title already exists
        if normalized_title in section_dict:
            # Merge bullets into existing section
            existing_section = section_dict[normalized_title]
            existing_bullets = [b["text"] for b in existing_section["bullets"]]

            # Use the most descriptive title (prefer longer, more specific titles)
            if len(section_title) > len(existing_section["title"]):
                existing_section["title"] = section_title

            # Add separator if needed (empty string for work experience sections)
            is_work_exp = normalized_title == "work experience"
            if is_work_exp and existing_bullets and existing_bullets[-1] != "":
                existing_bullets.append("")

            # Add new bullets, removing any duplicates (exact match and near-duplicates)
            existing_bullet_texts = {b.strip().lower() for b in existing_bullets if b.strip()}
            new_bullets_to_add = []
            for bullet in filtered_bullets:
                bullet_text = bullet.strip() if isinstance(bullet, str) else str(bullet).strip()
                bullet_lower = bullet_text.lower()

                # Skip empty bullets and exact duplicates (case-insensitive)
                if bullet_text and bullet_lower not in existing_bullet_texts:
                    # Check for near-duplicates (similarity > 90%)
                    is_duplicate = False
                    for existing_text in existing_bullet_texts:
                        # Simple similarity check - if one contains the other or vice versa with high overlap
                        if (bullet_lower in existing_text or existing_text in bullet_lower) and \
                           abs(len(bullet_lower) - len(existing_text)) < max(len(bullet_lower), len(existing_text)) * 0.2:
                            is_duplicate = True
                            break

                    if not is_duplicate:
                        existing_bullet_texts.add(bullet_lower)
                        new_bullets_to_add.append(bullet)

            if new_bullets_to_add:
                existing_bullets.extend(new_bullets_to_add)

            # Update section with merged bullets
            existing_section["bullets"] = [
                {
                    "id": f"{existing_section['id']}-{i}",
                    "text": bullet,
                    "params": {},
                }
                for i, bullet in enumerate(existing_bullets)
            ]
            logger.info(f"Merged duplicate section '{section_title}' (normalized: '{normalized_title}') with {len(new_bullets_to_add)} new bullets into existing section '{existing_section['title']}'")
        else:
            # Create new section - remove duplicates within the bullets list itself
            unique_bullets = []
            seen_bullet_texts = set()
            for bullet in filtered_bullets:
                bullet_text = bullet.strip() if isinstance(bullet, str) else str(bullet).strip()
                bullet_lower = bullet_text.lower()
                if bullet_text and bullet_lower not in seen_bullet_texts:
                    seen_bullet_texts.add(bullet_lower)
                    unique_bullets.append(bullet)

            if unique_bullets:
                new_section = {
                    "title": section_title,
                    "bullets": [
                        {
                            "id": f"{len(sections)}-{i}",
                            "text": bullet,
                            "params": {},
                        }
                        for i, bullet in enumerate(unique_bullets)
                    ],
                    "id": str(len(sections)),
                }
                section_dict[normalized_title] = new_section
                sections.append(new_section)

    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        line_original = line.strip()

        # More strict section header detection
        # A section header should:
        # 1. Be relatively short (section titles are usually < 50 chars)
        # 2. Not contain bullet points
        # 3. Match a section keyword at word boundaries
        # 4. Either be all caps, title case, or standalone
        # 5. Not be part of a longer sentence

        is_section = False
        # Strict section header detection - must be a standalone header line
        if (len(line_original) < 50 and
            not line_original.startswith(("•", "-", "*", "·")) and
            not any(char.isdigit() for char in line_original[:10]) and  # Not a date
            " / " not in line_original):  # Not a job entry

            # Check if line matches section keywords at word boundaries
            for keyword in section_keywords:
                # Use word boundary matching to avoid false positives
                pattern = r'\b' + re.escape(keyword) + r'\b'
                if re.search(pattern, line_lower):
                    # Additional checks to ensure it's a header, not content
                    words = line_lower.split()
                    keyword_words = keyword.split()

                    # Must start with keyword or be mostly the keyword
                    starts_with_keyword = line_lower.startswith(keyword)
                    is_mostly_keyword = (
                        len(words) <= len(keyword_words) + 2 and  # Max 2 extra words
                        all(w in keyword_words or len(w) < 5 for w in words)  # Short words only
                    )

                    if starts_with_keyword or is_mostly_keyword:
                        # Check if next line exists and is not empty (content should follow)
                        # Also check previous line - if it was content, this is likely a new section
                        prev_was_content = (
                            i > 0 and
                            current_section and
                            current_bullets and
                            len(current_bullets) > 0
                        )
                        has_next_content = i + 1 < len(lines) and lines[i + 1].strip()

                        if (has_next_content or prev_was_content):
                            is_section = True
                            break

        if is_section:
            # Save previous section before starting new one
            if current_section and current_bullets:
                save_section(current_section, current_bullets)

            # Start new section - use original line but normalize title
            current_section = line_original
            current_bullets = []

        elif current_section and line_original:
            # This is content for current section
            # Skip lines that look like section headers (to prevent false section starts)
            line_lower_check = line_lower
            looks_like_header = (
                len(line_original) < 50 and
                not line_original.startswith(("•", "-", "*", "·")) and
                any(keyword in line_lower_check for keyword in section_keywords) and
                i + 1 < len(lines) and lines[i + 1].strip()  # Has content after
            )

            if not looks_like_header:
                if line_original.startswith(("•", "-", "*", "·")):
                    bullet_text = line_original[1:].strip()
                    if bullet_text:
                        current_bullets.append(f"• {bullet_text}")
                elif " / " in line_original and any(char.isdigit() for char in line_original):
                    # Job entry
                    current_bullets.append(f"**{line_original}**")
                else:
                    current_bullets.append(line_original)

    # Save last section
    if current_section and current_bullets:
        save_section(current_section, current_bullets)

    return {
        "name": name,
        "title": title,
        "email": email,
        "phone": phone,
        "location": location,
        "summary": "",
        "sections": sections,
        "detected_variables": {},
    }
