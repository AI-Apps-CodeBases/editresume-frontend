"""Resume export service - PDF and DOCX generation"""

from __future__ import annotations

import logging
import re
from datetime import datetime
from io import BytesIO
from typing import Optional

from fastapi import HTTPException
from fastapi.responses import Response
from sqlalchemy.orm import Session

from app.api.models import ExportPayload
from app.models import ExportAnalytics, Resume, User
from app.utils.resume_formatting import (
    apply_replacements,
    format_regular_bullets,
    format_work_experience_bullets,
    strip_bullet_markers,
)
from app.utils.resume_templates import TEMPLATES

logger = logging.getLogger(__name__)


def replace_cover_letter_placeholders(content: str, payload: ExportPayload) -> str:
    """Replace common cover letter placeholders with actual values"""
    if not content:
        return content
    
    # Remove email, phone, and date placeholders completely
    placeholders_to_remove = [
        '[Your Email]', '[Your email]', '[your email]', '[YOUR EMAIL]',
        '[Your Phone Number]', '[Your Phone]', '[Your phone]', '[your phone]', '[YOUR PHONE]',
        '[Date]', '[date]', '[DATE]',
        '{{email}}', '{{phone}}', '{{date}}',
        '[Your Phone Number]', '[Phone Number]', '[Phone]'
    ]
    
    result = content
    for placeholder in placeholders_to_remove:
        result = result.replace(placeholder, '')
    
    # Replace name and location placeholders
    replacements_map = {
        '[Your Name]': payload.name or '',
        '[Your Address]': payload.location or '',
        '[City, State, Zip]': payload.location or '',
        '{{name}}': payload.name or '',
        '{{location}}': payload.location or '',
        '{{address}}': payload.location or '',
    }
    
    replacements_map.update({
        '[Your name]': payload.name or '',
        '[your name]': payload.name or '',
        '[YOUR NAME]': payload.name or '',
    })
    
    for placeholder, value in replacements_map.items():
        result = result.replace(placeholder, value)
    
    # Apply any additional replacements from payload.replacements
    if payload.replacements:
        result = apply_replacements(result, payload.replacements)
    
    # Clean up multiple spaces and empty lines
    result = re.sub(r'\n\s*\n\s*\n+', '\n\n', result)
    result = re.sub(r' +', ' ', result)
    result = result.strip()
    
    return result


async def export_pdf(
    payload: ExportPayload,
    user_email: Optional[str] = None,
    session_id: Optional[str] = None,
    db: Optional[Session] = None,
) -> Response:
    """
    Export resume as PDF.
    
    CRITICAL VISIBILITY FILTERING:
    - Sections with params.visible === False are excluded
    - Bullets with params.visible === False are excluded (via formatting functions)
    - Fields with fieldsVisible[field] === False are excluded (summary, name, title, etc.)
    - Company headers with visible === False hide all their associated bullets
    
    DO NOT remove these filters - they are essential for respecting user visibility settings.
    """
    """Export resume as PDF"""
    try:
        # Validate payload
        if not payload:
            raise ValueError("Export payload is required")
        
        # Check if we have at least some content to export
        has_content = bool(
            payload.name or 
            payload.summary or 
            (payload.sections and len(payload.sections) > 0) or
            payload.cover_letter
        )
        if not has_content:
            raise ValueError("Resume must have at least a name, summary, sections, or cover letter")
        
        import weasyprint
        from weasyprint import HTML
        
        weasyprint_version = weasyprint.__version__
        logger.info(f"WeasyPrint version: {weasyprint_version}")
        
        if weasyprint_version.startswith("62."):
            logger.error(
                f"WeasyPrint {weasyprint_version} has known bugs. Please upgrade to 63.0+ using: pip install --upgrade 'weasyprint>=63.0'"
            )

        replacements = payload.replacements or {}
        template_id = payload.template or "tech"
        template_style = TEMPLATES.get(template_id, TEMPLATES["tech"])
        template_config = payload.templateConfig if payload.templateConfig is not None else {}

        logger.info(f"Exporting PDF with template: {template_id}")
        logger.info(f"Template config provided: {bool(template_config)}")
        logger.info(f"Template config type: {type(template_config)}")
        logger.info(f"Template config keys: {list(template_config.keys()) if isinstance(template_config, dict) else 'N/A'}")
        if isinstance(template_config, dict) and template_config.get("typography"):
            typo = template_config.get("typography", {})
            font_sizes = typo.get("fontSize", {})
            logger.info(f"Font sizes in config: h1={font_sizes.get('h1')}, h2={font_sizes.get('h2')}, body={font_sizes.get('body')}")
        logger.info(
            f"Two-column settings: left={payload.two_column_left}, right={payload.two_column_right}, width={payload.two_column_left_width}"
        )

        # Use templateConfig if available and not empty, otherwise fall back to template lookup
        if template_config and isinstance(template_config, dict) and len(template_config) > 0:
            # Typography
            typography = template_config.get("typography", {})
            font_family_heading = typography.get("fontFamily", {}).get("heading", "Arial")
            font_family_body = typography.get("fontFamily", {}).get("body", "Arial")
            font_family = f"{font_family_heading}, {font_family_body}, sans-serif"
            
            # Design
            design = template_config.get("design", {})
            header_style = design.get("headerStyle", "left-aligned")
            # Match frontend logic: center if headerStyle is 'centered' OR template is one of: clean, two-column, compact, classic
            templates_that_center = ["clean", "two-column", "compact", "classic"]
            header_align = "center" if (header_style == "centered" or template_id in templates_that_center) else "left"
            header_border = "2px solid #000" if design.get("dividers", True) else "none"
            section_uppercase = "text-transform: uppercase;" if design.get("sectionUppercase", False) else ""
            
            # Bullet style
            bullet_style = design.get("bulletStyle", "circle")
            bullet_symbols = {
                "circle": "•",
                "square": "■",
                "dash": "—",
                "none": ""
            }
            bullet_symbol = bullet_symbols.get(bullet_style, "•")
            
            # Layout
            layout_config = template_config.get("layout", {})
            layout_columns = layout_config.get("columns", "single")
            layout = "two-column" if layout_columns in ["two-column", "asymmetric"] else "single"
            # Get column width from templateConfig if available
            template_column_width = layout_config.get("columnWidth")
            logger.info(f"Layout config: columns={layout_columns}, columnWidth from templateConfig={template_column_width}")
            
            # Colors
            colors = design.get("colors", {})
            primary_color = colors.get("primary", "#000000")
            secondary_color = colors.get("secondary", "#000000")
            accent_color = colors.get("accent", "#000000")
            text_color = colors.get("text", "#000000")
            
            # Check if we need gradient backgrounds (for new templates like vibrant, gradient)
            use_gradient_header = template_id in ["vibrant", "gradient", "creative"]
            gradient_css = ""
            if use_gradient_header and primary_color and accent_color:
                gradient_css = f"background: linear-gradient(135deg, {primary_color} 0%, {accent_color} 100%);"
            
            # Spacing
            spacing_config = template_config.get("spacing", {})
            section_gap = spacing_config.get("sectionGap", 15)
            item_gap = spacing_config.get("itemGap", 6)
            page_margin = spacing_config.get("pageMargin", 24)
            
            # Font sizes
            h1_size = typography.get("fontSize", {}).get("h1", 24)
            h2_size = typography.get("fontSize", {}).get("h2", 12)
            body_size = typography.get("fontSize", {}).get("body", 11)
            line_height = typography.get("lineHeight", 1.4)
            
            logger.info(f"Using templateConfig values:")
            logger.info(f"  Font sizes: h1={h1_size}px, h2={h2_size}px, body={body_size}px, lineHeight={line_height}")
            logger.info(f"  Spacing: sectionGap={section_gap}px, itemGap={item_gap}px, pageMargin={page_margin}px")
            logger.info(f"  Colors: primary={primary_color}, text={text_color}")
            logger.info(f"  Layout: columns={layout_columns}, columnWidth={template_column_width}")
        else:
            # Fallback to template lookup
            font_family = template_style["styles"]["font"]
            header_align = template_style["styles"]["header_align"]
            header_border = template_style["styles"]["header_border"]
            section_uppercase = (
                "text-transform: uppercase;"
                if template_style["styles"]["section_uppercase"]
                else ""
            )
            layout = template_style["styles"]["layout"]
            primary_color = "#000000"
            secondary_color = "#000000"
            accent_color = "#000000"
            text_color = "#000000"
            section_gap = 15
            item_gap = 6
            h1_size = 24
            h2_size = 12
            body_size = 11
            line_height = 1.4
            # Default bullet style for fallback
            bullet_style = "circle"
            bullet_symbols = {
                "circle": "•",
                "square": "■",
                "dash": "—",
                "none": ""
            }
            bullet_symbol = bullet_symbols.get(bullet_style, "•")
            use_gradient_header = False
            gradient_css = ""
            
            logger.info(f"Using fallback font sizes: h1={h1_size}px, h2={h2_size}px, body={body_size}px, lineHeight={line_height}")

        # Build HTML content sections
        # Handle different contact info formats based on template
        contact_items_html = ""  # For infographic template
        if template_id == "corporate-premium":
            # Corporate premium uses structured contact in sidebar
            contact_info = ""  # Will be handled in sidebar
        elif template_id == "infographic":
            # Infographic uses icons - will be handled in header
            contact_info = ""
            if payload.email:
                contact_items_html += f'<div style="display: flex; align-items: center; gap: 8px; font-size: {body_size}px; color: {text_color};"><span style="font-size: 16px;">📧</span><span>{apply_replacements(payload.email, replacements)}</span></div>'
            if payload.phone:
                contact_items_html += f'<div style="display: flex; align-items: center; gap: 8px; font-size: {body_size}px; color: {text_color};"><span style="font-size: 16px;">📱</span><span>{apply_replacements(payload.phone, replacements)}</span></div>'
            if payload.location:
                contact_items_html += f'<div style="display: flex; align-items: center; gap: 8px; font-size: {body_size}px; color: {text_color};"><span style="font-size: 16px;">📍</span><span>{apply_replacements(payload.location, replacements)}</span></div>'
        elif template_id == "timeline":
            # Timeline uses simple contact string
            contact_items = []
            if payload.email:
                contact_items.append(apply_replacements(payload.email, replacements))
            if payload.phone:
                contact_items.append(apply_replacements(payload.phone, replacements))
            if payload.location:
                contact_items.append(apply_replacements(payload.location, replacements))
            contact_info = " • ".join(contact_items)
        else:
            # Default: simple contact info string
            contact_info = apply_replacements(payload.email or "", replacements)
            if payload.phone:
                contact_info += " • " + apply_replacements(payload.phone, replacements)
            if payload.location:
                contact_info += " • " + apply_replacements(payload.location, replacements)

        # Check if this is a cover letter only export (no resume sections)
        is_cover_letter_only = bool(payload.cover_letter and not payload.sections and not payload.summary)
        
        # Add cover letter if provided
        cover_letter_html = ""
        if payload.cover_letter:
            # Replace placeholders with actual values
            cover_letter_content = replace_cover_letter_placeholders(payload.cover_letter, payload)
            
            # Parse cover letter structure
            lines = cover_letter_content.split('\n')
            
            # Extract title (first non-empty line, or construct from company/position)
            title = ""
            if payload.company_name and payload.position_title:
                title = f"{payload.position_title} at {payload.company_name}"
            elif payload.company_name:
                title = f"{payload.company_name} - Cover Letter"
            else:
                # Try to extract from content
                for line in lines[:3]:
                    if line.strip() and len(line.strip()) < 100:
                        title = line.strip()
                        break
                if not title:
                    title = "Cover Letter"
            
            # Build structured letter HTML
            letter_parts = []
            
            # Title
            letter_parts.append(f'<h1 class="cover-letter-title" style="font-size: 24px; font-weight: bold; margin-bottom: 24px; text-align: center;">{title}</h1>')
            
            # Parse content sections
            current_section = []
            in_body = False
            
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    if current_section:
                        section_text = ' '.join(current_section)
                        if section_text:
                            # Check if it's a date
                            if i < 5 and any(month in section_text for month in ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']):
                                letter_parts.append(f'<p style="margin-bottom: 8px; font-size: {body_size}px;">{section_text}</p>')
                            # Check if it's a salutation
                            elif section_text.startswith('Dear'):
                                letter_parts.append(f'<p style="margin-bottom: 12px; font-size: {body_size}px;">{section_text}</p>')
                            # Check if it's a closing
                            elif section_text in ['Sincerely,', 'Best regards,', 'Yours sincerely,']:
                                letter_parts.append(f'<p style="margin-top: 24px; margin-bottom: 8px; font-size: {body_size}px;">{section_text}</p>')
                            # Check if it's signature (typically after closing)
                            elif i > len(lines) - 3:
                                letter_parts.append(f'<p style="margin-bottom: 8px; font-size: {body_size}px;">{section_text}</p>')
                            # Body paragraph
                            else:
                                letter_parts.append(f'<p style="margin-bottom: 12px; font-size: {body_size}px; line-height: 1.6; text-align: justify;">{section_text}</p>')
                        current_section = []
                    continue
                
                # Check for title (skip if it's the title line)
                if line == title:
                    continue
                
                # Check for contact info patterns
                if '@' in line or any(char.isdigit() and len(line.replace(' ', '').replace('-', '').replace('(', '').replace(')', '')) >= 10 for char in line):
                    letter_parts.append(f'<p style="margin-bottom: 4px; font-size: {body_size - 1}px;">{line}</p>')
                    continue
                
                current_section.append(line)
            
            # Handle remaining section
            if current_section:
                section_text = ' '.join(current_section)
                if section_text:
                    letter_parts.append(f'<p style="margin-bottom: 12px; font-size: {body_size}px; line-height: 1.6; text-align: justify;">{section_text}</p>')
            
            # Fallback: if no structured parsing worked, use original method
            if len(letter_parts) <= 1:  # Only title
                paragraphs = []
                for para in cover_letter_content.split('\n\n'):
                    para = para.strip()
                    if para:
                        para = para.replace('\n', ' ').strip()
                        para = re.sub(r' +', ' ', para)
                        if para:
                            paragraphs.append(f'<p style="margin-bottom: 12px; font-size: {body_size}px; line-height: 1.6; text-align: justify;">{para}</p>')
                letter_parts.extend(paragraphs)
            
            cover_letter_body = '\n                '.join(letter_parts[1:]) if len(letter_parts) > 1 else ''
            
            cover_letter_html = f"""
            <div class="cover-letter-section" style="font-family: {font_family}; max-width: 800px; margin: 0 auto; padding: 40px;">
                {letter_parts[0] if letter_parts else ''}
                <div class="cover-letter-content" style="margin-top: 20px;">
                    {cover_letter_body}
                </div>
            </div>
            """

        # Determine layout from templateConfig or fallback
        # Special handling for corporate-premium (sidebar layout)
        is_corporate_premium = template_id == "corporate-premium"
        is_timeline = template_id == "timeline"
        is_infographic = template_id == "infographic"
        
        is_two_column = layout == "two-column" or (
            template_config and 
            template_config.get("layout", {}).get("columns") in ["two-column", "asymmetric"]
        ) or is_corporate_premium
        
        # Build sections HTML
        # Special handling for corporate-premium, timeline, and infographic templates
        if is_corporate_premium:
            # Corporate Premium: Sidebar layout with structured contact info
            sidebar_width = template_config.get("layout", {}).get("columnWidth", 30) if template_config else 30
            main_width = 100 - sidebar_width
            
            # Build contact items for sidebar
            contact_items = []
            if payload.email:
                contact_items.append(f'<div style="margin-bottom: 12px;"><div style="font-size: {body_size - 1}px; opacity: 0.8; margin-bottom: 4px; text-transform: uppercase; letter-spacing: 0.5px; font-weight: 600;">Email</div><div style="font-size: {body_size}px; font-weight: 400;">{apply_replacements(payload.email, replacements)}</div></div>')
            if payload.phone:
                contact_items.append(f'<div style="margin-bottom: 12px;"><div style="font-size: {body_size - 1}px; opacity: 0.8; margin-bottom: 4px; text-transform: uppercase; letter-spacing: 0.5px; font-weight: 600;">Phone</div><div style="font-size: {body_size}px; font-weight: 400;">{apply_replacements(payload.phone, replacements)}</div></div>')
            if payload.location:
                contact_items.append(f'<div style="margin-bottom: 12px;"><div style="font-size: {body_size - 1}px; opacity: 0.8; margin-bottom: 4px; text-transform: uppercase; letter-spacing: 0.5px; font-weight: 600;">Location</div><div style="font-size: {body_size}px; font-weight: 400;">{apply_replacements(payload.location, replacements)}</div></div>')
            
            sidebar_html = f"""
            <div style="width: {sidebar_width}%; background: {primary_color}; color: white; padding: 32px 24px; border-radius: 8px; min-height: 200px; box-sizing: border-box;">
                <h1 style="font-family: {font_family}; font-size: {h1_size}px; font-weight: bold; margin-bottom: 8px; letter-spacing: 0.5px; color: white;">{apply_replacements(payload.name, replacements)}</h1>
                {f'<p style="font-size: {body_size + 2}px; opacity: 0.9; margin-bottom: 24px; font-weight: 300;">{apply_replacements(payload.title, replacements)}</p>' if payload.title else ''}
                <div style="border-top: 1px solid rgba(255,255,255,0.2); padding-top: 20px; margin-top: 20px;">
                    {''.join(contact_items)}
                </div>
            </div>"""
            
            # Build main content sections
            main_sections_html = ""
            # CRITICAL: Check fieldsVisible for summary - only show if not explicitly hidden
            if payload.summary and (not payload.fieldsVisible or payload.fieldsVisible.get("summary") is not False):
                main_sections_html += f"""
                <section style="margin-bottom: {section_gap}px;">
                    <h2 style="font-family: {font_family}; font-size: {h2_size}px; font-weight: bold; color: {primary_color}; border-bottom: 2px solid {accent_color}; padding-bottom: 6px; display: inline-block; margin-bottom: 12px;">Professional Summary</h2>
                    <p style="font-size: {body_size}px; line-height: {line_height}; color: {text_color}; margin-top: 12px;">{apply_replacements(payload.summary, replacements)}</p>
                </section>"""
            
            # CRITICAL: Filter sections by visible property - skip sections with visible === False
            for s in payload.sections:
                if s.params and s.params.get("visible") is False:
                    continue
                section_title = s.title.lower()
                is_work_experience = "experience" in section_title or "work" in section_title or "employment" in section_title
                
                if is_work_experience:
                    bullets_html = format_work_experience_bullets(s.bullets, replacements)
                else:
                    bullets_html = format_regular_bullets(s.bullets, replacements, s.title)
                
                main_sections_html += f"""
                <section style="margin-bottom: {section_gap}px;">
                    <h2 style="font-family: {font_family}; font-size: {h2_size}px; font-weight: bold; color: {primary_color}; border-bottom: 2px solid {accent_color}; padding-bottom: 6px; display: inline-block; margin-bottom: 12px;">{apply_replacements(s.title, replacements)}</h2>
                    {bullets_html}
                </section>"""
            
            content_html = f"""
            {cover_letter_html}
            <div style="display: flex; gap: 24px; margin-bottom: 24px;">
                {sidebar_html}
                <div style="width: {main_width}%;">
                    {main_sections_html}
                </div>
            </div>
            <div style="width: 100%;">
                {main_sections_html if not payload.summary else ''}
            </div>"""
            
        elif is_timeline:
            # Timeline: Timeline markers and lines
            sections_html = ""
            # CRITICAL: Check fieldsVisible for summary
            if payload.summary and (not payload.fieldsVisible or payload.fieldsVisible.get("summary") is not False):
                sections_html += f"""
                <section style="margin-bottom: {section_gap}px;">
                    <div style="padding-left: 24px; border-left: 3px solid {accent_color}; margin-left: 8px;">
                        <p style="font-size: {body_size + 1}px; line-height: {line_height}; color: {text_color}; font-style: italic;">{apply_replacements(payload.summary, replacements)}</p>
                    </div>
                </section>"""
            
            # CRITICAL: Filter sections by visible property
            for s in payload.sections:
                if s.params and s.params.get("visible") is False:
                    continue
                section_title = s.title.lower()
                is_work_experience = "experience" in section_title or "work" in section_title or "employment" in section_title
                
                if is_work_experience:
                    bullets_html = format_work_experience_bullets(s.bullets, replacements)
                else:
                    bullets_html = format_regular_bullets(s.bullets, replacements, s.title)
                
                sections_html += f"""
                <section style="margin-bottom: {section_gap}px; position: relative; padding-left: 32px;">
                    <div style="position: absolute; left: 8px; top: 0; bottom: -24px; width: 3px; background: {primary_color}; border-radius: 2px;"></div>
                    <div style="position: absolute; left: 0; top: 4px; width: 16px; height: 16px; border-radius: 50%; background: {primary_color}; border: 3px solid white; box-shadow: 0 0 0 2px {primary_color};"></div>
                    <h2 style="font-family: {font_family}; font-size: {h2_size}px; font-weight: bold; color: {primary_color}; margin-bottom: 12px;">{apply_replacements(s.title, replacements)}</h2>
                    <div style="padding-left: 8px;">
                        {bullets_html}
                    </div>
                </section>"""
            
            content_html = cover_letter_html + sections_html
            
        elif is_infographic:
            # Infographic: Card-based sections with icons
            sections_html = ""
            # CRITICAL: Check fieldsVisible for summary
            if payload.summary and (not payload.fieldsVisible or payload.fieldsVisible.get("summary") is not False):
                sections_html += f"""
                <section style="margin-bottom: {section_gap}px;">
                    <div style="background: {accent_color}10; border: 2px solid {accent_color}30; border-radius: 12px; padding: 20px;">
                        <div style="display: flex; align-items: start; gap: 12px;">
                            <span style="font-size: 24px; line-height: 1;">💡</span>
                            <p style="font-size: {body_size + 1}px; line-height: {line_height}; color: {text_color}; margin: 0; flex: 1;">{apply_replacements(payload.summary, replacements)}</p>
                        </div>
                    </div>
                </section>"""
            
            section_icons = {
                'experience': '💼', 'education': '🎓', 'skills': '⚡', 'projects': '🚀',
                'certifications': '🏆', 'awards': '⭐', 'languages': '🌐', 'publications': '📚'
            }
            
            # CRITICAL: Filter sections by visible property
            for s in payload.sections:
                if s.params and s.params.get("visible") is False:
                    continue
                section_title_lower = s.title.lower()
                icon = '📋'
                for key, icon_char in section_icons.items():
                    if key in section_title_lower:
                        icon = icon_char
                        break
                
                section_title = s.title.lower()
                is_work_experience = "experience" in section_title or "work" in section_title or "employment" in section_title
                
                if is_work_experience:
                    bullets_html = format_work_experience_bullets(s.bullets, replacements)
                else:
                    bullets_html = format_regular_bullets(s.bullets, replacements, s.title)
                
                sections_html += f"""
                <section style="margin-bottom: {section_gap}px; padding: 20px; background: {primary_color}05; border-radius: 12px; border: 1px solid {primary_color}20;">
                    <div style="display: flex; align-items: center; gap: 12px; margin-bottom: 16px;">
                        <div style="width: 48px; height: 48px; border-radius: 12px; background: linear-gradient(135deg, {primary_color} 0%, {accent_color} 100%); display: flex; align-items: center; justify-content: center; font-size: 24px; box-shadow: 0 2px 8px {primary_color}30;">{icon}</div>
                        <h2 style="font-family: {font_family}; font-size: {h2_size}px; font-weight: bold; color: {primary_color}; margin: 0; letter-spacing: 0.3px;">{apply_replacements(s.title, replacements)}</h2>
                    </div>
                    <div style="padding-left: 8px;">
                        {bullets_html}
                    </div>
                </section>"""
            
            content_html = cover_letter_html + sections_html
            
        elif is_two_column:
            left_sections_html = ""
            right_sections_html = ""

            # Use localStorage configuration if provided, otherwise use smart default distribution
            left_section_ids = set(payload.two_column_left or [])
            right_section_ids = set(payload.two_column_right or [])
            
            # Check if summary should be in left or right column
            summary_id = "__summary__"
            summary_in_left = summary_id in left_section_ids
            summary_in_right = summary_id in right_section_ids
            
            # CRITICAL: Add summary to appropriate column if it exists, is assigned, and not hidden
            if payload.summary and (not payload.fieldsVisible or payload.fieldsVisible.get("summary") is not False):
                if (summary_in_left or summary_in_right):
                    summary_html = f"""
                    <div class="section">
                        <h2>Professional Summary</h2>
                        <div class="summary">{apply_replacements(payload.summary, replacements)}</div>
                    </div>"""
                    if summary_in_left:
                        left_sections_html += summary_html
                    elif summary_in_right:
                        right_sections_html += summary_html
                elif not left_section_ids and not right_section_ids:
                    # Default: add summary to left column if no configuration exists
                    summary_html = f"""
                    <div class="section">
                        <h2>Professional Summary</h2>
                        <div class="summary">{apply_replacements(payload.summary, replacements)}</div>
                    </div>"""
                    left_sections_html += summary_html

            # If no configuration provided, use smart distribution: Skills, Certificates, Education on left
            if not left_section_ids and not right_section_ids:
                # Default distribution: Skills, Certificates, Education on left; everything else on right
                left_column_keywords = ['skill', 'certificate', 'certification', 'education', 'academic', 'qualification', 'award', 'honor']
                
                # CRITICAL: Filter sections by visible property
                for s in payload.sections:
                    if s.params and s.params.get("visible") is False:
                        continue
                    section_title_lower = s.title.lower()
                    # Skip summary section as it's already added above
                    if 'summary' in section_title_lower or 'professional summary' in section_title_lower:
                        continue
                    
                    is_left_column = any(keyword in section_title_lower for keyword in left_column_keywords)
                    
                    section_title = s.title.lower()
                    is_work_experience = (
                        "experience" in section_title
                        or "work" in section_title
                        or "employment" in section_title
                    )

                    # Determine section header style based on template (for two-column layout)
                    section_header_class = ""
                    section_header_style = ""
                    if template_id == "vibrant":
                        section_colors = ["#3b82f6", "#8b5cf6", "#f59e0b", "#10b981", "#ef4444", "#06b6d4"]
                        section_idx = len([x for x in payload.sections if payload.sections.index(x) < payload.sections.index(s)])
                        section_color = section_colors[section_idx % len(section_colors)]
                        section_header_class = "section-header-vibrant"
                        section_header_style = f"background: {section_color}; color: white; padding: 8px 12px; border-radius: 4px; margin-bottom: 12px;"
                    
                    if is_work_experience:
                        bullets_html = format_work_experience_bullets(
                            s.bullets, replacements
                        )
                        if section_header_style:
                            section_html = f"""
                        <div class="section">
                            <div class="{section_header_class}" style="{section_header_style}">
                                <h2 style="margin: 0; font-size: {h2_size}px; font-weight: bold;">{apply_replacements(s.title, replacements)}</h2>
                            </div>
                            {bullets_html}
                        </div>"""
                        else:
                            section_html = f"""
                        <div class="section">
                            <h2>{apply_replacements(s.title, replacements)}</h2>
                            {bullets_html}
                        </div>"""
                    else:
                        bullets_html = format_regular_bullets(s.bullets, replacements, s.title)
                        section_title_formatted = apply_replacements(s.title, replacements)
                        section_lower = s.title.lower()
                        is_skill_section = (
                            "skill" in section_lower
                            or "technical" in section_lower
                            or "technology" in section_lower
                            or "competencies" in section_lower
                            or "expertise" in section_lower
                            or "proficiencies" in section_lower
                        )
                        if section_header_style:
                            section_html = f"""
                            <div class="section">
                                <div class="{section_header_class}" style="{section_header_style}">
                                    <h2 style="margin: 0; font-size: {h2_size}px; font-weight: bold;">{section_title_formatted}</h2>
                                </div>
                                {'<ul>' if not is_skill_section else ''}
                                {bullets_html}
                                {'</ul>' if not is_skill_section else ''}
                            </div>"""
                        elif is_skill_section:
                            section_html = f"""
                            <div class="section">
                                <h2>{section_title_formatted}</h2>
                                {bullets_html}
                            </div>"""
                        else:
                            section_html = f"""
                            <div class="section">
                                <h2>{section_title_formatted}</h2>
                                <ul>
                                    {bullets_html}
                                </ul>
                            </div>"""

                    if is_left_column:
                        left_sections_html += section_html
                    else:
                        right_sections_html += section_html
            else:
                # Use provided configuration
                # CRITICAL: Filter sections by visible property
                for s in payload.sections:
                    if s.params and s.params.get("visible") is False:
                        continue
                    section_title = s.title.lower()
                    # Skip summary section as it's already added above
                    if 'summary' in section_title or 'professional summary' in section_title:
                        continue
                    
                    is_work_experience = (
                        "experience" in section_title
                        or "work" in section_title
                        or "employment" in section_title
                    )

                    if is_work_experience:
                        bullets_html = format_work_experience_bullets(
                            s.bullets, replacements
                        )
                        section_html = f"""
                        <div class="section">
                            <h2>{apply_replacements(s.title, replacements)}</h2>
                            {bullets_html}
                        </div>"""
                    else:
                        bullets_html = format_regular_bullets(s.bullets, replacements, s.title)
                        section_title = apply_replacements(s.title, replacements)
                        section_lower = s.title.lower()
                        is_skill_section = (
                            "skill" in section_lower
                            or "technical" in section_lower
                            or "technology" in section_lower
                            or "competencies" in section_lower
                            or "expertise" in section_lower
                            or "proficiencies" in section_lower
                        )
                        if is_skill_section:
                            section_html = f"""
                            <div class="section">
                                <h2>{section_title}</h2>
                                {bullets_html}
                            </div>"""
                        else:
                            section_html = f"""
                            <div class="section">
                                <h2>{section_title}</h2>
                                <ul>
                                    {bullets_html}
                                </ul>
                            </div>"""

                    if s.id and s.id in left_section_ids:
                        left_sections_html += section_html
                    elif s.id and s.id in right_section_ids:
                        right_sections_html += section_html

            # Calculate column widths with spacing
            # Use templateConfig columnWidth if available, otherwise use payload value, otherwise default to 50
            if template_config and isinstance(template_config, dict):
                layout_config = template_config.get("layout", {})
                template_column_width = layout_config.get("columnWidth")
                if template_column_width is not None:
                    left_width_percent = template_column_width
                    logger.info(f"Using column width from templateConfig: {left_width_percent}%")
                else:
                    left_width_percent = payload.two_column_left_width or 50
                    logger.info(f"templateConfig columnWidth not found, using payload value: {left_width_percent}%")
            else:
                left_width_percent = payload.two_column_left_width or 50
                logger.info(f"No templateConfig, using payload value: {left_width_percent}%")
            
            gap_percent = 1.5
            # Adjust right width to account for gap
            right_width_percent = 100 - left_width_percent - gap_percent
            logger.info(f"Final column widths: left={left_width_percent}%, right={right_width_percent}% (gap={gap_percent}%)")

            content_html = f"""
            {cover_letter_html}
            <table class="two-column" style="width: 100%; border-collapse: collapse; table-layout: fixed;">
                <tr>
                    <td class="column" style="width: {left_width_percent}%; padding-right: {gap_percent}%; vertical-align: top; overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; box-sizing: border-box;">{left_sections_html}</td>
                    <td class="column" style="width: {right_width_percent}%; vertical-align: top; overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; box-sizing: border-box;">{right_sections_html}</td>
                </tr>
            </table>"""
        else:
            # Single column layout
            # CRITICAL: Check fieldsVisible for summary
            summary_html = (
                f'<div class="summary">{apply_replacements(payload.summary, replacements)}</div>'
                if payload.summary and (not payload.fieldsVisible or payload.fieldsVisible.get("summary") is not False)
                else ""
            )
            sections_html = ""
            # CRITICAL: Filter sections by visible property
            for s in payload.sections:
                if s.params and s.params.get("visible") is False:
                    continue
                section_title = s.title.lower()
                is_work_experience = (
                    "experience" in section_title
                    or "work" in section_title
                    or "employment" in section_title
                )

                # Determine section header style based on template
                section_header_class = ""
                section_header_style = ""
                if template_id == "vibrant":
                    # Vibrant template uses colored block headers
                    section_colors = ["#3b82f6", "#8b5cf6", "#f59e0b", "#10b981", "#ef4444", "#06b6d4"]
                    section_idx = len([x for x in payload.sections if payload.sections.index(x) < payload.sections.index(s)])
                    section_color = section_colors[section_idx % len(section_colors)]
                    section_header_class = "section-header-vibrant"
                    section_header_style = f"background: {section_color}; color: white; padding: 8px 12px; border-radius: 4px; margin-bottom: 12px;"
                elif template_id == "gradient":
                    # Gradient template uses gradient dividers
                    section_header_class = ""
                    section_header_style = ""
                
                if is_work_experience:
                    bullets_html = format_work_experience_bullets(
                        s.bullets, replacements
                    )
                    if section_header_style:
                        sections_html += f"""
                    <div class="section">
                        <div class="{section_header_class}" style="{section_header_style}">
                            <h2 style="margin: 0; font-size: {h2_size}px; font-weight: bold;">{apply_replacements(s.title, replacements)}</h2>
                        </div>
                        {bullets_html}
                    </div>"""
                    else:
                        sections_html += f"""
                    <div class="section">
                        <h2>{apply_replacements(s.title, replacements)}</h2>
                        {bullets_html}
                    </div>"""
                else:
                    bullets_html = format_regular_bullets(s.bullets, replacements, s.title)
                    section_title = apply_replacements(s.title, replacements)
                    section_lower = s.title.lower()
                    is_skill_section = (
                        "skill" in section_lower
                        or "technical" in section_lower
                        or "technology" in section_lower
                        or "competencies" in section_lower
                        or "expertise" in section_lower
                        or "proficiencies" in section_lower
                    )
                    if section_header_style:
                        sections_html += f"""
                        <div class="section">
                            <div class="{section_header_class}" style="{section_header_style}">
                                <h2 style="margin: 0; font-size: {h2_size}px; font-weight: bold;">{section_title}</h2>
                            </div>
                            {'<ul>' if not is_skill_section else ''}
                            {bullets_html}
                            {'</ul>' if not is_skill_section else ''}
                        </div>"""
                    elif is_skill_section:
                        sections_html += f"""
                        <div class="section">
                            <h2>{section_title}</h2>
                            {bullets_html}
                        </div>"""
                    else:
                        sections_html += f"""
                        <div class="section">
                            <h2>{section_title}</h2>
                            <ul>
                                {bullets_html}
                            </ul>
                        </div>"""
            content_html = cover_letter_html + summary_html + sections_html

        # Set page margin based on export type
        if is_cover_letter_only:
            # Professional margins for cover letter (2.5cm = ~1 inch)
            page_margin_cm = "2.5cm"
            logger.info(f"Page margin set to 2.5cm for cover letter only export")
        else:
            # Minimal margin for resume to maximize content area
            page_margin_cm = "0.1cm"
            logger.info(f"Page margin set to 0.1cm to fill entire A4 page")
        
        # Build bullet CSS based on bullet style
        bullet_css = ""
        if bullet_style != 'none':
            bullet_css = f'''.section li::before {{ content: "{bullet_symbol}"; font-weight: bold; color: {primary_color}; position: absolute; left: 0; top: 0; }}
        .job-entry li::before {{ content: "{bullet_symbol}"; font-weight: bold; color: {primary_color}; position: absolute; left: 0; top: 0; }}'''
        bullet_padding = "padding-left: 14px;" if bullet_style != 'none' else "padding-left: 0;"
        
        # Build header HTML for special templates
        header_html = ""
        if not is_cover_letter_only:
            if is_corporate_premium:
                header_html = ""  # Corporate premium has header in sidebar
            elif is_timeline:
                title_html = f'<p style="font-size: {body_size + 3}px; color: {secondary_color}; margin-bottom: 16px; font-weight: 500;">{apply_replacements(payload.title, replacements)}</p>' if payload.title else ""
                header_html = f'<header style="border-bottom: 3px solid {primary_color}; padding-bottom: 20px; margin-bottom: 32px;"><h1 style="font-family: {font_family}; font-size: {h1_size}px; font-weight: bold; color: {primary_color}; margin-bottom: 8px; letter-spacing: 0.2px;">{apply_replacements(payload.name, replacements)}</h1>{title_html}<div style="color: {secondary_color}; font-size: {body_size}px;">{contact_info}</div></header>'
            elif is_infographic:
                title_html = f'<p style="font-size: {body_size + 3}px; color: {secondary_color}; font-weight: 500;">{apply_replacements(payload.title, replacements)}</p>' if payload.title else ""
                contact_div = f'<div style="display: flex; flex-wrap: wrap; gap: 16px; margin-top: 16px;">{contact_items_html}</div>' if contact_items_html else ""
                header_html = f'<header style="background: linear-gradient(135deg, {primary_color}15 0%, {accent_color}15 100%); padding: 32px; border-radius: 16px; border: 2px solid {primary_color}30; margin-bottom: 32px;"><div style="display: flex; align-items: center; gap: 16px; margin-bottom: 16px;"><div style="width: 80px; height: 80px; border-radius: 50%; background: linear-gradient(135deg, {primary_color} 0%, {accent_color} 100%); display: flex; align-items: center; justify-content: center; font-size: 36px; box-shadow: 0 4px 12px {primary_color}40;">👤</div><div style="flex: 1;"><h1 style="font-family: {font_family}; font-size: {h1_size}px; font-weight: bold; color: {primary_color}; margin-bottom: 8px; letter-spacing: 0.3px;">{apply_replacements(payload.name, replacements)}</h1>{title_html}</div></div>{contact_div}</header>'
            else:
                header_html = f'<div class="header"><h1>{apply_replacements(payload.name, replacements)}</h1><div class="title">{apply_replacements(payload.title, replacements)}</div><div class="contact">{contact_info}</div></div>'
        
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        @page {{ size: A4; margin: {page_margin_cm}; }}
        body {{ margin: 0; padding: 0 0.1cm 0 0.1cm; width: 100%; font-family: {font_family}; font-size: {body_size}px; line-height: {line_height}; color: {text_color}; {'text-align: justify;' if is_cover_letter_only else ''} box-sizing: border-box; overflow-wrap: break-word; word-wrap: break-word; }}
        .header {{ text-align: {header_align}; border-bottom: {header_border}; padding-bottom: 10px; margin-bottom: {section_gap}px; {gradient_css if use_gradient_header else ''} }}
        .header h1 {{ margin: 0; font-size: {h1_size}px; font-weight: bold; color: {'white' if use_gradient_header else primary_color}; }}
        .header .title {{ font-size: {h2_size + 2}px; margin: 5px 0; color: {'white' if use_gradient_header else text_color}; }}
        .header .contact {{ font-size: {body_size}px; color: {'white' if use_gradient_header else text_color}; margin-top: 5px; }}
        .summary {{ margin-bottom: {section_gap}px; font-size: {body_size}px; line-height: {line_height}; color: {text_color}; }}
        .section {{ margin-bottom: {section_gap}px; }}
        .section h2 {{ font-size: {h2_size}px; font-weight: bold; color: {primary_color}; {section_uppercase}
                       border-bottom: 1px solid {primary_color}; padding-bottom: 3px; margin-bottom: 8px; }}
        /* Support for vibrant template - colored section headers */
        .section-header-vibrant {{ background: {primary_color}; color: white; padding: 8px 12px; border-radius: 4px; margin-bottom: 12px; }}
        /* Support for gradient template - gradient section dividers */
        .section-divider-gradient {{ height: 3px; background: linear-gradient(90deg, {primary_color} 0%, {accent_color} 100%); border-radius: 2px; margin-bottom: 8px; }}
        .section ul {{ margin: 0; padding-left: 0; list-style: none; }}
        .section li {{ margin-bottom: {item_gap}px; font-size: {body_size}px; color: {text_color}; position: relative; {bullet_padding} }}
        .job-entry ul {{ margin: 0; padding-left: 0; list-style: none; }}
        .job-entry li {{ margin-bottom: {item_gap}px; font-size: {body_size}px; color: {text_color}; position: relative; {bullet_padding} }}
        {bullet_css}
        .section li * {{ display: inline; }}
        .section li strong {{ font-weight: bold; }}
        .company-name-line strong {{ font-weight: bold; }}
        .job-title strong {{ font-weight: bold; }}
        .job-entry {{ margin-bottom: 20px; }}
        .company-header {{ margin-bottom: 8px; }}
        .company-name-line {{ font-weight: bold; font-size: 1.1em; color: {primary_color}; margin-bottom: 3px; }}
        .company-title-line {{ display: flex; justify-content: space-between; align-items: center; font-size: {body_size}px; }}
        .job-title {{ font-weight: 500; color: {text_color}; }}
        .job-date {{ color: #666666; font-size: {body_size - 1}px; }}
        .skills-section {{ font-size: {body_size}px; color: {text_color}; line-height: {line_height}; }}
        .job-separator {{ height: 10px; }}
        .two-column {{ 
            width: 100%; 
            border-collapse: collapse;
            table-layout: fixed;
        }}
        .two-column td.column {{
            vertical-align: top;
            padding: 0;
            overflow-wrap: break-word;
            word-wrap: break-word;
            word-break: break-word;
            box-sizing: border-box;
        }}
        .two-column td.column:first-child {{
            padding-right: 1%;
        }}
        .section {{ overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; }}
        .summary {{ overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; }}
        .job-entry {{ overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; }}
        .clearfix {{ clear: both; }}
        .cover-letter-section {{ margin-bottom: 20px; {'page-break-after: always;' if not is_cover_letter_only else ''} }}
        .cover-letter-section h2 {{ 
            font-size: {h2_size + 4 if is_cover_letter_only else h2_size + 2}px; 
            font-weight: bold; 
            margin-bottom: {20 if is_cover_letter_only else 15}px; 
            {'border-bottom: 2px solid ' + primary_color + ';' if not is_cover_letter_only else ''} 
            padding-bottom: {10 if is_cover_letter_only else 5}px; 
            {'color: ' + primary_color + ';' if not is_cover_letter_only else ''} 
        }}
        .cover-letter-title {{
            font-size: {h1_size}px;
            font-weight: bold;
            margin-bottom: 30px;
            margin-top: 0;
            text-align: center;
            color: {primary_color};
        }}
        .cover-letter-content {{ 
            font-size: {body_size + 2 if is_cover_letter_only else body_size + 1}px; 
            line-height: {line_height + 0.3 if is_cover_letter_only else line_height + 0.1}; 
            text-align: justify; 
            margin-top: {20 if is_cover_letter_only else 0}px; 
        }}
        .cover-letter-content p {{
            margin: 0 0 {12 if is_cover_letter_only else 8}px 0;
            padding: 0;
        }}
        .cover-letter-content p:last-child {{
            margin-bottom: 0;
        }}
    </style>
</head>
<body>
    {header_html}
    {content_html}
</body>
</html>"""

        logger.info(f"Generated HTML length: {len(html_content)}")
        logger.info(f"Layout type: {layout}")
        
        # Final cleanup: remove any remaining ** characters that might have slipped through
        # This ensures no ** characters appear in the final PDF
        # Optimized: single pass is usually enough, second pass only if needed
        html_content = html_content.replace('**', '')
        if '**' in html_content:
            html_content = html_content.replace('**', '')

        try:
            pdf_bytes = HTML(string=html_content).write_pdf()
            
            if not pdf_bytes or len(pdf_bytes) == 0:
                logger.error("PDF generation returned empty bytes")
                raise HTTPException(
                    status_code=500,
                    detail="PDF generation failed: Empty PDF file was generated. Please check your resume content."
                )
            
            if not isinstance(pdf_bytes, bytes):
                logger.error(f"PDF generation returned invalid type: {type(pdf_bytes)}")
                raise HTTPException(
                    status_code=500,
                    detail="PDF generation failed: Invalid PDF data type."
                )
            
            if not pdf_bytes.startswith(b"%PDF-"):
                logger.error(f"PDF validation failed: Invalid PDF header. First 20 bytes: {pdf_bytes[:20]}")
                raise HTTPException(
                    status_code=500,
                    detail="PDF generation failed: Generated file is not a valid PDF. Please try again or contact support."
                )
            
            logger.info(f"PDF generated successfully, size: {len(pdf_bytes)} bytes, header: {pdf_bytes[:8]}")
            
            # Return PDF immediately, track analytics asynchronously
            # Create response first
            response = Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename=resume.pdf",
                    "Content-Length": str(len(pdf_bytes)),
                },
            )
            
            # Track export analytics in background (non-blocking)
            if user_email and db:
                try:
                    # Use a separate try-except to not block the response
                    user = db.query(User).filter(User.email == user_email).first()
                    if user:
                        # Find or create resume record
                        resume = (
                            db.query(Resume)
                            .filter(Resume.user_id == user.id, Resume.name == payload.name)
                            .first()
                        )

                        if not resume:
                            resume = Resume(
                                user_id=user.id,
                                name=payload.name,
                                title=payload.title,
                                email=payload.email,
                                phone=payload.phone,
                                location=payload.location,
                                summary=payload.summary,
                                template=template_id,
                            )
                            db.add(resume)
                            db.flush()  # Get ID without full commit

                        # Create export analytics record
                        export_analytics = ExportAnalytics(
                            user_id=user.id if user else None,
                            session_id=session_id if not user else None,
                            resume_id=resume.id if resume else None,
                            export_format="pdf",
                            template_used=template_id,
                            file_size=len(pdf_bytes),
                            export_success=True,
                        )
                        db.add(export_analytics)
                        db.commit()  # Single commit for both operations

                        logger.info(
                            f"Export analytics tracked for user {user_email}: PDF export"
                        )
                except Exception as e:
                    logger.warning(f"Export analytics tracking failed (non-blocking): {e}")
                    try:
                        db.rollback()  # Rollback on error
                    except:
                        pass
            
            return response
        except AttributeError as e:
            if "'super' object has no attribute" in str(e) or "transform" in str(e):
                logger.error(f"WeasyPrint compatibility error: {str(e)}")
                raise HTTPException(
                    status_code=500,
                    detail="PDF generation failed due to a compatibility issue. Please upgrade WeasyPrint to 63.0+ or contact support."
                )
            raise
        except Exception as pdf_error:
            error_msg = str(pdf_error)
            logger.error(f"WeasyPrint PDF generation error: {error_msg}")
            
            if "transform" in error_msg.lower() or "super" in error_msg.lower():
                raise HTTPException(
                    status_code=500,
                    detail="PDF generation failed due to a compatibility issue. Please upgrade WeasyPrint to 63.0+ or contact support."
                )
            elif "font" in error_msg.lower() or "glyph" in error_msg.lower():
                raise HTTPException(
                    status_code=500,
                    detail="PDF generation failed due to font issues. Please try a different font or contact support."
                )
            else:
                raise HTTPException(
                    status_code=500,
                    detail=f"PDF generation failed: {error_msg}. Please check your resume content and try again."
                )
    except HTTPException:
        raise
    except ValueError as e:
        logger.error(f"PDF export validation error: {str(e)}")
        raise HTTPException(
            status_code=400,
            detail=f"Invalid resume data: {str(e)}"
        )
    except Exception as e:
        error_msg = str(e)
        logger.error(f"PDF export error: {error_msg}")
        import traceback
        logger.error(traceback.format_exc())
        
        # Provide more specific error messages
        if "template" in error_msg.lower():
            raise HTTPException(
                status_code=400,
                detail="Invalid template configuration. Please select a valid template and try again."
            )
        elif "weasyprint" in error_msg.lower() or "html" in error_msg.lower():
            raise HTTPException(
                status_code=500,
                detail=f"PDF generation error: {error_msg}. Please check your resume content and try again."
            )
        elif "missing" in error_msg.lower() or "required" in error_msg.lower():
            raise HTTPException(
                status_code=400,
                detail=f"Missing required data: {error_msg}"
            )
        else:
            # Include the actual error message for debugging
            raise HTTPException(
                status_code=500,
                detail=f"PDF generation failed: {error_msg}"
            )


async def export_docx(
    payload: ExportPayload,
    user_email: Optional[str] = None,
    session_id: Optional[str] = None,
    db: Optional[Session] = None,
) -> Response:
    """
    Export resume as DOCX.
    
    CRITICAL VISIBILITY FILTERING:
    - Sections with params.visible === False are excluded
    - Bullets with params.visible === False should be filtered (check b.params.get("visible"))
    - Fields with fieldsVisible[field] === False are excluded (summary, name, title, etc.)
    
    DO NOT remove these filters - they are essential for respecting user visibility settings.
    """
    """Export resume as DOCX"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        replacements = payload.replacements or {}
        template_id = payload.template or "tech"
        template_style = TEMPLATES.get(template_id, TEMPLATES["tech"])
        template_config = payload.templateConfig or {}

        logger.info(f"Exporting DOCX with template: {template_id}")
        logger.info(f"Template config provided: {bool(template_config)}")
        logger.info(
            f"Two-column settings: left={payload.two_column_left}, right={payload.two_column_right}, width={payload.two_column_left_width}"
        )

        # Use templateConfig if available, otherwise fall back to template lookup
        if template_config:
            layout_config = template_config.get("layout", {})
            layout_columns = layout_config.get("columns", "single")
            layout = "two-column" if layout_columns in ["two-column", "asymmetric"] else "single"
            
            typography = template_config.get("typography", {})
            font_family_heading = typography.get("fontFamily", {}).get("heading", "Arial")
            font_family_body = typography.get("fontFamily", {}).get("body", "Arial")
            h1_size = typography.get("fontSize", {}).get("h1", 24)
            h2_size = typography.get("fontSize", {}).get("h2", 12)
            body_size = typography.get("fontSize", {}).get("body", 11)
            
            design = template_config.get("design", {})
            header_style = design.get("headerStyle", "left-aligned")
            # Match frontend logic: center if headerStyle is 'centered' OR template is one of: clean, two-column, compact, classic
            templates_that_center = ["clean", "two-column", "compact", "classic"]
            header_align = "center" if (header_style == "centered" or template_id in templates_that_center) else "left"
            
            spacing_config = template_config.get("spacing", {})
            page_margin = spacing_config.get("pageMargin", 24)
            page_margin_inches = Inches(page_margin / 96)  # Convert px to inches (assuming 96 DPI)
        else:
            layout = template_style["styles"]["layout"]
            font_family_heading = "Arial" if "Arial" in template_style["styles"]["font"] else "Georgia"
            font_family_body = font_family_heading
            h1_size = 18
            h2_size = 12
            body_size = 11
            header_align = template_style["styles"]["header_align"]
            page_margin_inches = Inches(1)

        doc = Document()

        section = doc.sections[0]
        # Check if this is a cover letter only export (no resume sections)
        is_cover_letter_only = bool(payload.cover_letter and not payload.sections and not payload.summary)
        
        # For cover letter only, use larger margins for professional appearance
        if is_cover_letter_only:
            page_margin_inches = Inches(1.5)  # 1.5 inches margins for cover letters
        
        section.top_margin = page_margin_inches
        section.bottom_margin = page_margin_inches
        section.left_margin = page_margin_inches
        section.right_margin = page_margin_inches

        # Only add header (name, title, contact) if not cover letter only
        # CRITICAL: Check fieldsVisible for each field
        if not is_cover_letter_only:
            if payload.name and (not payload.fieldsVisible or payload.fieldsVisible.get("name") is not False):
                name_para = doc.add_paragraph()
                name_run = name_para.add_run(apply_replacements(payload.name, replacements))
                name_run.font.size = Pt(h1_size)
                name_run.font.bold = True
                name_run.font.name = font_family_heading
                if header_align == "center":
                    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if payload.title and (not payload.fieldsVisible or payload.fieldsVisible.get("title") is not False):
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(apply_replacements(payload.title, replacements))
                title_run.font.size = Pt(h2_size)
                title_run.font.name = font_family_body
                if header_align == "center":
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            contact_parts = []
            if payload.email and (not payload.fieldsVisible or payload.fieldsVisible.get("email") is not False):
                contact_parts.append(apply_replacements(payload.email, replacements))
            if payload.phone and (not payload.fieldsVisible or payload.fieldsVisible.get("phone") is not False):
                contact_parts.append(apply_replacements(payload.phone, replacements))
            if payload.location and (not payload.fieldsVisible or payload.fieldsVisible.get("location") is not False):
                contact_parts.append(apply_replacements(payload.location, replacements))

            if contact_parts:
                contact_para = doc.add_paragraph(" • ".join(contact_parts))
                contact_para.runs[0].font.size = Pt(body_size)
                contact_para.runs[0].font.name = font_family_body
                if header_align == "center":
                    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

        # Add cover letter if provided
        if payload.cover_letter:
            # Replace placeholders with actual values
            cover_letter_content = replace_cover_letter_placeholders(payload.cover_letter, payload)
            
            # Determine title
            if payload.company_name and payload.position_title:
                cover_letter_title = f"{payload.position_title} at {payload.company_name}"
            elif payload.company_name:
                cover_letter_title = f"{payload.company_name} - Cover Letter"
            else:
                cover_letter_title = "Cover Letter"
            
            # Add title
            if cover_letter_title:
                cover_letter_heading = doc.add_paragraph()
                cover_letter_heading_run = cover_letter_heading.add_run(cover_letter_title)
                cover_letter_heading_run.font.size = Pt(18 if is_cover_letter_only else 14)
                cover_letter_heading_run.font.bold = True
                cover_letter_heading_run.font.name = font_family_heading
                cover_letter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_cover_letter_only else WD_ALIGN_PARAGRAPH.LEFT
                doc.add_paragraph()
            
            # Parse and format cover letter content
            lines = cover_letter_content.split('\n')
            current_paragraph = []
            
            for i, line in enumerate(lines):
                line = line.strip()
                
                # Skip title if it matches
                if line == cover_letter_title:
                    continue
                
                # Empty line indicates paragraph break
                if not line:
                    if current_paragraph:
                        para_text = ' '.join(current_paragraph)
                        para_text = re.sub(r' +', ' ', para_text)
                        if para_text:
                            # Check for special formatting
                            para = doc.add_paragraph(para_text)
                            
                            # Date formatting (smaller, no spacing before)
                            if i < 5 and any(month in para_text for month in ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']):
                                para.runs[0].font.size = Pt(11)
                                para.paragraph_format.space_before = Pt(0)
                                para.paragraph_format.space_after = Pt(6)
                            # Salutation formatting
                            elif para_text.startswith('Dear'):
                                para.runs[0].font.size = Pt(12 if is_cover_letter_only else 11)
                                para.paragraph_format.space_before = Pt(0)
                                para.paragraph_format.space_after = Pt(12)
                            # Closing formatting
                            elif para_text in ['Sincerely,', 'Best regards,', 'Yours sincerely,']:
                                para.runs[0].font.size = Pt(12 if is_cover_letter_only else 11)
                                para.paragraph_format.space_before = Pt(24)
                                para.paragraph_format.space_after = Pt(6)
                            # Signature formatting (usually last line)
                            elif i > len(lines) - 3 and len(current_paragraph) == 1:
                                para.runs[0].font.size = Pt(12 if is_cover_letter_only else 11)
                                para.paragraph_format.space_before = Pt(0)
                                para.paragraph_format.space_after = Pt(6)
                            # Contact info formatting
                            elif '@' in para_text or any(char.isdigit() and len(para_text.replace(' ', '').replace('-', '').replace('(', '').replace(')', '')) >= 10 for char in para_text):
                                para.runs[0].font.size = Pt(10)
                                para.paragraph_format.space_before = Pt(0)
                                para.paragraph_format.space_after = Pt(4)
                            # Body paragraph formatting
                            else:
                                para.runs[0].font.size = Pt(12 if is_cover_letter_only else 11)
                                para.runs[0].font.name = font_family_body
                                para.paragraph_format.space_after = Pt(12 if is_cover_letter_only else 8)
                                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                para.paragraph_format.line_spacing = 1.15
                            
                            para.runs[0].font.name = font_family_body
                        current_paragraph = []
                    continue
                
                current_paragraph.append(line)
            
            # Handle remaining paragraph
            if current_paragraph:
                para_text = ' '.join(current_paragraph)
                para_text = re.sub(r' +', ' ', para_text)
                if para_text and para_text != cover_letter_title:
                    para = doc.add_paragraph(para_text)
                    para.runs[0].font.size = Pt(12 if is_cover_letter_only else 11)
                    para.runs[0].font.name = font_family_body
                    para.paragraph_format.space_after = Pt(12 if is_cover_letter_only else 8)
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.line_spacing = 1.15

            doc.add_paragraph()  # Add spacing

        # Determine layout from templateConfig or fallback
        is_two_column = layout == "two-column" or (
            template_config and 
            template_config.get("layout", {}).get("columns") in ["two-column", "asymmetric"]
        )
        
        logger.info(
            f"Layout decision: layout='{layout}', is_two_column={is_two_column}"
        )
        if is_two_column:
            logger.info("Creating two-column layout")
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            # Create a table for two-column layout
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False

            # Remove table borders
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl.insert(0, tblPr)

            tblBorders = OxmlElement("w:tblBorders")
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "none")
                border.set(qn("w:sz"), "0")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "auto")
                tblBorders.append(border)
            tblPr.append(tblBorders)

            # Set column widths based on user preference
            left_width_percent = payload.two_column_left_width or 50
            right_width_percent = 100 - left_width_percent

            # Convert percentages to inches (assuming 8.5" total width)
            left_width_inches = (left_width_percent / 100) * 8.5
            right_width_inches = (right_width_percent / 100) * 8.5

            table.columns[0].width = Inches(left_width_inches)
            table.columns[1].width = Inches(right_width_inches)

            logger.info(
                f'Column widths: left={left_width_inches:.1f}", right={right_width_inches:.1f}"'
            )

            # Left column
            left_cell = table.cell(0, 0)
            # Clear existing paragraphs in left cell
            for paragraph in left_cell.paragraphs:
                paragraph.clear()

            # Right column
            right_cell = table.cell(0, 1)
            # Clear existing paragraphs in right cell
            for paragraph in right_cell.paragraphs:
                paragraph.clear()

            # Use localStorage configuration if provided, otherwise fallback to alternating
            left_section_ids = set(payload.two_column_left or [])
            right_section_ids = set(payload.two_column_right or [])

            logger.info(
                f"Section IDs from payload: left={left_section_ids}, right={right_section_ids}"
            )
            logger.info(
                f"Available sections: {[s.id for s in payload.sections if s.id]}"
            )

            if not left_section_ids and not right_section_ids:
                # Use smart default distribution: Skills, Certificates, Education on left
                logger.info(
                    "Using smart default distribution - no localStorage configuration provided"
                )
                left_column_keywords = ['skill', 'certificate', 'certification', 'education', 'academic', 'qualification', 'award', 'honor']
                
                left_sections = [
                    s for s in payload.sections 
                    if any(keyword in s.title.lower() for keyword in left_column_keywords)
                ]
                right_sections = [
                    s for s in payload.sections 
                    if not any(keyword in s.title.lower() for keyword in left_column_keywords)
                ]
            else:
                # Use provided configuration
                logger.info("Using localStorage configuration")
                left_sections = [
                    s for s in payload.sections if s.id and s.id in left_section_ids
                ]
                right_sections = [
                    s for s in payload.sections if s.id and s.id in right_section_ids
                ]

            logger.info(
                f"Final section distribution: left={[s.title for s in left_sections]}, right={[s.title for s in right_sections]}"
            )
            logger.info(
                f"Left sections count: {len(left_sections)}, Right sections count: {len(right_sections)}"
            )

            # Add summary to appropriate column if assigned
            summary_id = "__summary__"
            summary_in_left = summary_id in left_section_ids
            summary_in_right = summary_id in right_section_ids
            
            if payload.summary and (summary_in_left or summary_in_right):
                summary_title = "Professional Summary"
                if template_style.get("styles", {}).get("section_uppercase", False):
                    summary_title = summary_title.upper()
                
                target_cell = left_cell if summary_in_left else right_cell
                summary_para = target_cell.add_paragraph()
                summary_heading = summary_para.add_run(summary_title)
                summary_heading.font.size = Pt(h2_size)
                summary_heading.font.bold = True
                summary_heading.font.name = font_family_heading
                
                summary_content_para = target_cell.add_paragraph()
                summary_content = summary_content_para.add_run(apply_replacements(payload.summary, replacements))
                summary_content.font.size = Pt(body_size)
                summary_content.font.name = font_family_body
                
                target_cell.add_paragraph()  # Add spacing
            elif payload.summary and not left_section_ids and not right_section_ids:
                # Default: add summary to left column
                summary_title = "Professional Summary"
                if template_style.get("styles", {}).get("section_uppercase", False):
                    summary_title = summary_title.upper()
                
                summary_para = left_cell.add_paragraph()
                summary_heading = summary_para.add_run(summary_title)
                summary_heading.font.size = Pt(h2_size)
                summary_heading.font.bold = True
                summary_heading.font.name = font_family_heading
                
                summary_content_para = left_cell.add_paragraph()
                summary_content = summary_content_para.add_run(apply_replacements(payload.summary, replacements))
                summary_content.font.size = Pt(body_size)
                summary_content.font.name = font_family_body
                
                left_cell.add_paragraph()  # Add spacing

            # Add sections to left column
            for s in left_sections:
                section_title = apply_replacements(s.title, replacements)
                if template_style["styles"]["section_uppercase"]:
                    section_title = section_title.upper()

                heading_para = left_cell.add_paragraph()
                heading_run = heading_para.add_run(section_title)
                heading_run.font.size = Pt(12)
                heading_run.font.bold = True
                heading_run.font.name = (
                    "Arial"
                    if "Arial" in template_style["styles"]["font"]
                    else "Georgia"
                )

                # Check section type
                section_lower = s.title.lower()
                is_work_experience = (
                    "experience" in section_lower
                    or "work" in section_lower
                    or "employment" in section_lower
                )
                is_skill_section = (
                    "skill" in section_lower
                    or "technical" in section_lower
                    or "technology" in section_lower
                    or "competencies" in section_lower
                    or "expertise" in section_lower
                    or "proficiencies" in section_lower
                )

                # Create content for this section
                if s.bullets:
                    if is_work_experience:
                        # Work experience - handle company headers
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            
                            # Check if company header
                            if bullet_text.startswith("**") and "**" in bullet_text[2:]:
                                header_text = bullet_text.replace("**", "").strip()
                                parts = header_text.split(' / ')
                                company_name = parts[0] if parts else ''
                                location = parts[1] if len(parts) >= 4 else ''
                                title = parts[2] if len(parts) >= 4 else (parts[1] if len(parts) >= 2 else '')
                                date_range = parts[3] if len(parts) >= 4 else (parts[2] if len(parts) >= 3 else '')
                                
                                # Company header - Line 1
                                company_para = left_cell.add_paragraph()
                                company_line = company_name
                                if location:
                                    company_line += f' / {location}'
                                company_run = company_para.add_run(company_line)
                                company_run.font.size = Pt(11)
                                company_run.font.bold = True
                                
                                # Company header - Line 2
                                title_para = left_cell.add_paragraph()
                                title_run = title_para.add_run(title)
                                title_run.font.size = Pt(10)
                                title_run.font.bold = True
                                date_run = title_para.add_run(f'\t{date_range}')
                                date_run.font.size = Pt(9)
                                date_run.font.italic = True
                            else:
                                # Regular bullet - strip all bullet markers
                                clean_text = strip_bullet_markers(bullet_text)
                                if clean_text:
                                    bullet_para = left_cell.add_paragraph()
                                    bullet_para.style = "List Bullet"
                                    run = bullet_para.add_run(clean_text)
                                    run.font.size = Pt(10)
                    elif is_skill_section:
                        # Skills section - comma-separated, no bullets
                        skill_items = []
                        for b in s.bullets:
                            if b.text.strip():
                                bullet_text = apply_replacements(b.text, replacements)
                                clean_text = strip_bullet_markers(bullet_text)
                                if clean_text:
                                    skill_items.append(clean_text)
                        if skill_items:
                            skills_para = left_cell.add_paragraph(", ".join(skill_items))
                            skills_para.runs[0].font.size = Pt(10)
                    else:
                        # Regular section with bullets
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            # Strip all bullet markers
                            clean_text = strip_bullet_markers(bullet_text)
                            if clean_text:
                                bullet_para = left_cell.add_paragraph()
                                bullet_para.style = "List Bullet"
                                run = bullet_para.add_run(clean_text)
                                run.font.size = Pt(10)

                left_cell.add_paragraph()

            # Add sections to right column
            for s in right_sections:
                section_title = apply_replacements(s.title, replacements)
                if template_style["styles"]["section_uppercase"]:
                    section_title = section_title.upper()

                heading_para = right_cell.add_paragraph()
                heading_run = heading_para.add_run(section_title)
                heading_run.font.size = Pt(12)
                heading_run.font.bold = True
                heading_run.font.name = (
                    "Arial"
                    if "Arial" in template_style["styles"]["font"]
                    else "Georgia"
                )

                # Check section type
                section_lower = s.title.lower()
                is_work_experience = (
                    "experience" in section_lower
                    or "work" in section_lower
                    or "employment" in section_lower
                )
                is_skill_section = (
                    "skill" in section_lower
                    or "technical" in section_lower
                    or "technology" in section_lower
                    or "competencies" in section_lower
                    or "expertise" in section_lower
                    or "proficiencies" in section_lower
                )

                # Create content for this section
                if s.bullets:
                    if is_work_experience:
                        # Work experience - handle company headers
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            
                            # Check if company header
                            if bullet_text.startswith("**") and "**" in bullet_text[2:]:
                                header_text = bullet_text.replace("**", "").strip()
                                parts = header_text.split(' / ')
                                company_name = parts[0] if parts else ''
                                location = parts[1] if len(parts) >= 4 else ''
                                title = parts[2] if len(parts) >= 4 else (parts[1] if len(parts) >= 2 else '')
                                date_range = parts[3] if len(parts) >= 4 else (parts[2] if len(parts) >= 3 else '')
                                
                                # Company header - Line 1
                                company_para = right_cell.add_paragraph()
                                company_line = company_name
                                if location:
                                    company_line += f' / {location}'
                                company_run = company_para.add_run(company_line)
                                company_run.font.size = Pt(11)
                                company_run.font.bold = True
                                
                                # Company header - Line 2
                                title_para = right_cell.add_paragraph()
                                title_run = title_para.add_run(title)
                                title_run.font.size = Pt(10)
                                title_run.font.bold = True
                                date_run = title_para.add_run(f'\t{date_range}')
                                date_run.font.size = Pt(9)
                                date_run.font.italic = True
                            else:
                                # Regular bullet - strip existing bullets
                                clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                                clean_text = clean_text.lstrip('•*- ').strip()
                                if clean_text:
                                    bullet_para = right_cell.add_paragraph()
                                    bullet_para.style = "List Bullet"
                                    run = bullet_para.add_run(clean_text)
                                    run.font.size = Pt(10)
                    elif is_skill_section:
                        # Skills section - comma-separated, no bullets
                        skill_items = []
                        for b in s.bullets:
                            if b.text.strip():
                                bullet_text = apply_replacements(b.text, replacements)
                                clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                                clean_text = clean_text.lstrip('•*- ').strip()
                                if clean_text:
                                    skill_items.append(clean_text)
                        if skill_items:
                            skills_para = right_cell.add_paragraph(", ".join(skill_items))
                            skills_para.runs[0].font.size = Pt(10)
                    else:
                        # Regular section with bullets
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            # Strip existing bullets
                            clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                            clean_text = clean_text.lstrip('•*- ').strip()
                            if clean_text:
                                bullet_para = right_cell.add_paragraph()
                                bullet_para.style = "List Bullet"
                                run = bullet_para.add_run(clean_text)
                                run.font.size = Pt(10)

                right_cell.add_paragraph()
        else:
            # Single column layout
            # CRITICAL: Check fieldsVisible for summary
            if payload.summary and (not payload.fieldsVisible or payload.fieldsVisible.get("summary") is not False):
                summary_para = doc.add_paragraph(
                    apply_replacements(payload.summary, replacements)
                )
                summary_para.runs[0].font.size = Pt(10)
                doc.add_paragraph()

            # CRITICAL: Filter sections by visible property
            for s in payload.sections:
                if s.params and s.params.get("visible") is False:
                    continue
                section_title = apply_replacements(s.title, replacements)
                if template_style["styles"]["section_uppercase"]:
                    section_title = section_title.upper()

                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run(section_title)
                heading_run.font.size = Pt(12)
                heading_run.font.bold = True
                heading_run.font.name = (
                    "Arial"
                    if "Arial" in template_style["styles"]["font"]
                    else "Georgia"
                )

                # Check section type
                section_lower = s.title.lower()
                is_work_experience = (
                    "experience" in section_lower
                    or "work" in section_lower
                    or "employment" in section_lower
                )
                is_skill_section = (
                    "skill" in section_lower
                    or "technical" in section_lower
                    or "technology" in section_lower
                    or "competencies" in section_lower
                    or "expertise" in section_lower
                    or "proficiencies" in section_lower
                )

                # Create content for this section
                if s.bullets:
                    if is_work_experience:
                        # Work experience - handle company headers
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            
                            # Check if company header
                            if bullet_text.startswith("**") and "**" in bullet_text[2:]:
                                header_text = bullet_text.replace("**", "").strip()
                                parts = header_text.split(' / ')
                                company_name = parts[0] if parts else ''
                                location = parts[1] if len(parts) >= 4 else ''
                                title = parts[2] if len(parts) >= 4 else (parts[1] if len(parts) >= 2 else '')
                                date_range = parts[3] if len(parts) >= 4 else (parts[2] if len(parts) >= 3 else '')
                                
                                # Company header - Line 1
                                company_para = doc.add_paragraph()
                                company_line = company_name
                                if location:
                                    company_line += f' / {location}'
                                company_run = company_para.add_run(company_line)
                                company_run.font.size = Pt(11)
                                company_run.font.bold = True
                                
                                # Company header - Line 2
                                title_para = doc.add_paragraph()
                                title_run = title_para.add_run(title)
                                title_run.font.size = Pt(10)
                                title_run.font.bold = True
                                date_run = title_para.add_run(f'\t{date_range}')
                                date_run.font.size = Pt(9)
                                date_run.font.italic = True
                            else:
                                # Regular bullet - strip existing bullets
                                clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                                clean_text = clean_text.lstrip('•*- ').strip()
                                if clean_text:
                                    bullet_para = doc.add_paragraph()
                                    bullet_para.style = "List Bullet"
                                    run = bullet_para.add_run(clean_text)
                                    run.font.size = Pt(10)
                    elif is_skill_section:
                        # Skills section - comma-separated, no bullets
                        skill_items = []
                        for b in s.bullets:
                            if b.text.strip():
                                bullet_text = apply_replacements(b.text, replacements)
                                clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                                clean_text = clean_text.lstrip('•*- ').strip()
                                if clean_text:
                                    skill_items.append(clean_text)
                        if skill_items:
                            skills_para = doc.add_paragraph(", ".join(skill_items))
                            skills_para.runs[0].font.size = Pt(10)
                    else:
                        # Regular section with bullets
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            # Strip existing bullets
                            clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                            clean_text = clean_text.lstrip('•*- ').strip()
                            if clean_text:
                                bullet_para = doc.add_paragraph()
                                bullet_para.style = "List Bullet"
                                run = bullet_para.add_run(clean_text)
                                run.font.size = Pt(10)

                doc.add_paragraph()

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        docx_bytes = buffer.getvalue()
        logger.info(f"DOCX generated successfully, size: {len(docx_bytes)} bytes")

        # Track export analytics
        if user_email and db:
            try:
                user = db.query(User).filter(User.email == user_email).first()
                if user:
                    # Find or create resume record
                    resume = (
                        db.query(Resume)
                        .filter(Resume.user_id == user.id, Resume.name == payload.name)
                        .first()
                    )

                    if not resume:
                        resume = Resume(
                            user_id=user.id,
                            name=payload.name,
                            title=payload.title,
                            email=payload.email,
                            phone=payload.phone,
                            location=payload.location,
                            summary=payload.summary,
                            template=template_id,
                        )
                        db.add(resume)
                        db.commit()
                        db.refresh(resume)

                    # Create export analytics record
                    export_analytics = ExportAnalytics(
                        user_id=user.id if user else None,
                        session_id=session_id if not user else None,
                        resume_id=resume.id if resume else None,
                        export_format="docx",
                        template_used=template_id,
                        file_size=len(docx_bytes),
                        export_success=True,
                    )
                    db.add(export_analytics)
                    db.commit()

                    logger.info(
                        f"Export analytics tracked for user {user_email}: DOCX export"
                    )
            except Exception as e:
                logger.error(f"Failed to track export analytics: {e}")

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=resume.docx"},
        )
    except Exception as e:
        logger.error(f"DOCX export error: {str(e)}")
        import traceback

        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

