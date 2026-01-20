"""Extract DOCX content while preserving styles and structure."""

from __future__ import annotations

import logging
from io import BytesIO
from typing import Any

logger = logging.getLogger(__name__)


def extract_docx_with_structure(file_bytes: bytes) -> dict[str, Any]:
    """
    Extract DOCX content while preserving:
    - Paragraph styles (bold, italic, size, alignment)
    - Table structure
    - Header/footer content

    Returns:
        {
            'paragraphs': [
                {
                    'text': str,
                    'style': str,
                    'bold': bool,
                    'italic': bool,
                    'size': float,
                    'alignment': int
                }
            ],
            'tables': [list of extracted tables],
            'metadata': {'has_tables': bool, 'has_styles': bool}
        }
    """
    try:
        from docx import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        logger.error("python-docx not installed")
        raise ImportError("python-docx is required for DOCX extraction")

    docx_file = BytesIO(file_bytes)
    paragraphs_data = []
    tables_data = []
    has_styles = False

    try:
        doc = Document(docx_file)

        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                if para.text.strip():
                    # Extract style information
                    style_name = para.style.name if para.style else ""
                    runs = para.runs
                    
                    # Check for bold/italic in runs
                    is_bold = any(run.bold for run in runs if run.bold is not None)
                    is_italic = any(run.italic for run in runs if run.italic is not None)
                    
                    # Get font size (use first run's size)
                    font_size = None
                    if runs:
                        for run in runs:
                            if run.font.size:
                                font_size = run.font.size.pt
                                has_styles = True
                                break
                    
                    # Get alignment
                    alignment = para.alignment
                    alignment_int = alignment.value if alignment else 0

                    para_data = {
                        'text': para.text,
                        'style': style_name,
                        'bold': is_bold,
                        'italic': is_italic,
                        'size': font_size,
                        'alignment': alignment_int,
                    }
                    paragraphs_data.append(para_data)

            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = "\n".join(
                            para.text for para in cell.paragraphs if para.text.strip()
                        )
                        row_data.append(cell_text)
                    table_data.append(row_data)
                tables_data.append(table_data)

        # Extract header/footer content
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        paragraphs_data.insert(0, {
                            'text': para.text,
                            'style': 'Header',
                            'bold': False,
                            'italic': False,
                            'size': None,
                            'alignment': 0,
                        })
            
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        paragraphs_data.append({
                            'text': para.text,
                            'style': 'Footer',
                            'bold': False,
                            'italic': False,
                            'size': None,
                            'alignment': 0,
                        })

        logger.info(
            f"DOCX extraction complete: {len(paragraphs_data)} paragraphs, "
            f"{len(tables_data)} tables extracted"
        )

        return {
            'paragraphs': paragraphs_data,
            'tables': tables_data,
            'metadata': {
                'has_tables': len(tables_data) > 0,
                'has_styles': has_styles,
            },
        }

    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise


def extract_docx_text_only(file_bytes: bytes) -> str:
    """Extract plain text from DOCX (fallback method)."""
    try:
        from docx import Document
        docx_file = BytesIO(file_bytes)
        doc = Document(docx_file)
        
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            text += para.text + "\n"
        
        return text
    except Exception as e:
        logger.error(f"DOCX text extraction error: {e}")
        raise
