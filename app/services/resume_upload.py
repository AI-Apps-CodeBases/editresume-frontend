"""Resume upload service - file parsing and text extraction"""

from __future__ import annotations

import json
import logging
import re
from io import BytesIO
from typing import Dict, List, Optional

from app.core.openai_client import OPENAI_MAX_TOKENS, USE_AI_PARSER, openai_client
from app.utils.resume_parsing import parse_resume_with_regex

logger = logging.getLogger(__name__)


def extract_docx_text_from_bytes(file_content: bytes) -> str:
    """Extract text from DOCX file bytes"""
    try:
        from docx import Document
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        text = ""
        docx_file = BytesIO(file_content)
        doc = Document(docx_file)

        def extract_text_from_element(element):
            elem_text = ""
            if isinstance(element, Paragraph):
                if element.text.strip():
                    elem_text += element.text + "\n"
            elif isinstance(element, Table):
                for row in element.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                elem_text += para.text + "\n"
            return elem_text

        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                if para.text.strip():
                    text += para.text + "\n"
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                text += para.text + "\n"

        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"

        logger.info(
            f"DOCX extracted {len(text)} characters from {len(doc.paragraphs)} paragraphs"
        )
        return text
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        raise


def extract_pdf_text_from_bytes(file_content: bytes) -> str:
    """Extract text from PDF file bytes"""
    try:
        import pdfplumber

        text = ""
        pdf_file = BytesIO(file_content)
        with pdfplumber.open(pdf_file) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"PDF has {total_pages} pages")

            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num} ---\n"
                    text += page_text + "\n"
                    logger.info(
                        f"Extracted {len(page_text)} characters from page {page_num}"
                    )

        logger.info(
            f"PDF extraction complete: {total_pages} pages, {len(text)} total characters"
        )
        return text
    except Exception as e:
        error_msg = str(e)
        logger.error(f"PDF extraction error: {error_msg}")
        if "Root object" in error_msg or "really a PDF" in error_msg:
            raise ValueError(
                "PDF file appears to be corrupted or invalid. Try:\n1. Re-saving the PDF from its original application\n2. Converting to DOCX first\n3. Or use manual entry"
            )
        raise


async def upload_and_parse_resume(file_content: bytes, filename: str, content_type: Optional[str] = None) -> Dict:
    """Upload and parse a resume file"""
    try:
        file_type = filename.split(".")[-1].lower() if "." in filename else ""

        logger.info(
            f"Upload: filename={filename}, type={file_type}, content_type={content_type}, size={len(file_content)}"
        )

        magic_bytes = file_content[:4]
        is_docx = magic_bytes == b"PK\x03\x04"
        is_pdf = file_content[:5] == b"%PDF-"

        logger.info(f"File magic: is_docx={is_docx}, is_pdf={is_pdf}")

        text = ""
        actual_type = ""

        if (
            is_docx
            or file_type in ["docx", "doc"]
            or (content_type and "wordprocessingml" in content_type)
        ):
            logger.info("Processing as DOCX")
            actual_type = "DOCX"
            try:
                text = extract_docx_text_from_bytes(file_content)
            except Exception as docx_error:
                logger.error(f"DOCX error: {str(docx_error)}")
                return {
                    "success": False,
                    "error": f"Could not read DOCX file: {str(docx_error)}",
                }

        elif is_pdf or file_type == "pdf" or content_type == "application/pdf":
            logger.info("Processing as PDF")
            actual_type = "PDF"
            try:
                text = extract_pdf_text_from_bytes(file_content)
            except ValueError as pdf_error:
                return {
                    "success": False,
                    "error": str(pdf_error),
                }
            except Exception as pdf_error:
                raise

        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return {
                "success": False,
                "error": f"Unsupported file type: {file_type}. Please upload PDF or DOCX.",
            }

        if not text.strip():
            logger.warning(f"No text extracted from {actual_type} file")
            return {
                "success": False,
                "error": f"Could not extract text from {actual_type} file. The file might be:\n• Empty\n• Using special formatting (text boxes)\n• Password protected\n\nPlease try exporting from Word as a new DOCX, or use manual entry.",
            }

        logger.info("Using AI-powered parsing for better resume organization...")
        parsed_data = parse_resume_with_ai(text)

        logger.info(
            "AI parsing successful: text_length=%s, sections=%s, summary_len=%s",
            len(text),
            len(parsed_data.get("sections", [])) if isinstance(parsed_data, dict) else "n/a",
            len(parsed_data.get("summary", "")) if isinstance(parsed_data, dict) else "n/a",
        )
        return {
            "success": True,
            "data": parsed_data,
            "raw_text": text[:1000],
            "debug": {
                "file_type": file_type,
                "text_length": len(text),
                "has_content": bool(text.strip()),
            },
        }
    except Exception as e:
        return {"success": False, "error": f"Upload failed: {str(e)}"}


def parse_resume_with_ai(text: str) -> Dict:
    """Use AI to intelligently parse and structure resume content"""
    if not USE_AI_PARSER:
        logger.info(
            "AI resume parser disabled via USE_AI_PARSER flag, using basic parser"
        )
        return parse_resume_text(text)

    if not openai_client:
        logger.warning("AI parsing not available, falling back to basic parsing")
        return parse_resume_text(text)

    try:
        prompt = f"""Parse this resume text and extract structured information. This resume may have been exported from a resume builder app, so pay special attention to preserving the exact structure and content.

CRITICAL: You MUST return valid JSON with ALL property names enclosed in double quotes. Do not use single quotes or unquoted property names.

Return a JSON object with this exact structure:

{{
  "name": "Full Name",
  "title": "Professional Title or Current Role",
  "email": "email@example.com",
  "phone": "+1-234-567-8900",
  "location": "City, State/Country",
  "summary": "Professional summary or objective copied verbatim from the resume",
  "sections": [
    {{
      "title": "Work Experience",
      "bullets": [
        "**Company Name / Job Title / Start Date - End Date**",
        "• Achievement or task for this role",
        "• Another achievement with metrics",
        "• Key responsibility or accomplishment",
        "",
        "**Another Company / Job Title / Start Date - End Date**",
        "• Achievement at this company",
        "• Another task or responsibility"
      ]
    }},
    {{
      "title": "Education",
      "bullets": [
        "Degree, Institution, Year",
        "GPA or honors if applicable"
      ]
    }},
    {{
      "title": "Skills",
      "bullets": [
        "Skill category: specific skills",
        "Another category: more skills"
      ]
    }}
  ]
}}

CRITICAL PARSING RULES FOR APP-EXPORTED RESUMES:
1. PRESERVE ALL CONTENT - Don't skip, summarize, or paraphrase anything
2. Maintain exact bullet point structure and wording
3. For work experience sections, preserve company names, job titles, dates, and ALL achievements
4. Keep quantified metrics exactly as written (e.g., "30% reduction", "over 100 applications")
5. Preserve technical skills and tools mentioned (AWS, Kubernetes, Docker, etc.)
6. Maintain project descriptions with all details
7. Extract contact information from header (name, email, phone, location)
8. Handle multi-page content by combining everything
9. If you see patterns like "reduced deployment time by 30%" repeated, keep each instance
10. For DevOps/Engineering resumes, preserve all technical details and metrics
11. Don't paraphrase or shorten content - extract exactly as written
12. Group content into logical sections but preserve all bullet points

SPECIAL ATTENTION TO:
- Quantified achievements (percentages, numbers, metrics)
- Technical skills and tools (AWS, Kubernetes, Docker, Jenkins, etc.)
- Company names and job titles
- Project descriptions and outcomes
- All bullet points with their complete text
- Work experience entries with exact dates and roles

CRITICAL RULES FOR WORK EXPERIENCE:
1. Company/Job line MUST be: **Company Name / Job Title / Date Range** (with ** for bold)
2. Tasks/achievements MUST start with "• " (bullet point)
3. Separate different jobs with empty string ""
4. Extract complete date ranges (e.g., "Jan 2020 - Dec 2022" or "2020-2022")
5. Include ALL jobs from the resume

OTHER RULES:
6. Identify ALL sections (Education, Skills, Projects, Certifications, Awards, etc.)
7. Preserve important details (dates, technologies, metrics, achievements)
8. If something is missing (like phone or email), leave it as empty string
9. Professional summary MUST be an exact copy of the resume's Professional Summary section without rewording
10. Every skill listed in the resume must appear in the JSON output exactly as written (split into multiple bullets if needed)
11. Return ONLY valid JSON, no markdown code blocks
12. This may be a MULTI-PAGE resume - extract ALL information from ALL pages
13. Include everything: all work experience, education, skills, projects, certifications

Resume Text (Full Content):
{text[:12000]}
"""

        # For very long resumes, use higher token limit
        resume_length = len(text)
        max_tokens_for_resume = (
            min(4000, OPENAI_MAX_TOKENS) if resume_length > 8000 else OPENAI_MAX_TOKENS
        )

        logger.info(
            f"Parsing resume: {resume_length} characters, using {max_tokens_for_resume} max tokens"
        )

        response = openai_client["requests"].post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {openai_client['api_key']}",
                "Content-Type": "application/json",
            },
            json={
                "model": openai_client["model"],
                "messages": [
                    {
                        "role": "system",
                        "content": "You are a resume parsing expert specializing in app-exported resumes. Extract structured information from multi-page resumes accurately. Capture ALL sections and information without paraphrasing or summarizing. Preserve exact wording, metrics, and technical details. Always return valid JSON.",
                    },
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.3,
                "max_tokens": max_tokens_for_resume,
            },
            timeout=120,
        )

        if response.status_code != 200:
            logger.error(f"OpenAI API error: {response.status_code} - {response.text}")
            return parse_resume_text(text)

        result = response.json()
        ai_response = result["choices"][0]["message"]["content"].strip()

        # Clean up the response
        ai_response = re.sub(r"^```json\s*", "", ai_response)
        ai_response = re.sub(r"\s*```$", "", ai_response)

        # Try to fix common JSON formatting issues
        try:
            parsed_data = json.loads(ai_response)
        except json.JSONDecodeError as e:
            logger.warning(f"Initial JSON parse failed: {e}")
            # Try to fix common JSON issues
            ai_response = re.sub(
                r"(\w+):", r'"\1":', ai_response
            )  # Add quotes to unquoted keys
            ai_response = re.sub(
                r':\s*([^",\[\{][^,}\]]*?)([,}\]])', r': "\1"\2', ai_response
            )  # Add quotes to unquoted string values
            try:
                parsed_data = json.loads(ai_response)
            except json.JSONDecodeError as e2:
                logger.error(f"JSON cleanup failed: {e2}")
                raise e2

        sections_raw = parsed_data.get("sections") or []
        normalized_sections: List[Dict] = []
        summary_section_text = ""
        summary_titles = {
            "professional summary",
            "summary",
            "executive summary",
            "profile",
            "about me",
        }

        for raw_section in sections_raw:
            section_title = ""
            section_params: Dict = {}
            bullets_source: List = []

            if isinstance(raw_section, dict):
                section_title = str(raw_section.get("title", "")).strip()
                section_params = raw_section.get("params") or {}
                bullets_source = raw_section.get("bullets", []) or []
            else:
                section_title = str(raw_section).strip()
                bullets_source = []

            bullet_entries: List[Dict] = []
            for bullet in bullets_source:
                if isinstance(bullet, dict):
                    bullet_entries.append(
                        {
                            "text": str(bullet.get("text", "")),
                            "params": bullet.get("params") or {},
                        }
                    )
                else:
                    bullet_entries.append({"text": str(bullet), "params": {}})

            title_lower = section_title.lower()
            if title_lower in summary_titles and bullet_entries and not summary_section_text:
                combined_summary = " ".join(
                    entry["text"].strip() for entry in bullet_entries if entry["text"].strip()
                ).strip()
                if combined_summary:
                    summary_section_text = combined_summary
                # Skip adding this section; the content will populate parsed_data["summary"]
                continue

            normalized_sections.append(
                {
                    "title": section_title,
                    "params": section_params,
                    "bullets": bullet_entries,
                }
            )

        existing_summary = (parsed_data.get("summary") or "").strip()
        if summary_section_text:
            if not existing_summary or existing_summary != summary_section_text:
                parsed_data["summary"] = summary_section_text

        skills_section = next(
            (section for section in normalized_sections if "skill" in (section.get("title") or "").lower()),
            None,
        )
        skills_bullets_present = bool(
            skills_section
            and any(entry.get("text", "").strip() for entry in skills_section.get("bullets", []))
        )
        if not skills_bullets_present:
            extracted_skills = _extract_skills_from_text(text)
            if extracted_skills:
                fallback_bullets = [{"text": skill, "params": {}} for skill in extracted_skills]
                if skills_section:
                    skills_section["bullets"] = fallback_bullets
                else:
                    normalized_sections.append(
                        {
                            "title": "Skills",
                            "params": {},
                            "bullets": fallback_bullets,
                        }
                    )

        final_sections: List[Dict] = []
        for idx, section in enumerate(normalized_sections):
            bullets: List[Dict] = []
            for bullet_idx, entry in enumerate(section.get("bullets", [])):
                bullet_text = str(entry.get("text", ""))
                bullet_params = entry.get("params") or {}
                bullets.append(
                    {
                        "id": f"{idx}-{bullet_idx}",
                        "text": bullet_text,
                        "params": bullet_params,
                    }
                )

            final_sections.append(
                {
                    "id": str(idx),
                    "title": section.get("title") or f"Section {idx + 1}",
                    "bullets": bullets,
                    "params": section.get("params") or {},
                }
            )

        parsed_data["sections"] = final_sections

        parsed_data.setdefault("detected_variables", {})

        logger.info(
            f"AI parsing successful: {len(parsed_data.get('sections', []))} sections extracted"
        )
        return parsed_data

    except json.JSONDecodeError as e:
        logger.error(f"AI returned invalid JSON: {e}")
        logger.info("Falling back to basic text parsing")
        return parse_resume_text(text)
    except Exception as e:
        logger.error(f"AI parsing failed: {str(e)}")
        logger.info("Falling back to basic text parsing")
        return parse_resume_text(text)


def _extract_skills_from_text(raw_text: str) -> List[str]:
    """Extract skills from text"""
    if not raw_text:
        return []

    skill_headings = [
        "skills",
        "technical skills",
        "technical skills & tools",
        "core competencies",
        "core skills",
        "technologies",
        "tools & technologies",
        "tech stack",
    ]
    heading_prefixes = tuple(skill.lower() for skill in skill_headings)

    lines = raw_text.splitlines()
    captured_lines: List[str] = []
    capturing = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if capturing and captured_lines:
                break
            continue

        lower_line = stripped.lower()
        if not capturing and lower_line.startswith(heading_prefixes):
            capturing = True
            if ":" in stripped:
                after_colon = stripped.split(":", 1)[1].strip()
                if after_colon:
                    captured_lines.append(after_colon)
            continue

        if capturing:
            if stripped == stripped.upper() and len(stripped.split()) <= 6:
                break
            if re.match(r"^[A-Z][A-Za-z0-9/&\-\s]{2,40}:$", stripped):
                break
            if lower_line.startswith(heading_prefixes):
                continue
            captured_lines.append(stripped)

    if not captured_lines:
        return []

    skills: List[str] = []
    for entry in captured_lines:
        cleaned = entry.strip("•·▪∙●○-–— ")
        if not cleaned:
            continue
        if "," in cleaned:
            parts = [part.strip() for part in cleaned.split(",") if part.strip()]
            skills.extend(parts)
        else:
            skills.append(cleaned)

    deduped: List[str] = []
    seen = set()
    for skill in skills:
        lowered = skill.lower()
        if lowered not in seen:
            seen.add(lowered)
            deduped.append(skill)

    return deduped


def parse_resume_text(text: str) -> Dict:
    """Basic text parsing fallback"""
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    data = {
        "name": "",
        "title": "",
        "email": "",
        "phone": "",
        "location": "",
        "summary": "",
        "sections": [],
        "detected_variables": {},
    }

    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    if email_match:
        data["email"] = email_match.group()

    phone_match = re.search(r"[\+\(]?[0-9][0-9 .\-\(\)]{8,}[0-9]", text)
    if phone_match:
        data["phone"] = phone_match.group()

    if lines:
        data["name"] = lines[0]
        if len(lines) > 1:
            data["title"] = lines[1]

    section_headers = [
        "experience",
        "education",
        "skills",
        "projects",
        "certifications",
        "summary",
    ]
    current_section = None
    current_bullets = []
    summary_lines = []

    for i, line in enumerate(lines[2:], start=2):
        line_lower = line.lower()

        if any(header in line_lower for header in section_headers):
            if current_section:
                data["sections"].append(
                    {
                        "id": str(len(data["sections"])),
                        "title": current_section,
                        "bullets": [
                            {"id": str(j), "text": b, "params": {}}
                            for j, b in enumerate(current_bullets)
                        ],
                    }
                )
            current_section = line
            current_bullets = []
        elif current_section:
            if line.startswith("•") or line.startswith("-") or line.startswith("*"):
                bullet_text = re.sub(r"^[•\-\*]\s*", "", line)
                current_bullets.append(bullet_text)
            elif len(line) > 20 and not any(
                header in line_lower for header in section_headers
            ):
                current_bullets.append(line)
        elif "summary" in line_lower:
            pass
        elif not current_section and len(line) > 30:
            summary_lines.append(line)

    if current_section:
        data["sections"].append(
            {
                "id": str(len(data["sections"])),
                "title": current_section,
                "bullets": [
                    {"id": str(j), "text": b, "params": {}}
                    for j, b in enumerate(current_bullets)
                ],
            }
        )

    if summary_lines:
        data["summary"] = " ".join(summary_lines)

    data["detected_variables"] = detect_parameterization(text)

    return data


def detect_parameterization(text: str) -> Dict[str, str]:
    """Detect parameterization variables in text"""
    variables = {}

    companies = re.findall(
        r"\b(?:at|for|with)\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)\b", text
    )
    if companies:
        most_common = max(set(companies), key=companies.count)
        variables["{{company}}"] = most_common

    percentages = re.findall(r"(\d+)%", text)
    if percentages:
        variables["{{metric}}"] = percentages[0]

    tech_keywords = [
        "AWS",
        "Azure",
        "GCP",
        "Kubernetes",
        "Docker",
        "Python",
        "Java",
        "React",
        "Node.js",
    ]
    found_tech = [tech for tech in tech_keywords if tech in text]
    if found_tech:
        variables["{{tech}}"] = found_tech[0]

    return variables

