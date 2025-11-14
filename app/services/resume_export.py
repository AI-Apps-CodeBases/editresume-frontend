"""Resume export service - PDF and DOCX generation"""

from __future__ import annotations

import logging
from io import BytesIO
from typing import Optional

from fastapi import HTTPException
from fastapi.responses import Response
from sqlalchemy.orm import Session

from app.api.models import ExportPayload
from app.models import ExportAnalytics, Resume, User
from app.utils.resume_formatting import (
    apply_replacements,
    format_regular_bullets,
    format_work_experience_bullets,
    strip_bullet_markers,
)
from app.utils.resume_templates import TEMPLATES

logger = logging.getLogger(__name__)


async def export_pdf(
    payload: ExportPayload,
    user_email: Optional[str] = None,
    db: Optional[Session] = None,
) -> Response:
    """Export resume as PDF"""
    try:
        from weasyprint import HTML

        replacements = payload.replacements or {}
        template_id = payload.template or "tech"
        template_style = TEMPLATES.get(template_id, TEMPLATES["tech"])

        logger.info(f"Exporting PDF with template: {template_id}")
        logger.info(f"Template style: {template_style}")
        logger.info(
            f"Two-column settings: left={payload.two_column_left}, right={payload.two_column_right}, width={payload.two_column_left_width}"
        )

        font_family = template_style["styles"]["font"]
        header_align = template_style["styles"]["header_align"]
        header_border = template_style["styles"]["header_border"]
        section_uppercase = (
            "text-transform: uppercase;"
            if template_style["styles"]["section_uppercase"]
            else ""
        )
        layout = template_style["styles"]["layout"]

        # Build HTML content sections
        contact_info = apply_replacements(payload.email or "", replacements)
        if payload.phone:
            contact_info += " • " + apply_replacements(payload.phone, replacements)
        if payload.location:
            contact_info += " • " + apply_replacements(payload.location, replacements)

        # Add cover letter if provided
        cover_letter_html = ""
        if payload.cover_letter:
            cover_letter_content = payload.cover_letter.replace("\n", "<br>")
            cover_letter_html = f"""
            <div class="cover-letter-section">
                <h2>Cover Letter</h2>
                <div class="cover-letter-content">
                    {cover_letter_content}
                </div>
            </div>
            """

        # Build sections HTML
        if layout == "two-column":
            left_sections_html = ""
            right_sections_html = ""

            # Use localStorage configuration if provided, otherwise fallback to alternating
            left_section_ids = set(payload.two_column_left or [])
            right_section_ids = set(payload.two_column_right or [])

            # If no configuration provided, use alternating logic
            if not left_section_ids and not right_section_ids:
                for i, s in enumerate(payload.sections):
                    section_title = s.title.lower()
                    is_work_experience = (
                        "experience" in section_title
                        or "work" in section_title
                        or "employment" in section_title
                    )

                    if is_work_experience:
                        bullets_html = format_work_experience_bullets(
                            s.bullets, replacements
                        )
                        section_html = f"""
                        <div class="section">
                            <h2>{apply_replacements(s.title, replacements)}</h2>
                            {bullets_html}
                        </div>"""
                    else:
                        bullets_html = format_regular_bullets(s.bullets, replacements, s.title)
                        section_title = apply_replacements(s.title, replacements)
                        section_lower = s.title.lower()
                        is_skill_section = (
                            "skill" in section_lower
                            or "technical" in section_lower
                            or "technology" in section_lower
                            or "competencies" in section_lower
                            or "expertise" in section_lower
                            or "proficiencies" in section_lower
                        )
                        if is_skill_section:
                            section_html = f"""
                            <div class="section">
                                <h2>{section_title}</h2>
                                {bullets_html}
                            </div>"""
                        else:
                            section_html = f"""
                            <div class="section">
                                <h2>{section_title}</h2>
                                <ul>
                                    {bullets_html}
                                </ul>
                            </div>"""

                    if i % 2 == 0:
                        left_sections_html += section_html
                    else:
                        right_sections_html += section_html
            else:
                # Use provided configuration
                for s in payload.sections:
                    section_title = s.title.lower()
                    is_work_experience = (
                        "experience" in section_title
                        or "work" in section_title
                        or "employment" in section_title
                    )

                    if is_work_experience:
                        bullets_html = format_work_experience_bullets(
                            s.bullets, replacements
                        )
                        section_html = f"""
                        <div class="section">
                            <h2>{apply_replacements(s.title, replacements)}</h2>
                            {bullets_html}
                        </div>"""
                    else:
                        bullets_html = format_regular_bullets(s.bullets, replacements, s.title)
                        section_title = apply_replacements(s.title, replacements)
                        section_lower = s.title.lower()
                        is_skill_section = (
                            "skill" in section_lower
                            or "technical" in section_lower
                            or "technology" in section_lower
                            or "competencies" in section_lower
                            or "expertise" in section_lower
                            or "proficiencies" in section_lower
                        )
                        if is_skill_section:
                            section_html = f"""
                            <div class="section">
                                <h2>{section_title}</h2>
                                {bullets_html}
                            </div>"""
                        else:
                            section_html = f"""
                            <div class="section">
                                <h2>{section_title}</h2>
                                <ul>
                                    {bullets_html}
                                </ul>
                            </div>"""

                    if s.id and s.id in left_section_ids:
                        left_sections_html += section_html
                    elif s.id and s.id in right_section_ids:
                        right_sections_html += section_html

            # Calculate column widths with spacing
            left_width = payload.two_column_left_width or 50
            right_width = 100 - left_width

            # Adjust for spacing between columns (2% gap)
            gap = 2
            left_width_adjusted = left_width - (gap / 2)
            right_width_adjusted = right_width - (gap / 2)

            content_html = f"""
            {cover_letter_html}
            <div class="two-column">
                <div class="column" style="width: {left_width_adjusted}%; margin-right: {gap}%;">{left_sections_html}</div>
                <div class="column" style="width: {right_width_adjusted}%;">{right_sections_html}</div>
                <div class="clearfix"></div>
            </div>"""
        else:
            # Single column layout
            summary_html = (
                f'<div class="summary">{apply_replacements(payload.summary, replacements)}</div>'
                if payload.summary
                else ""
            )
            sections_html = ""
            for s in payload.sections:
                section_title = s.title.lower()
                is_work_experience = (
                    "experience" in section_title
                    or "work" in section_title
                    or "employment" in section_title
                )

                if is_work_experience:
                    bullets_html = format_work_experience_bullets(
                        s.bullets, replacements
                    )
                    sections_html += f"""
                    <div class="section">
                        <h2>{apply_replacements(s.title, replacements)}</h2>
                        {bullets_html}
                    </div>"""
                else:
                    bullets_html = format_regular_bullets(s.bullets, replacements, s.title)
                    section_title = apply_replacements(s.title, replacements)
                    section_lower = s.title.lower()
                    is_skill_section = (
                        "skill" in section_lower
                        or "technical" in section_lower
                        or "technology" in section_lower
                        or "competencies" in section_lower
                        or "expertise" in section_lower
                        or "proficiencies" in section_lower
                    )
                    if is_skill_section:
                        sections_html += f"""
                        <div class="section">
                            <h2>{section_title}</h2>
                            {bullets_html}
                        </div>"""
                    else:
                        sections_html += f"""
                        <div class="section">
                            <h2>{section_title}</h2>
                            <ul>
                                {bullets_html}
                            </ul>
                        </div>"""
            content_html = cover_letter_html + summary_html + sections_html

        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        @page {{ size: A4; margin: 2cm; }}
        body {{ font-family: {font_family}; font-size: 11pt; line-height: 1.4; color: #000000; }}
        .header {{ text-align: {header_align}; border-bottom: {header_border}; padding-bottom: 10px; margin-bottom: 15px; }}
        .header h1 {{ margin: 0; font-size: 24pt; font-weight: bold; color: #000000; }}
        .header .title {{ font-size: 14pt; margin: 5px 0; color: #000000; }}
        .header .contact {{ font-size: 10pt; color: #000000; margin-top: 5px; }}
        .summary {{ margin-bottom: 15px; font-size: 10pt; line-height: 1.5; color: #000000; }}
        .section {{ margin-bottom: 15px; }}
        .section h2 {{ font-size: 12pt; font-weight: bold; color: #000000; {section_uppercase}
                       border-bottom: 1px solid #333; padding-bottom: 3px; margin-bottom: 8px; }}
        .section ul {{ margin: 0; padding-left: 0; list-style: none; }}
        .section li {{ margin-bottom: 6px; font-size: 10pt; color: #000000; position: relative; padding-left: 14px; }}
        .section li::before {{ content: "•"; font-weight: bold; color: #000000; position: absolute; left: 0; top: 0; }}
        .job-entry {{ margin-bottom: 20px; }}
        .company-header {{ margin-bottom: 8px; }}
        .company-name-line {{ font-weight: bold; font-size: 1.1em; color: #000000; margin-bottom: 3px; }}
        .company-title-line {{ display: flex; justify-content: space-between; align-items: center; font-size: 10pt; }}
        .job-title {{ font-weight: 500; color: #000000; }}
        .job-date {{ color: #666666; font-size: 9pt; }}
        .skills-section {{ font-size: 10pt; color: #000000; line-height: 1.6; }}
        .job-separator {{ height: 10px; }}
        .two-column {{ width: 100%; }}
        .column {{ float: left; }}
        .clearfix {{ clear: both; }}
        .cover-letter-section {{ margin-bottom: 20px; page-break-after: always; }}
        .cover-letter-section h2 {{ font-size: 14pt; font-weight: bold; margin-bottom: 10px; border-bottom: 2px solid #333; padding-bottom: 5px; }}
        .cover-letter-content {{ font-size: 11pt; line-height: 1.6; text-align: justify; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{apply_replacements(payload.name, replacements)}</h1>
        <div class="title">{apply_replacements(payload.title, replacements)}</div>
        <div class="contact">{contact_info}</div>
    </div>
    {content_html}
</body>
</html>"""

        logger.info(f"Generated HTML length: {len(html_content)}")
        logger.info(f"Layout type: {layout}")

        pdf_bytes = HTML(string=html_content).write_pdf()
        logger.info(f"PDF generated successfully, size: {len(pdf_bytes)} bytes")

        # Track export analytics
        if user_email and db:
            try:
                user = db.query(User).filter(User.email == user_email).first()
                if user:
                    # Find or create resume record
                    resume = (
                        db.query(Resume)
                        .filter(Resume.user_id == user.id, Resume.name == payload.name)
                        .first()
                    )

                    if not resume:
                        resume = Resume(
                            user_id=user.id,
                            name=payload.name,
                            title=payload.title,
                            email=payload.email,
                            phone=payload.phone,
                            location=payload.location,
                            summary=payload.summary,
                            template=template_id,
                        )
                        db.add(resume)
                        db.commit()
                        db.refresh(resume)

                    # Create export analytics record
                    export_analytics = ExportAnalytics(
                        user_id=user.id,
                        resume_id=resume.id,
                        export_format="pdf",
                        template_used=template_id,
                        file_size=len(pdf_bytes),
                        export_success=True,
                    )
                    db.add(export_analytics)
                    db.commit()

                    logger.info(
                        f"Export analytics tracked for user {user_email}: PDF export"
                    )
            except Exception as e:
                logger.error(f"Failed to track export analytics: {e}")

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=resume.pdf"},
        )
    except Exception as e:
        logger.error(f"PDF export error: {str(e)}")
        import traceback

        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


async def export_docx(
    payload: ExportPayload,
    user_email: Optional[str] = None,
    db: Optional[Session] = None,
) -> Response:
    """Export resume as DOCX"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        replacements = payload.replacements or {}
        template_id = payload.template or "tech"
        template_style = TEMPLATES.get(template_id, TEMPLATES["tech"])

        logger.info(f"Exporting DOCX with template: {template_id}")
        logger.info(f"Template layout: {template_style['styles']['layout']}")
        logger.info(
            f"Two-column settings: left={payload.two_column_left}, right={payload.two_column_right}, width={payload.two_column_left_width}"
        )

        layout = template_style["styles"]["layout"]

        doc = Document()

        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        name_para = doc.add_paragraph()
        name_run = name_para.add_run(apply_replacements(payload.name, replacements))
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.name = (
            "Arial" if "Arial" in template_style["styles"]["font"] else "Georgia"
        )
        if template_style["styles"]["header_align"] == "center":
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_para = doc.add_paragraph()
        title_run = title_para.add_run(apply_replacements(payload.title, replacements))
        title_run.font.size = Pt(12)
        title_run.font.name = (
            "Arial" if "Arial" in template_style["styles"]["font"] else "Georgia"
        )
        if template_style["styles"]["header_align"] == "center":
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact_parts = []
        if payload.email:
            contact_parts.append(apply_replacements(payload.email, replacements))
        if payload.phone:
            contact_parts.append(apply_replacements(payload.phone, replacements))
        if payload.location:
            contact_parts.append(apply_replacements(payload.location, replacements))

        if contact_parts:
            contact_para = doc.add_paragraph(" • ".join(contact_parts))
            contact_para.runs[0].font.size = Pt(10)
            if template_style["styles"]["header_align"] == "center":
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Add cover letter if provided
        if payload.cover_letter:
            cover_letter_heading = doc.add_paragraph()
            cover_letter_heading_run = cover_letter_heading.add_run("Cover Letter")
            cover_letter_heading_run.font.size = Pt(14)
            cover_letter_heading_run.font.bold = True

            cover_letter_para = doc.add_paragraph(payload.cover_letter)
            cover_letter_para.runs[0].font.size = Pt(11)
            cover_letter_para.runs[0].font.name = "Arial"

            doc.add_paragraph()  # Add spacing

        logger.info(
            f"Layout decision: layout='{layout}', is_two_column={layout == 'two-column'}"
        )
        if layout == "two-column":
            logger.info("Creating two-column layout")
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            # Create a table for two-column layout
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False

            # Remove table borders
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl.insert(0, tblPr)

            tblBorders = OxmlElement("w:tblBorders")
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "none")
                border.set(qn("w:sz"), "0")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "auto")
                tblBorders.append(border)
            tblPr.append(tblBorders)

            # Set column widths based on user preference
            left_width_percent = payload.two_column_left_width or 50
            right_width_percent = 100 - left_width_percent

            # Convert percentages to inches (assuming 8.5" total width)
            left_width_inches = (left_width_percent / 100) * 8.5
            right_width_inches = (right_width_percent / 100) * 8.5

            table.columns[0].width = Inches(left_width_inches)
            table.columns[1].width = Inches(right_width_inches)

            logger.info(
                f'Column widths: left={left_width_inches:.1f}", right={right_width_inches:.1f}"'
            )

            # Left column
            left_cell = table.cell(0, 0)
            # Clear existing paragraphs in left cell
            for paragraph in left_cell.paragraphs:
                paragraph.clear()

            # Right column
            right_cell = table.cell(0, 1)
            # Clear existing paragraphs in right cell
            for paragraph in right_cell.paragraphs:
                paragraph.clear()

            # Use localStorage configuration if provided, otherwise fallback to alternating
            left_section_ids = set(payload.two_column_left or [])
            right_section_ids = set(payload.two_column_right or [])

            logger.info(
                f"Section IDs from payload: left={left_section_ids}, right={right_section_ids}"
            )
            logger.info(
                f"Available sections: {[s.id for s in payload.sections if s.id]}"
            )

            if not left_section_ids and not right_section_ids:
                # Use alternating logic if no configuration
                logger.info(
                    "Using alternating logic - no localStorage configuration provided"
                )
                left_sections = [
                    s for i, s in enumerate(payload.sections) if i % 2 == 0
                ]
                right_sections = [
                    s for i, s in enumerate(payload.sections) if i % 2 == 1
                ]
            else:
                # Use provided configuration
                logger.info("Using localStorage configuration")
                left_sections = [
                    s for s in payload.sections if s.id and s.id in left_section_ids
                ]
                right_sections = [
                    s for s in payload.sections if s.id and s.id in right_section_ids
                ]

            logger.info(
                f"Final section distribution: left={[s.title for s in left_sections]}, right={[s.title for s in right_sections]}"
            )
            logger.info(
                f"Left sections count: {len(left_sections)}, Right sections count: {len(right_sections)}"
            )

            # Add sections to left column
            for s in left_sections:
                section_title = apply_replacements(s.title, replacements)
                if template_style["styles"]["section_uppercase"]:
                    section_title = section_title.upper()

                heading_para = left_cell.add_paragraph()
                heading_run = heading_para.add_run(section_title)
                heading_run.font.size = Pt(12)
                heading_run.font.bold = True
                heading_run.font.name = (
                    "Arial"
                    if "Arial" in template_style["styles"]["font"]
                    else "Georgia"
                )

                # Check section type
                section_lower = s.title.lower()
                is_work_experience = (
                    "experience" in section_lower
                    or "work" in section_lower
                    or "employment" in section_lower
                )
                is_skill_section = (
                    "skill" in section_lower
                    or "technical" in section_lower
                    or "technology" in section_lower
                    or "competencies" in section_lower
                    or "expertise" in section_lower
                    or "proficiencies" in section_lower
                )

                # Create content for this section
                if s.bullets:
                    if is_work_experience:
                        # Work experience - handle company headers
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            
                            # Check if company header
                            if bullet_text.startswith("**") and "**" in bullet_text[2:]:
                                header_text = bullet_text.replace("**", "").strip()
                                parts = header_text.split(' / ')
                                company_name = parts[0] if parts else ''
                                location = parts[1] if len(parts) >= 4 else ''
                                title = parts[2] if len(parts) >= 4 else (parts[1] if len(parts) >= 2 else '')
                                date_range = parts[3] if len(parts) >= 4 else (parts[2] if len(parts) >= 3 else '')
                                
                                # Company header - Line 1
                                company_para = left_cell.add_paragraph()
                                company_line = company_name
                                if location:
                                    company_line += f' / {location}'
                                company_run = company_para.add_run(company_line)
                                company_run.font.size = Pt(11)
                                company_run.font.bold = True
                                
                                # Company header - Line 2
                                title_para = left_cell.add_paragraph()
                                title_run = title_para.add_run(title)
                                title_run.font.size = Pt(10)
                                title_run.font.bold = True
                                date_run = title_para.add_run(f'\t{date_range}')
                                date_run.font.size = Pt(9)
                                date_run.font.italic = True
                            else:
                                # Regular bullet - strip all bullet markers
                                clean_text = strip_bullet_markers(bullet_text)
                                if clean_text:
                                    bullet_para = left_cell.add_paragraph()
                                    bullet_para.style = "List Bullet"
                                    run = bullet_para.add_run(clean_text)
                                    run.font.size = Pt(10)
                    elif is_skill_section:
                        # Skills section - comma-separated, no bullets
                        skill_items = []
                        for b in s.bullets:
                            if b.text.strip():
                                bullet_text = apply_replacements(b.text, replacements)
                                clean_text = strip_bullet_markers(bullet_text)
                                if clean_text:
                                    skill_items.append(clean_text)
                        if skill_items:
                            skills_para = left_cell.add_paragraph(", ".join(skill_items))
                            skills_para.runs[0].font.size = Pt(10)
                    else:
                        # Regular section with bullets
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            # Strip all bullet markers
                            clean_text = strip_bullet_markers(bullet_text)
                            if clean_text:
                                bullet_para = left_cell.add_paragraph()
                                bullet_para.style = "List Bullet"
                                run = bullet_para.add_run(clean_text)
                                run.font.size = Pt(10)

                left_cell.add_paragraph()

            # Add sections to right column
            for s in right_sections:
                section_title = apply_replacements(s.title, replacements)
                if template_style["styles"]["section_uppercase"]:
                    section_title = section_title.upper()

                heading_para = right_cell.add_paragraph()
                heading_run = heading_para.add_run(section_title)
                heading_run.font.size = Pt(12)
                heading_run.font.bold = True
                heading_run.font.name = (
                    "Arial"
                    if "Arial" in template_style["styles"]["font"]
                    else "Georgia"
                )

                # Check section type
                section_lower = s.title.lower()
                is_work_experience = (
                    "experience" in section_lower
                    or "work" in section_lower
                    or "employment" in section_lower
                )
                is_skill_section = (
                    "skill" in section_lower
                    or "technical" in section_lower
                    or "technology" in section_lower
                    or "competencies" in section_lower
                    or "expertise" in section_lower
                    or "proficiencies" in section_lower
                )

                # Create content for this section
                if s.bullets:
                    if is_work_experience:
                        # Work experience - handle company headers
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            
                            # Check if company header
                            if bullet_text.startswith("**") and "**" in bullet_text[2:]:
                                header_text = bullet_text.replace("**", "").strip()
                                parts = header_text.split(' / ')
                                company_name = parts[0] if parts else ''
                                location = parts[1] if len(parts) >= 4 else ''
                                title = parts[2] if len(parts) >= 4 else (parts[1] if len(parts) >= 2 else '')
                                date_range = parts[3] if len(parts) >= 4 else (parts[2] if len(parts) >= 3 else '')
                                
                                # Company header - Line 1
                                company_para = right_cell.add_paragraph()
                                company_line = company_name
                                if location:
                                    company_line += f' / {location}'
                                company_run = company_para.add_run(company_line)
                                company_run.font.size = Pt(11)
                                company_run.font.bold = True
                                
                                # Company header - Line 2
                                title_para = right_cell.add_paragraph()
                                title_run = title_para.add_run(title)
                                title_run.font.size = Pt(10)
                                title_run.font.bold = True
                                date_run = title_para.add_run(f'\t{date_range}')
                                date_run.font.size = Pt(9)
                                date_run.font.italic = True
                            else:
                                # Regular bullet - strip existing bullets
                                clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                                clean_text = clean_text.lstrip('•*- ').strip()
                                if clean_text:
                                    bullet_para = right_cell.add_paragraph()
                                    bullet_para.style = "List Bullet"
                                    run = bullet_para.add_run(clean_text)
                                    run.font.size = Pt(10)
                    elif is_skill_section:
                        # Skills section - comma-separated, no bullets
                        skill_items = []
                        for b in s.bullets:
                            if b.text.strip():
                                bullet_text = apply_replacements(b.text, replacements)
                                clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                                clean_text = clean_text.lstrip('•*- ').strip()
                                if clean_text:
                                    skill_items.append(clean_text)
                        if skill_items:
                            skills_para = right_cell.add_paragraph(", ".join(skill_items))
                            skills_para.runs[0].font.size = Pt(10)
                    else:
                        # Regular section with bullets
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            # Strip existing bullets
                            clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                            clean_text = clean_text.lstrip('•*- ').strip()
                            if clean_text:
                                bullet_para = right_cell.add_paragraph()
                                bullet_para.style = "List Bullet"
                                run = bullet_para.add_run(clean_text)
                                run.font.size = Pt(10)

                right_cell.add_paragraph()
        else:
            # Single column layout
            if payload.summary:
                summary_para = doc.add_paragraph(
                    apply_replacements(payload.summary, replacements)
                )
                summary_para.runs[0].font.size = Pt(10)
                doc.add_paragraph()

            for s in payload.sections:
                section_title = apply_replacements(s.title, replacements)
                if template_style["styles"]["section_uppercase"]:
                    section_title = section_title.upper()

                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run(section_title)
                heading_run.font.size = Pt(12)
                heading_run.font.bold = True
                heading_run.font.name = (
                    "Arial"
                    if "Arial" in template_style["styles"]["font"]
                    else "Georgia"
                )

                # Check section type
                section_lower = s.title.lower()
                is_work_experience = (
                    "experience" in section_lower
                    or "work" in section_lower
                    or "employment" in section_lower
                )
                is_skill_section = (
                    "skill" in section_lower
                    or "technical" in section_lower
                    or "technology" in section_lower
                    or "competencies" in section_lower
                    or "expertise" in section_lower
                    or "proficiencies" in section_lower
                )

                # Create content for this section
                if s.bullets:
                    if is_work_experience:
                        # Work experience - handle company headers
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            
                            # Check if company header
                            if bullet_text.startswith("**") and "**" in bullet_text[2:]:
                                header_text = bullet_text.replace("**", "").strip()
                                parts = header_text.split(' / ')
                                company_name = parts[0] if parts else ''
                                location = parts[1] if len(parts) >= 4 else ''
                                title = parts[2] if len(parts) >= 4 else (parts[1] if len(parts) >= 2 else '')
                                date_range = parts[3] if len(parts) >= 4 else (parts[2] if len(parts) >= 3 else '')
                                
                                # Company header - Line 1
                                company_para = doc.add_paragraph()
                                company_line = company_name
                                if location:
                                    company_line += f' / {location}'
                                company_run = company_para.add_run(company_line)
                                company_run.font.size = Pt(11)
                                company_run.font.bold = True
                                
                                # Company header - Line 2
                                title_para = doc.add_paragraph()
                                title_run = title_para.add_run(title)
                                title_run.font.size = Pt(10)
                                title_run.font.bold = True
                                date_run = title_para.add_run(f'\t{date_range}')
                                date_run.font.size = Pt(9)
                                date_run.font.italic = True
                            else:
                                # Regular bullet - strip existing bullets
                                clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                                clean_text = clean_text.lstrip('•*- ').strip()
                                if clean_text:
                                    bullet_para = doc.add_paragraph()
                                    bullet_para.style = "List Bullet"
                                    run = bullet_para.add_run(clean_text)
                                    run.font.size = Pt(10)
                    elif is_skill_section:
                        # Skills section - comma-separated, no bullets
                        skill_items = []
                        for b in s.bullets:
                            if b.text.strip():
                                bullet_text = apply_replacements(b.text, replacements)
                                clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                                clean_text = clean_text.lstrip('•*- ').strip()
                                if clean_text:
                                    skill_items.append(clean_text)
                        if skill_items:
                            skills_para = doc.add_paragraph(", ".join(skill_items))
                            skills_para.runs[0].font.size = Pt(10)
                    else:
                        # Regular section with bullets
                        for b in s.bullets:
                            if not b.text.strip():
                                continue
                            bullet_text = apply_replacements(b.text, replacements)
                            # Strip existing bullets
                            clean_text = bullet_text.replace("•", "").replace("*", "").replace("-", "").strip()
                            clean_text = clean_text.lstrip('•*- ').strip()
                            if clean_text:
                                bullet_para = doc.add_paragraph()
                                bullet_para.style = "List Bullet"
                                run = bullet_para.add_run(clean_text)
                                run.font.size = Pt(10)

                doc.add_paragraph()

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        docx_bytes = buffer.getvalue()
        logger.info(f"DOCX generated successfully, size: {len(docx_bytes)} bytes")

        # Track export analytics
        if user_email and db:
            try:
                user = db.query(User).filter(User.email == user_email).first()
                if user:
                    # Find or create resume record
                    resume = (
                        db.query(Resume)
                        .filter(Resume.user_id == user.id, Resume.name == payload.name)
                        .first()
                    )

                    if not resume:
                        resume = Resume(
                            user_id=user.id,
                            name=payload.name,
                            title=payload.title,
                            email=payload.email,
                            phone=payload.phone,
                            location=payload.location,
                            summary=payload.summary,
                            template=template_id,
                        )
                        db.add(resume)
                        db.commit()
                        db.refresh(resume)

                    # Create export analytics record
                    export_analytics = ExportAnalytics(
                        user_id=user.id,
                        resume_id=resume.id,
                        export_format="docx",
                        template_used=template_id,
                        file_size=len(docx_bytes),
                        export_success=True,
                    )
                    db.add(export_analytics)
                    db.commit()

                    logger.info(
                        f"Export analytics tracked for user {user_email}: DOCX export"
                    )
            except Exception as e:
                logger.error(f"Failed to track export analytics: {e}")

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=resume.docx"},
        )
    except Exception as e:
        logger.error(f"DOCX export error: {str(e)}")
        import traceback

        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

