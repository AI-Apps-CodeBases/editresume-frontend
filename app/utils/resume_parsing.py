"""Resume parsing utilities - extracted from legacy_app.py"""

from __future__ import annotations

import logging
import re
from typing import Tuple

logger = logging.getLogger(__name__)


def extract_pdf_text(file_content: bytes) -> Tuple[str, list[str]]:
    """Extract text from PDF using multiple methods"""
    text = ""
    methods = []

    # Method 1: PyMuPDF (fitz) - Most reliable
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_content, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        methods.append("PyMuPDF")
        logger.info(f"PyMuPDF extracted {len(text)} characters")
    except Exception as e:
        logger.warning(f"PyMuPDF failed: {e}")

    # Method 2: pdfplumber - Good for complex layouts
    if len(text.strip()) < 100:
        try:
            import pdfplumber
            import io

            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            methods.append("pdfplumber")
            logger.info(f"pdfplumber extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

    # Method 3: PyPDF2 - Fallback
    if len(text.strip()) < 100:
        try:
            import PyPDF2
            import io

            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            methods.append("PyPDF2")
            logger.info(f"PyPDF2 extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")

    return text, methods


def extract_docx_text(file_content: bytes) -> Tuple[str, list[str]]:
    """Extract text from DOCX using multiple methods"""
    text = ""
    methods = []

    # Method 1: python-docx - Standard method
    try:
        from docx import Document
        import io

        doc = Document(io.BytesIO(file_content))

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

        methods.append("python-docx")
        logger.info(f"python-docx extracted {len(text)} characters")
    except Exception as e:
        logger.warning(f"python-docx failed: {e}")

    # Method 2: docx2txt - Alternative method
    if len(text.strip()) < 100:
        try:
            import docx2txt
            import io

            text = docx2txt.process(io.BytesIO(file_content))
            methods.append("docx2txt")
            logger.info(f"docx2txt extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"docx2txt failed: {e}")

    return text, methods


def extract_doc_text(file_content: bytes) -> Tuple[str, list[str]]:
    """Extract text from DOC files"""
    text = ""
    methods = []

    # Method 1: antiword (if available)
    try:
        import subprocess
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name

        try:
            result = subprocess.run(
                ["antiword", temp_file_path],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0:
                text = result.stdout
                methods.append("antiword")
                logger.info(f"antiword extracted {len(text)} characters")
        finally:
            os.unlink(temp_file_path)
    except Exception as e:
        logger.warning(f"antiword failed: {e}")

    # Method 2: Try reading as text (fallback)
    if len(text.strip()) < 100:
        try:
            text = file_content.decode("utf-8", errors="ignore")
            methods.append("UTF-8 decode")
            logger.info(f"UTF-8 decode extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"UTF-8 decode failed: {e}")

    return text, methods


def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""

    # Remove excessive whitespace
    text = re.sub(r"\n\s*\n\s*\n", "\n\n", text)  # Max 2 newlines
    text = re.sub(r"[ \t]+", " ", text)  # Multiple spaces to single

    # Remove common PDF artifacts
    text = re.sub(r"[^\x00-\x7F]+", " ", text)  # Remove non-ASCII
    text = re.sub(r"\f", "\n", text)  # Form feeds to newlines

    # Clean up bullet points
    text = re.sub(r"[•·▪▫‣⁃]", "•", text)  # Normalize bullets

    return text.strip()


def parse_resume_with_regex(text: str) -> dict:
    """Simple, reliable resume parser that actually works"""
    logger.info("Using simple regex-based resume parsing")

    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # Extract basic info
    name = ""
    title = ""
    email = ""
    phone = ""
    location = ""

    # Find name (first line that looks like a name)
    for line in lines[:3]:
        if (
            len(line) > 2
            and not any(char.isdigit() for char in line)
            and "@" not in line
            and "•" not in line
        ):
            name = line
            break

    # Find title (line after name)
    for i, line in enumerate(lines[1:4]):
        if (
            len(line) > 2
            and not any(char.isdigit() for char in line)
            and "@" not in line
            and "•" not in line
            and line != name
        ):
            title = line
            break

    # Find email
    email_match = re.search(
        r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", text
    )
    if email_match:
        email = email_match.group()

    # Find phone
    phone_match = re.search(
        r"(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})", text
    )
    if phone_match:
        phone = phone_match.group().strip()

    # Find location
    location_match = re.search(r"([A-Za-z\s]+),\s*([A-Z]{2})", text)
    if location_match:
        location = location_match.group().strip()

    # Extract sections
    sections = []
    current_section = None
    current_bullets = []

    section_keywords = [
        "work experience",
        "professional experience",
        "employment",
        "experience",
        "projects",
        "education",
        "skills",
        "certifications",
    ]

    for line in lines:
        line_lower = line.lower()

        # Check if this is a section header
        is_section = any(keyword in line_lower for keyword in section_keywords)

        if is_section:
            # Save previous section
            if current_section and current_bullets:
                sections.append(
                    {
                        "title": current_section,
                        "bullets": [
                            {
                                "id": f"{len(sections)}-{i}",
                                "text": bullet,
                                "params": {},
                            }
                            for i, bullet in enumerate(current_bullets)
                        ],
                        "id": str(len(sections)),
                    }
                )

            # Start new section
            current_section = line
            current_bullets = []

        elif current_section and line:
            # This is content for current section
            if line.startswith(("•", "-", "*")):
                bullet_text = line[1:].strip()
                if bullet_text:
                    current_bullets.append(f"• {bullet_text}")
            elif " / " in line and any(char.isdigit() for char in line):
                # Job entry
                current_bullets.append(f"**{line}**")
            else:
                current_bullets.append(line)

    # Save last section
    if current_section and current_bullets:
        sections.append(
            {
                "title": current_section,
                "bullets": [
                    {"id": f"{len(sections)}-{i}", "text": bullet, "params": {}}
                    for i, bullet in enumerate(current_bullets)
                ],
                "id": str(len(sections)),
            }
        )

    return {
        "name": name,
        "title": title,
        "email": email,
        "phone": phone,
        "location": location,
        "summary": "",
        "sections": sections,
        "detected_variables": {},
    }

