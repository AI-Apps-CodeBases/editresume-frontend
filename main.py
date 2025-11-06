from fastapi import FastAPI, UploadFile, File, HTTPException, WebSocket, WebSocketDisconnect, Depends, Request, Query, Header
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import re
import json
import logging
import secrets
from datetime import datetime
from openai import OpenAI
from keyword_extractor import KeywordExtractor
from grammar_checker import GrammarStyleChecker
from sqlalchemy.orm import Session
from database import get_db, create_tables, migrate_schema, User, Resume, ResumeVersion, ExportAnalytics, JobMatch, SharedResume, ResumeView, SharedResumeComment, DATABASE_URL, JobDescription, MatchSession
from version_control import VersionControlService

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import ATS checker with fallback
try:
    from ats_checker import ATSChecker
    ats_checker = ATSChecker()
    logger.info("ATS checker initialized successfully")
except ImportError as e:
    logger.warning(f"ATS checker not available: {e}")
    ats_checker = None
except Exception as e:
    logger.warning(f"ATS checker failed to initialize: {e}")
    ats_checker = None

# Import Enhanced ATS checker
try:
    from enhanced_ats_checker import EnhancedATSChecker
    enhanced_ats_checker = EnhancedATSChecker()
    logger.info("Enhanced ATS checker initialized successfully")
except ImportError as e:
    logger.warning(f"Enhanced ATS checker not available: {e}")
    enhanced_ats_checker = None
except Exception as e:
    logger.warning(f"Enhanced ATS checker failed to initialize: {e}")
    enhanced_ats_checker = None

# Import AI Improvement Engine
try:
    from ai_improvement_engine import AIResumeImprovementEngine
    ai_improvement_engine = AIResumeImprovementEngine()
    logger.info("AI Improvement Engine initialized successfully")
except ImportError as e:
    logger.warning(f"AI Improvement Engine not available: {e}")
    ai_improvement_engine = None
except Exception as e:
    logger.warning(f"AI Improvement Engine failed to initialize: {e}")
    ai_improvement_engine = None

# OpenAI Configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
OPENAI_MAX_TOKENS = int(os.getenv("OPENAI_MAX_TOKENS", "2000"))

openai_client = None
if OPENAI_API_KEY and OPENAI_API_KEY != "sk-your-openai-api-key-here":
    try:
        # Use requests library directly to avoid version compatibility issues
        import requests
        openai_client = {
            'api_key': OPENAI_API_KEY,
            'model': OPENAI_MODEL,
            'requests': requests
        }
        logger.info(f"OpenAI client initialized successfully with model: {OPENAI_MODEL}")
    except Exception as e:
        logger.error(f"Failed to initialize OpenAI client: {e}")
        logger.warning("AI features will be disabled due to client initialization error")
        openai_client = None
else:
    if not OPENAI_API_KEY:
        logger.warning("OpenAI API key not found. AI features will be disabled.")
    else:
        logger.warning("OpenAI API key is still the placeholder value. Please set a real API key.")

# Initialize keyword extractor and grammar checker
keyword_extractor = KeywordExtractor()
grammar_checker = GrammarStyleChecker()

# Initialize database
create_tables()
migrate_schema()

app = FastAPI(title="editresume.io API", version="0.1.0")

# CORS Configuration
# Base allowed origins
BASE_ALLOWED_ORIGINS = [
    "http://localhost:3000",  # Local development
    "http://localhost:3001",  # Alternative local port
    "https://staging.editresume.io",  # Staging frontend
    "https://editresume.io",  # Production frontend
    "https://www.editresume.io",  # Production with www
    "https://editresume-staging.onrender.com",  # Staging backend (for testing)
    "https://editresume-staging-d4ang4wye-hasans-projects-d7f2163d.vercel.app",  # Vercel staging frontend
    "https://editresume-staging-git-fixuploadissue-hasans-projects-d7f2163d.vercel.app",  # Vercel staging frontend (branch)
]

# Add additional origins from environment variable
ADDITIONAL_ORIGINS = os.getenv("ADDITIONAL_CORS_ORIGINS", "").split(",")
ADDITIONAL_ORIGINS = [origin.strip() for origin in ADDITIONAL_ORIGINS if origin.strip()]

# Combine all allowed origins
ALLOWED_ORIGINS = BASE_ALLOWED_ORIGINS + ADDITIONAL_ORIGINS

# Log CORS configuration for debugging
print(f"CORS Allowed Origins: {ALLOWED_ORIGINS}")

# Custom function to check if origin is allowed
def is_origin_allowed(origin: str) -> bool:
    if not origin:
        return False
    
    # Check exact matches first
    if origin in ALLOWED_ORIGINS:
        return True
    
    # For staging environment, allow Vercel domains with our project ID
    ENVIRONMENT = os.getenv("ENVIRONMENT", "development")
    if ENVIRONMENT == "staging":
        if ("vercel.app" in origin and 
            "hasans-projects-d7f2163d" in origin and 
            ("editresume-staging" in origin or "editresume-staging-git" in origin)):
            return True
    
    return False

# For staging environment, use more permissive CORS
ENVIRONMENT = os.getenv("ENVIRONMENT", "development")
if ENVIRONMENT == "staging":
    # Allow all origins for staging (more permissive)
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],  # Allow all origins in staging
        allow_credentials=False,  # Must be False when allow_origins=["*"]
        allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
        allow_headers=["*"],
        expose_headers=["*"],
    )
else:
    # Production uses strict CORS
    app.add_middleware(
        CORSMiddleware,
        allow_origins=ALLOWED_ORIGINS,
        allow_credentials=True,
        allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
        allow_headers=["*"],
        expose_headers=["*"],
    )

# Add explicit OPTIONS handler for CORS preflight requests
@app.options("/{path:path}")
async def options_handler(path: str, request: Request):
    origin = request.headers.get("origin")
    headers = {
        "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS, PATCH",
        "Access-Control-Allow-Headers": "*",
    }
    
    # Allow chrome-extension origins (extensions bypass CORS but we handle it here)
    if origin and origin.startswith("chrome-extension://"):
        headers["Access-Control-Allow-Origin"] = origin
        headers["Access-Control-Allow-Credentials"] = "true"
    elif ENVIRONMENT == "staging":
        headers["Access-Control-Allow-Origin"] = "*"
        headers["Access-Control-Allow-Credentials"] = "false"
    elif origin and is_origin_allowed(origin):
        headers["Access-Control-Allow-Origin"] = origin
        headers["Access-Control-Allow-Credentials"] = "true"
    else:
        headers["Access-Control-Allow-Origin"] = "*"
        headers["Access-Control-Allow-Credentials"] = "false"
    
    return Response(status_code=200, headers=headers)

class BulletParam(BaseModel):
    id: Optional[str] = None
    text: str
    params: Optional[Dict[str, str]] = {}

class Section(BaseModel):
    id: Optional[str] = None
    title: str
    bullets: List[BulletParam] = []

class ResumePayload(BaseModel):
    name: str
    title: str
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    summary: Optional[str] = None
    sections: List[Section] = []
    variant: Optional[str] = None

class ExportPayload(BaseModel):
    name: str
    title: str
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    summary: Optional[str] = None
    sections: List[Section] = []
    replacements: Optional[Dict[str, str]] = None
    template: Optional[str] = "tech"
    two_column_left: Optional[List[str]] = []
    two_column_right: Optional[List[str]] = []
    two_column_left_width: Optional[int] = 50
    cover_letter: Optional[str] = None

class LoginPayload(BaseModel):
    email: str
    password: str

class SignupPayload(BaseModel):
    email: str
    password: str
    name: str

class ImproveBulletPayload(BaseModel):
    bullet: str
    context: Optional[str] = None
    tone: Optional[str] = "professional"

class GenerateBulletPointsPayload(BaseModel):
    role: str
    company: str
    skills: str
    count: int = 5
    tone: Optional[str] = "professional"

class GenerateSummaryPayload(BaseModel):
    role: str
    years_experience: int
    skills: str
    achievements: Optional[str] = None

class JobDescriptionMatchPayload(BaseModel):
    job_description: str
    resume_data: ResumePayload

class CoverLetterPayload(BaseModel):
    job_description: str
    resume_data: ResumePayload
    company_name: str
    position_title: str
    tone: str = "professional"
    custom_requirements: Optional[str] = None

class GrammarCheckPayload(BaseModel):
    text: str
    check_type: str = "all"  # "grammar", "style", "all"

class EnhancedATSPayload(BaseModel):
    resume_data: ResumePayload
    job_description: Optional[str] = None
    target_role: Optional[str] = None
    industry: Optional[str] = None

class AIImprovementPayload(BaseModel):
    resume_data: ResumePayload
    job_description: Optional[str] = None
    target_role: Optional[str] = None
    industry: Optional[str] = None
    strategy: Optional[str] = None  # Specific improvement strategy to focus on

# Version Control Models
class CreateVersionPayload(BaseModel):
    resume_id: int
    resume_data: Dict[str, Any]
    change_summary: Optional[str] = None
    is_auto_save: bool = False

class RollbackVersionPayload(BaseModel):
    version_id: int

class CompareVersionsPayload(BaseModel):
    version1_id: int
    version2_id: int

class SaveResumePayload(BaseModel):
    name: str
    title: str
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    summary: Optional[str] = None
    sections: List[Section] = []
    template: Optional[str] = "tech"

users_db = {}
user_stats = {}
payment_history = {}
PREMIUM_MODE = os.getenv("PREMIUM_MODE", "false").lower() == "true"

@app.get("/health")
async def health():
    openai_status = {
        "configured": OPENAI_API_KEY is not None,
        "model": OPENAI_MODEL if OPENAI_API_KEY else None,
        "client_ready": openai_client is not None
    }
    return {
        "status": "ok", 
        "db": DATABASE_URL, 
        "premium_mode": PREMIUM_MODE,
        "openai": openai_status
    }

@app.get("/api/test/db")
async def test_db(db: Session = Depends(get_db)):
    """Test database connection and table creation"""
    try:
        # Test if we can query users table
        user_count = db.query(User).count()
        
        # Test if we can create a test user
        test_user = User(
            email="test-db@example.com",
            name="Test DB User",
            password="test123",
            is_premium=False
        )
        db.add(test_user)
        db.commit()
        db.refresh(test_user)
        
        # Clean up test user
        db.delete(test_user)
        db.commit()
        
        return {
            "status": "success",
            "message": "Database connection working",
            "user_count": user_count,
            "test_user_created": True
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"Database error: {str(e)}"
        }

@app.post("/api/auth/signup")
async def signup(payload: SignupPayload, db: Session = Depends(get_db)):
    # Check if user already exists
    existing_user = db.query(User).filter(User.email == payload.email).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create new user
    new_user = User(
        email=payload.email,
        name=payload.name,
        password=payload.password,  # In production, hash this password
        is_premium=not PREMIUM_MODE
    )
    
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    
    token = secrets.token_urlsafe(32)
    
    logger.info(f"New user signup: {payload.email} (Premium: {new_user.is_premium})")
    return {
        "token": token,
        "user": {
            "email": new_user.email,
            "name": new_user.name,
            "isPremium": new_user.is_premium,
            "createdAt": new_user.created_at.isoformat()
        },
        "message": "Account created successfully"
    }

@app.post("/api/auth/login")
async def login(payload: LoginPayload, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == payload.email).first()
    
    if not user or user.password != payload.password:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = secrets.token_urlsafe(32)
    logger.info(f"User login: {payload.email}")
    
    return {
        "token": token,
        "user": {
            "email": user.email,
            "name": user.name,
            "isPremium": user.is_premium,
            "createdAt": user.created_at.isoformat()
        },
        "message": "Login successful"
    }

@app.get("/api/user/profile")
async def get_profile(email: str):
    user = users_db.get(email)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    stats = user_stats.get(email, {
        "resumesCreated": 0,
        "exportsThisMonth": 0,
        "totalExports": 0
    })
    
    return {
        "user": {
            "email": user["email"],
            "name": user["name"],
            "isPremium": user["isPremium"],
            "createdAt": user.get("created_at")
        },
        "stats": stats
    }

@app.get("/api/user/payment-history")
async def get_payment_history(email: str):
    user = users_db.get(email)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    history = payment_history.get(email, [])
    
    if user["isPremium"] and not history:
        history = [{
            "id": "1",
            "date": datetime.now().strftime("%Y-%m-%d"),
            "amount": "$9.99",
            "status": "Paid",
            "plan": "Premium Monthly"
        }]
    
    return {"payments": history}

@app.post("/api/user/upgrade")
async def upgrade_to_premium(payload: dict):
    email = payload.get("email")
    user = users_db.get(email)
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if user["isPremium"]:
        raise HTTPException(status_code=400, detail="Already premium")
    
    user["isPremium"] = True
    
    if email not in payment_history:
        payment_history[email] = []
    
    payment_history[email].append({
        "id": str(len(payment_history[email]) + 1),
        "date": datetime.now().strftime("%Y-%m-%d"),
        "amount": "$9.99",
        "status": "Paid",
        "plan": "Premium Monthly"
    })
    
    logger.info(f"User upgraded to premium: {email}")
    
    return {
        "user": {
            "email": user["email"],
            "name": user["name"],
            "isPremium": user["isPremium"]
        },
        "message": "Upgraded to premium successfully"
    }

@app.post("/api/user/track-export")
async def track_export(payload: dict):
    email = payload.get("email")
    if not email:
        return {"status": "ok"}
    
    if email not in user_stats:
        user_stats[email] = {
            "resumesCreated": 0,
            "exportsThisMonth": 0,
            "totalExports": 0
        }
    
    user_stats[email]["exportsThisMonth"] += 1
    user_stats[email]["totalExports"] += 1
    
    return {"status": "ok", "stats": user_stats[email]}

@app.delete("/api/user/account")
async def delete_account(email: str):
    if email not in users_db:
        raise HTTPException(status_code=404, detail="User not found")
    
    del users_db[email]
    if email in user_stats:
        del user_stats[email]
    if email in payment_history:
        del payment_history[email]
    
    logger.info(f"User account deleted: {email}")
    
    return {"message": "Account deleted successfully"}

# OpenAI AI Endpoints
@app.get("/api/openai/status")
async def get_openai_status():
    """Check OpenAI connection status"""
    if not OPENAI_API_KEY:
        return {
            "status": "disabled",
            "message": "OpenAI API key not configured",
            "configured": False
        }
    
    if not openai_client:
        return {
            "status": "error",
            "message": "OpenAI client not initialized",
            "configured": False
        }
    
    try:
        # Test the connection with a simple request using requests
        headers = {
            "Authorization": f"Bearer {openai_client['api_key']}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": openai_client['model'],
            "messages": [{"role": "user", "content": "Hello"}],
            "max_tokens": 5
        }
        
        response = openai_client['requests'].post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            tokens_used = result.get('usage', {}).get('total_tokens', 0)
            return {
                "status": "connected",
                "message": "OpenAI client is working",
                "configured": True,
                "model": OPENAI_MODEL,
                "tokens_used": tokens_used
            }
        else:
            return {
                "status": "error",
                "message": f"OpenAI API error: {response.status_code}",
                "configured": False
            }
    except Exception as e:
        return {
            "status": "error",
            "message": f"OpenAI connection failed: {str(e)}",
            "configured": False
        }

@app.post("/api/openai/improve-bullet")
async def improve_bullet(payload: ImproveBulletPayload):
    """Improve a bullet point using AI"""
    if not openai_client:
        raise HTTPException(status_code=503, detail="OpenAI service not available")
    
    try:
        tone_instructions = {
            "professional": "Use professional, corporate language with strong action verbs",
            "technical": "Use technical terminology and methodologies, focus on tools and processes",
            "formal": "Use formal, executive-level language with strategic focus",
            "casual": "Use conversational but workplace-appropriate language"
        }
        
        tone_instruction = tone_instructions.get(payload.tone, tone_instructions["professional"])
        
        context = f"Context: {payload.context}" if payload.context else ""
        
        # Generate diverse improvement prompts to avoid repetitive language
        improvement_templates = [
            f"""Transform this resume bullet point into a powerful, achievement-focused statement that will impress recruiters and pass ATS systems.

Current bullet: "{payload.bullet}"

{context}

Enhancement Guidelines:
- Quantify achievements with specific metrics, percentages, or numbers
- Use dynamic action verbs that demonstrate leadership and impact
- Highlight measurable business outcomes and results
- Incorporate relevant industry keywords naturally
- Maintain professional tone while being compelling
- {tone_instruction}
- Ensure ATS compatibility with standard formatting

Return ONLY the enhanced bullet point, no explanations.""",

            f"""Elevate this resume bullet point to showcase your professional impact and expertise.

Current bullet: "{payload.bullet}"

{context}

Optimization Requirements:
- Add concrete data points and quantifiable results
- Use industry-specific terminology and technical keywords
- Demonstrate problem-solving and value creation
- Show progression, growth, or improvement
- {tone_instruction}
- Format for maximum ATS visibility
- Keep concise but impactful

Return ONLY the optimized bullet point, no explanations.""",

            f"""Refine this resume bullet point to maximize its impact on hiring managers and applicant tracking systems.

Current bullet: "{payload.bullet}"

{context}

Improvement Focus:
- Include specific metrics, KPIs, or performance indicators
- Utilize powerful action verbs that convey leadership and expertise
- Emphasize business value and strategic contributions
- Integrate relevant technical skills and methodologies
- {tone_instruction}
- Ensure keyword optimization for ATS scanning
- Maintain clarity and professionalism

Return ONLY the refined bullet point, no explanations."""
        ]
        
        # Randomly select a template to ensure variety
        import random
        prompt = random.choice(improvement_templates)

        headers = {
            "Authorization": f"Bearer {openai_client['api_key']}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": openai_client['model'],
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": OPENAI_MAX_TOKENS,
            "temperature": 0.8,  # Increased temperature for more creative and diverse outputs
            "top_p": 0.9,  # Add top_p for better variety
            "frequency_penalty": 0.3,  # Reduce repetition
            "presence_penalty": 0.2  # Encourage new topics/words
        }
        
        response = openai_client['requests'].post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code}")
        
        result = response.json()
        improved_bullet = result['choices'][0]['message']['content'].strip()
        
        return {
            "success": True,
            "original": payload.bullet,
            "improved": improved_bullet,
            "tokens_used": result.get('usage', {}).get('total_tokens', 0),
            "tone": payload.tone
        }
        
    except Exception as e:
        logger.error(f"OpenAI improve bullet error: {str(e)}")
        error_message = "Failed to improve bullet: " + str(e)
        raise HTTPException(status_code=500, detail=error_message)

@app.post("/api/ai/generate_bullet_points")
async def generate_bullet_points(payload: GenerateBulletPointsPayload):
    """Generate bullet points from scratch using AI"""
    if not openai_client:
        raise HTTPException(status_code=503, detail="OpenAI service not available")
    
    try:
        tone_instructions = {
            "professional": "Use professional, corporate language with strong action verbs",
            "technical": "Use technical terminology and methodologies, focus on tools and processes",
            "formal": "Use formal, executive-level language with strategic focus",
            "casual": "Use conversational but workplace-appropriate language"
        }
        
        tone_instruction = tone_instructions.get(payload.tone, tone_instructions["professional"])
        
        prompt = f"""Generate {payload.count} professional resume bullet points for this role:

Role: {payload.role}
Company: {payload.company}
Skills: {payload.skills}

Requirements:
- Include specific metrics/numbers where possible (use realistic examples)
- Use strong action verbs
- Focus on achievements and impact
- Each bullet should be 1-2 lines
- {tone_instruction}
- ATS-friendly format
- Make them diverse and cover different aspects of the role

Return ONLY the bullet points, one per line, no numbering or explanations."""

        headers = {
            "Authorization": f"Bearer {openai_client['api_key']}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": openai_client['model'],
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": OPENAI_MAX_TOKENS,
            "temperature": 0.7
        }
        
        response = openai_client['requests'].post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code}")
        
        result = response.json()
        bullet_points_text = result['choices'][0]['message']['content'].strip()
        bullet_points = [line.strip() for line in bullet_points_text.split('\n') if line.strip()]
        
        return {
            "success": True,
            "bullet_points": bullet_points,
            "count": len(bullet_points),
            "tokens_used": result.get('usage', {}).get('total_tokens', 0),
            "role": payload.role,
            "company": payload.company,
            "tone": payload.tone
        }
        
    except Exception as e:
        logger.error(f"OpenAI generate bullet points error: {str(e)}")
        error_message = "Failed to generate bullet points: " + str(e)
        raise HTTPException(status_code=500, detail=error_message)

@app.post("/api/ai/generate_summary")
async def generate_summary(payload: GenerateSummaryPayload):
    """Generate professional resume summary using AI"""
    if not openai_client:
        raise HTTPException(status_code=503, detail="OpenAI service not available")
    
    try:
        achievements = f"Key achievements: {payload.achievements}" if payload.achievements else ""
        
        context = f"""
Role: {payload.role}
Experience: {payload.years_experience} years
Skills: {payload.skills}
{achievements}

Requirements:
- Concise and impactful (50-80 words)
- Highlight key strengths and value proposition
- Include relevant technical skills and experience
- ATS-optimized with industry keywords
- Professional tone
- Focus on achievements and impact

Return ONLY the summary text, no explanations or labels."""

        headers = {
            "Authorization": f"Bearer {openai_client['api_key']}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": openai_client['model'],
            "messages": [{"role": "user", "content": context}],
            "max_tokens": OPENAI_MAX_TOKENS,
            "temperature": 0.7
        }
        
        response = openai_client['requests'].post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code}")
        
        result = response.json()
        summary = result['choices'][0]['message']['content'].strip()
        
        return {
            "success": True,
            "summary": summary,
            "tokens_used": result.get('usage', {}).get('total_tokens', 0)
        }
    except Exception as e:
        logger.error(f"OpenAI generate summary error: {str(e)}")
        error_message = "Failed to generate summary: " + str(e)
        raise HTTPException(status_code=500, detail=error_message)

@app.post("/api/ai/generate_bullet_from_keywords")
async def generate_bullet_from_keywords(payload: dict):
    """Generate bullet points from keywords and company context"""
    if not openai_client:
        raise HTTPException(status_code=503, detail="OpenAI service not available")
    
    try:
        keywords = payload.get('keywords', '')
        company_title = payload.get('company_title', '')
        job_title = payload.get('job_title', '')
        
        if not keywords:
            raise HTTPException(status_code=400, detail="Keywords are required")
        
        context = f"""Generate a professional bullet point for a resume based on these keywords and context:

Company: {company_title}
Job Title: {job_title}
Keywords: {keywords}

Requirements:
- Create ONE professional bullet point
- Use action-oriented language (managed, implemented, developed, etc.)
- Include specific technologies/tools if mentioned in keywords
- Make it achievement-focused with metrics if possible
- 1-2 lines maximum
- Professional tone
- ATS-optimized

Return ONLY the bullet point text, no explanations or labels."""

        headers = {
            "Authorization": f"Bearer {openai_client['api_key']}",
            "Content-Type": "application/json"
        }
        
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json={
                "model": "gpt-3.5-turbo",
                "messages": [
                    {"role": "system", "content": "You are a professional resume writer. Create compelling bullet points that highlight achievements and technical skills."},
                    {"role": "user", "content": context}
                ],
                "max_tokens": 150,
                "temperature": 0.7
            },
            timeout=60
        )
        
        if response.status_code != 200:
            logger.error(f"OpenAI API error: {response.status_code} - {response.text}")
            raise HTTPException(status_code=500, detail="AI service error")
        
        result = response.json()
        bullet_text = result['choices'][0]['message']['content'].strip()
        
        # Clean up the response
        bullet_text = bullet_text.replace('•', '').replace('*', '').strip()
        
        return {
            "success": True,
            "bullet_text": bullet_text,
            "keywords_used": keywords,
            "company_title": company_title,
            "job_title": job_title,
            "tokens_used": result.get('usage', {}).get('total_tokens', 0)
        }
        
    except Exception as e:
        logger.error(f"OpenAI generate bullet from keywords error: {str(e)}")
        error_message = "AI generation failed: " + str(e)
        raise HTTPException(status_code=500, detail=error_message)

@app.post("/api/ai/generate_bullets_from_keywords")
async def generate_bullets_from_keywords(payload: dict):
    """Generate multiple bullet points from selected keywords for a work experience entry"""
    if not openai_client:
        raise HTTPException(status_code=503, detail="OpenAI service not available")
    
    try:
        keywords_list = payload.get('keywords', [])
        job_description = payload.get('job_description', '')
        company_title = payload.get('company_title', '')
        job_title = payload.get('job_title', '')
        
        if not keywords_list or len(keywords_list) == 0:
            raise HTTPException(status_code=400, detail="Keywords list is required")
        
        keywords_str = ', '.join(keywords_list)
        
        context = f"""Generate professional resume bullet points based on these selected keywords and job description context:

Job Description Context:
{job_description[:500] if job_description else 'Not provided'}

Company: {company_title or 'Not specified'}
Job Title: {job_title or 'Not specified'}
Selected Keywords: {keywords_str}

Requirements:
- Generate 3-5 professional bullet points that incorporate ALL {len(keywords_list)} selected keywords together
- Each bullet should naturally incorporate 2-3 of the keywords in a cohesive way
- Use ALL keywords across the bullets - distribute them naturally across the generated bullets
- Use action-oriented language (Led, Implemented, Developed, Optimized, etc.)
- Include specific technologies/tools mentioned in keywords
- Make them achievement-focused with metrics when possible
- Each bullet should be 1-2 lines maximum
- Professional, ATS-optimized tone
- Ensure keywords are naturally integrated, not forced
- Make bullets diverse and cover different aspects of the role
- Keywords should be used naturally in context, not just listed

Return ONLY a JSON array of bullet point strings, no explanations or labels.
Example format: ["Implemented scalable microservices architecture using Kubernetes and Docker", "Optimized CI/CD pipelines with Jenkins reducing deployment time by 40%", "Developed REST APIs using Python and Flask handling 10K+ requests daily"]"""

        headers = {
            "Authorization": f"Bearer {openai_client['api_key']}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": openai_client['model'],
            "messages": [
                {"role": "system", "content": "You are a professional resume writer. Create compelling bullet points that highlight achievements and technical skills while naturally incorporating keywords."},
                {"role": "user", "content": context}
            ],
            "max_tokens": 800,
            "temperature": 0.7
        }
        
        response = openai_client['requests'].post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code != 200:
            logger.error(f"OpenAI API error: {response.status_code} - {response.text}")
            raise HTTPException(status_code=500, detail="AI service error")
        
        result = response.json()
        bullets_text = result['choices'][0]['message']['content'].strip()
        
        # Try to parse as JSON array
        try:
            import json
            bullets = json.loads(bullets_text)
            if not isinstance(bullets, list):
                # If not a list, try to extract bullet points from text
                bullets = [line.strip().replace('•', '').replace('*', '').replace('-', '').strip() 
                          for line in bullets_text.split('\n') if line.strip()]
        except json.JSONDecodeError:
            # Parse as plain text - extract bullet points
            bullets = [line.strip().replace('•', '').replace('*', '').replace('-', '').strip() 
                      for line in bullets_text.split('\n') if line.strip() and len(line.strip()) > 10]
        
        # Clean up bullets
        cleaned_bullets = []
        for bullet in bullets:
            bullet = bullet.strip()
            if bullet and len(bullet) > 10:
                # Remove common prefixes
                for prefix in ['•', '*', '-', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.']:
                    if bullet.startswith(prefix):
                        bullet = bullet[len(prefix):].strip()
                cleaned_bullets.append(bullet)
        
        return {
            "success": True,
            "bullets": cleaned_bullets[:len(keywords_list) + 2],  # Limit to reasonable number
            "keywords_used": keywords_list,
            "company_title": company_title,
            "job_title": job_title,
            "tokens_used": result.get('usage', {}).get('total_tokens', 0)
        }
        
    except Exception as e:
        logger.error(f"OpenAI generate bullets from keywords error: {str(e)}")
        error_message = "AI generation failed: " + str(e)
        raise HTTPException(status_code=500, detail=error_message)

@app.post("/api/ai/generate_summary_from_experience")
async def generate_summary_from_experience(payload: dict):
    """Generate ATS-optimized professional summary by analyzing work experience"""
    if not openai_client:
        raise HTTPException(status_code=503, detail="OpenAI service not available")
    
    try:
        name = payload.get('name', 'Professional')
        title = payload.get('title', '')
        sections = payload.get('sections', [])
        
        # Extract work experience
        work_experience_text = ""
        skills_text = ""
        
        for section in sections:
            section_title = section.get('title', '').lower()
            bullets = section.get('bullets', [])
            
            if 'experience' in section_title or 'work' in section_title or 'employment' in section_title:
                work_experience_text += f"\n{section.get('title')}:\n"
                for bullet in bullets:
                    bullet_text = bullet.get('text', '') if isinstance(bullet, dict) else str(bullet)
                    if bullet_text.strip():
                        work_experience_text += f"  {bullet_text}\n"
            
            if 'skill' in section_title:
                for bullet in bullets:
                    bullet_text = bullet.get('text', '') if isinstance(bullet, dict) else str(bullet)
                    if bullet_text.strip():
                        skills_text += f"{bullet_text}, "
        
        context = f"""Analyze this professional's work experience and create a compelling ATS-optimized professional summary.

Professional Title: {title if title else 'Not specified'}

Work Experience:
{work_experience_text if work_experience_text else 'Limited information provided'}

Skills:
{skills_text if skills_text else 'To be extracted from experience'}

Requirements for the Professional Summary:
1. Length: 6-7 sentences (approximately 100-120 words)
2. ATS-Optimized: Include relevant keywords from their experience and industry
3. Structure:
   - Sentence 1: Opening statement with years of experience and core expertise
   - Sentences 2-4: Key achievements, skills, and value proposition with specific metrics when available
   - Sentences 5-6: Technical competencies and areas of expertise
   - Sentence 7: Career objective or unique value add
4. Include specific technologies, tools, and methodologies mentioned in experience
5. Use action-oriented language and quantifiable achievements
6. Professional, confident tone
7. Third-person perspective (avoid "I")
8. Focus on impact and results
9. Include industry-specific keywords for ATS systems

Return ONLY the professional summary paragraph, no labels, explanations, or formatting markers."""

        headers = {
            "Authorization": f"Bearer {openai_client['api_key']}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": openai_client['model'],
            "messages": [
                {"role": "system", "content": "You are an expert resume writer specializing in ATS-optimized professional summaries. You analyze work experience to create compelling, keyword-rich summaries that pass ATS systems and impress recruiters."},
                {"role": "user", "content": context}
            ],
            "max_tokens": 500,
            "temperature": 0.7
        }
        
        logger.info(f"Generating summary from experience for: {name}")
        
        response = openai_client['requests'].post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code}")
        
        result = response.json()
        summary = result['choices'][0]['message']['content'].strip()
        
        # Remove any quotes or formatting markers
        summary = summary.strip('"\'')
        
        logger.info(f"Generated summary: {len(summary)} characters")
        
        return {
            "success": True,
            "summary": summary,
            "tokens_used": result.get('usage', {}).get('total_tokens', 0),
            "word_count": len(summary.split())
        }
    except Exception as e:
        logger.error(f"OpenAI generate summary from experience error: {str(e)}")
        error_message = "Failed to generate summary: " + str(e)
        raise HTTPException(status_code=500, detail=error_message)

@app.post("/api/ai/match_job_description")
async def match_job_description(payload: JobDescriptionMatchPayload, user_email: str = None, db: Session = Depends(get_db)):
    """Match job description with resume and calculate similarity score"""
    try:
        # Convert resume data to text for analysis
        resume_text = f"{payload.resume_data.name} — {payload.resume_data.title}\n\n"
        if payload.resume_data.summary:
            resume_text += payload.resume_data.summary + "\n\n"
        
        for section in payload.resume_data.sections:
            resume_text += f"{section.title}\n"
            for bullet in section.bullets:
                resume_text += f"• {bullet.text}\n"
            resume_text += "\n"
        
        # Calculate similarity using keyword extractor
        similarity_result = keyword_extractor.calculate_similarity(
            payload.job_description, 
            resume_text
        )
        
        # Get keyword suggestions for missing keywords
        suggestions = keyword_extractor.get_keyword_suggestions(
            similarity_result.get('missing_keywords', [])
        )
        
        # Generate AI-powered improvement suggestions
        improvement_suggestions = []
        if openai_client and similarity_result.get('similarity_score', 0) < 70:
            try:
                improvement_prompt = f"""As an expert resume strategist, analyze this job-resume match and provide specific, actionable improvements.

JOB DESCRIPTION:
{payload.job_description[:1000]}

CURRENT RESUME:
{resume_text[:1000]}

ANALYSIS:
- Missing Keywords: {', '.join(similarity_result.get('missing_keywords', [])[:10])}
- Match Score: {similarity_result.get('similarity_score', 0)}%
- Technical Skills Gap: {', '.join(similarity_result.get('technical_missing', [])[:5])}

REQUIREMENTS:
Provide 4-6 highly specific, actionable suggestions that will:
1. Bridge keyword gaps with natural integration
2. Add missing technical skills to relevant sections
3. Quantify achievements with metrics
4. Enhance ATS compatibility
5. Strengthen job-relevant experience descriptions
6. Optimize professional summary for this role

Each suggestion should be:
- Specific and implementable
- Include exact wording examples
- Target specific resume sections
- Address the identified gaps

Return as JSON array: [{{"category": "Technical Skills", "suggestion": "Add Python and SQL to skills section and include in work experience bullets"}}, ...]"""

                headers = {
                    "Authorization": f"Bearer {openai_client['api_key']}",
                    "Content-Type": "application/json"
                }
                
                data = {
                    "model": openai_client['model'],
                    "messages": [{"role": "user", "content": improvement_prompt}],
                    "max_tokens": 800,  # Increased for more detailed suggestions
                    "temperature": 0.6,  # Balanced creativity and consistency
                    "top_p": 0.9,
                    "frequency_penalty": 0.2,
                    "presence_penalty": 0.1
                }
                
                response = openai_client['requests'].post(
                    "https://api.openai.com/v1/chat/completions",
                    headers=headers,
                    json=data,
                    timeout=60
                )
                
                if response.status_code == 200:
                    result = response.json()
                    suggestions_text = result['choices'][0]['message']['content'].strip()
                    
                    # Try to parse as JSON
                    try:
                        import json
                        improvement_suggestions = json.loads(suggestions_text)
                    except json.JSONDecodeError:
                        # If not JSON, create structured suggestions
                        lines = [line.strip() for line in suggestions_text.split('\n') if line.strip()]
                        improvement_suggestions = [
                            {"category": "General", "suggestion": line} 
                            for line in lines[:5]
                        ]
                        
            except Exception as e:
                logger.error(f"Error generating AI suggestions: {e}")
        
        return {
            "success": True,
            "match_analysis": {
                "similarity_score": similarity_result.get('similarity_score', 0),
                "technical_score": similarity_result.get('technical_score', 0),
                "matching_keywords": similarity_result.get('matching_keywords', []),
                "missing_keywords": similarity_result.get('missing_keywords', []),
                "technical_matches": similarity_result.get('technical_matches', []),
                "technical_missing": similarity_result.get('technical_missing', []),
                "total_job_keywords": similarity_result.get('total_job_keywords', 0),
                "match_count": similarity_result.get('match_count', 0),
                "missing_count": similarity_result.get('missing_count', 0)
            },
            "keyword_suggestions": suggestions,
            "improvement_suggestions": improvement_suggestions,
            "analysis_summary": {
                "overall_match": "Excellent" if similarity_result.get('similarity_score', 0) >= 80 else
                               "Good" if similarity_result.get('similarity_score', 0) >= 60 else
                               "Fair" if similarity_result.get('similarity_score', 0) >= 40 else
                               "Needs Improvement",
                "technical_match": "Strong" if similarity_result.get('technical_score', 0) >= 70 else
                                 "Moderate" if similarity_result.get('technical_score', 0) >= 40 else
                                 "Weak"
            }
        }
        
        # Track job match analytics
        if user_email and db:
            try:
                user = db.query(User).filter(User.email == user_email).first()
                if user:
                    # Find or create resume record
                    resume = db.query(Resume).filter(
                        Resume.user_id == user.id,
                        Resume.name == payload.resume_data.name
                    ).first()
                    
                    if not resume:
                        resume = Resume(
                            user_id=user.id,
                            name=payload.resume_data.name,
                            title=payload.resume_data.title,
                            email=payload.resume_data.email,
                            phone=payload.resume_data.phone,
                            location=payload.resume_data.location,
                            summary=payload.resume_data.summary,
                            template="tech"
                        )
                        db.add(resume)
                        db.commit()
                        db.refresh(resume)
                    
                    # Get latest version for this resume
                    version_service = VersionControlService(db)
                    latest_version = version_service.get_latest_version(resume.id, user.id)
                    
                    # Create job match record
                    job_match = JobMatch(
                        user_id=user.id,
                        resume_id=resume.id,
                        resume_version_id=latest_version.id if latest_version else None,
                        job_description=payload.job_description,
                        match_score=similarity_result.get('similarity_score', 0),
                        keyword_matches=similarity_result.get('matching_keywords', []),
                        missing_keywords=similarity_result.get('missing_keywords', []),
                        improvement_suggestions=improvement_suggestions
                    )
                    
                    db.add(job_match)
                    db.commit()
                    
                    logger.info(f"Job match analytics tracked for user {user_email}: score {similarity_result.get('similarity_score', 0)}")
            except Exception as e:
                logger.error(f"Failed to track job match analytics: {e}")
        
        return {
            "success": True,
            "match_analysis": {
                "similarity_score": similarity_result.get('similarity_score', 0),
                "technical_score": similarity_result.get('technical_score', 0),
                "matching_keywords": similarity_result.get('matching_keywords', []),
                "missing_keywords": similarity_result.get('missing_keywords', []),
                "technical_matches": similarity_result.get('technical_matches', []),
                "technical_missing": similarity_result.get('technical_missing', []),
                "total_job_keywords": similarity_result.get('total_job_keywords', 0),
                "match_count": similarity_result.get('match_count', 0),
                "missing_count": similarity_result.get('missing_count', 0)
            },
            "keyword_suggestions": suggestions,
            "improvement_suggestions": improvement_suggestions,
            "analysis_summary": {
                "overall_match": "Excellent" if similarity_result.get('similarity_score', 0) >= 80 else
                               "Good" if similarity_result.get('similarity_score', 0) >= 60 else
                               "Fair" if similarity_result.get('similarity_score', 0) >= 40 else
                               "Needs Improvement",
                "technical_match": "Strong" if similarity_result.get('technical_score', 0) >= 70 else
                                 "Moderate" if similarity_result.get('technical_score', 0) >= 40 else
                                 "Weak"
            }
        }
        
    except Exception as e:
        logger.error(f"Job description matching error: {str(e)}")
        error_message = "Failed to analyze job match: " + str(e)
        raise HTTPException(status_code=500, detail=error_message)

@app.post("/api/ai/cover_letter")
async def generate_cover_letter(payload: CoverLetterPayload):
    """Generate a tailored cover letter using AI"""
    if not openai_client:
        raise HTTPException(status_code=503, detail="OpenAI service not available")
    
    try:
        # Convert resume data to text for context
        resume_text = f"{payload.resume_data.name} — {payload.resume_data.title}\n\n"
        if payload.resume_data.summary:
            resume_text += payload.resume_data.summary + "\n\n"
        
        for section in payload.resume_data.sections:
            resume_text += f"{section.title}\n"
            for bullet in section.bullets:
                resume_text += f"• {bullet.text}\n"
            resume_text += "\n"
        
        # Define tone instructions
        tone_instructions = {
            "professional": "Use formal, corporate language with strong action verbs and industry terminology",
            "friendly": "Use warm, approachable language while maintaining professionalism",
            "concise": "Use direct, clear language with short sentences and bullet points where appropriate"
        }
        
        tone_instruction = tone_instructions.get(payload.tone, tone_instructions["professional"])
        
        # Build custom requirements context
        custom_context = f"\nAdditional Requirements: {payload.custom_requirements}" if payload.custom_requirements else ""
        
        prompt = f"""Generate a professional cover letter for this job application.

Job Details:
- Company: {payload.company_name}
- Position: {payload.position_title}
- Job Description: {payload.job_description[:2000]}

Candidate Information:
{resume_text[:2000]}

Requirements:
- Write a compelling cover letter that connects the candidate's experience to the job requirements
- Use {payload.tone} tone: {tone_instruction}
- Include specific examples from the candidate's background that match job requirements
- Address key requirements from the job description with concrete examples
- Keep it professional and engaging
- Length: 3-4 paragraphs (approximately 250-350 words)
- Start with a strong opening that shows enthusiasm and mentions the specific position
- Use action verbs and quantifiable achievements where possible
- End with a clear call to action
- Avoid generic phrases - be specific to this role and company
{custom_context}

Structure the response as JSON with these fields:
- opening: Opening paragraph (1-2 sentences) - mention the specific position and company
- body: Main body paragraphs (2-3 paragraphs) - highlight relevant experience and achievements
- closing: Closing paragraph (1-2 sentences) - express enthusiasm and next steps
- full_letter: Complete formatted letter with proper spacing

Return ONLY valid JSON, no markdown formatting."""

        headers = {
            "Authorization": f"Bearer {openai_client['api_key']}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": openai_client['model'],
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 1000,
            "temperature": 0.7
        }
        
        response = openai_client['requests'].post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code}")
        
        result = response.json()
        cover_letter_content = result['choices'][0]['message']['content'].strip()
        
        # Try to parse as JSON
        try:
            import json
            parsed_content = json.loads(cover_letter_content)
        except json.JSONDecodeError:
            # If not JSON, create structured response
            parsed_content = {
                "opening": "I am writing to express my strong interest in the " + payload.position_title + " position at " + payload.company_name + ".",
                "body": cover_letter_content,
                "closing": "I would welcome the opportunity to discuss how my experience can contribute to your team. Thank you for your consideration.",
                "full_letter": cover_letter_content
            }
        
        return {
            "success": True,
            "cover_letter": parsed_content,
            "tokens_used": result.get('usage', {}).get('total_tokens', 0),
            "company": payload.company_name,
            "position": payload.position_title,
            "tone": payload.tone
        }
        
    except Exception as e:
        logger.error(f"Cover letter generation error: {str(e)}")
        error_message = "Failed to generate cover letter: " + str(e)
        raise HTTPException(status_code=500, detail=error_message)

@app.post("/api/ai/grammar_check")
async def check_grammar_style(payload: GrammarCheckPayload):
    """Check grammar and style of text"""
    try:
        text = payload.text.strip()
        if not text:
            return {
                "success": False,
                "error": "No text provided for checking"
            }
        
        logger.info(f"Grammar/style check requested for {len(text)} characters")
        
        if payload.check_type in ["grammar", "all"]:
            grammar_issues = grammar_checker.check_grammar(text)
        else:
            grammar_issues = []
        
        if payload.check_type in ["style", "all"]:
            passive_issues = grammar_checker.check_passive_voice(text)
            weak_verb_issues = grammar_checker.check_weak_verbs(text)
            readability_score, readability_issues = grammar_checker.check_readability(text)
            strength_score, strength_issues = grammar_checker.check_action_verbs(text)
            style_score = grammar_checker.calculate_style_score(text)
            improvement_suggestions = grammar_checker.get_improvement_suggestions(text)
        else:
            passive_issues = []
            weak_verb_issues = []
            readability_issues = []
            strength_issues = []
            style_score = None
            improvement_suggestions = []
            readability_score = 100
            strength_score = 100
        
        # Format grammar issues
        formatted_grammar_issues = []
        for issue in grammar_issues:
            formatted_grammar_issues.append({
                "message": issue.message,
                "replacements": issue.replacements,
                "offset": issue.offset,
                "length": issue.length,
                "rule_id": issue.rule_id,
                "category": issue.category,
                "severity": issue.severity
            })
        
        # Format style issues
        all_style_issues = passive_issues + weak_verb_issues + readability_issues + strength_issues
        formatted_style_issues = []
        for issue in all_style_issues:
            formatted_style_issues.append({
                "type": issue.type,
                "message": issue.message,
                "suggestion": issue.suggestion,
                "severity": issue.severity,
                "score_impact": issue.score_impact
            })
        
        response = {
            "success": True,
            "text_length": len(text),
            "grammar_issues": formatted_grammar_issues,
            "style_issues": formatted_style_issues,
            "improvement_suggestions": improvement_suggestions
        }
        
        if style_score:
            response.update({
                "style_score": {
                    "overall_score": style_score.overall_score,
                    "grammar_score": style_score.grammar_score,
                    "readability_score": style_score.readability_score,
                    "strength_score": style_score.strength_score,
                    "issues_count": style_score.issues_count,
                    "suggestions": style_score.suggestions
                }
            })
        
        return response
        
    except Exception as e:
        logger.error(f"Grammar/style check error: {str(e)}")
        error_message = "Failed to check grammar and style: " + str(e)
        raise HTTPException(status_code=500, detail=error_message)

@app.post("/api/ai/generate_resume_content")
async def generate_resume_content(payload: dict):
    """Generate resume content based on user requirements"""
    if not openai_client:
        raise HTTPException(status_code=503, detail="OpenAI service not available")
    
    try:
        content_type = payload.get('contentType', 'job')
        requirements = payload.get('requirements', '')
        position = payload.get('position', 'end')
        target_section = payload.get('targetSection', '')
        existing_data = payload.get('existingData', {})
        context = payload.get('context', {})
        
        # Build context from existing resume data
        existing_context = f"""
        Current Resume:
        Name: {context.get('name', '')}
        Title: {context.get('title', '')}
        Existing Sections: {', '.join(context.get('currentSections', []))}
        """
        
        # Generate content based on type
        if content_type == 'job':
            prompt = f"""
            Based on the following requirements, generate a complete work experience entry:
            
            Requirements: {requirements}
            Position: {position}
            {existing_context}
            
            Generate a REALISTIC work experience entry with:
            1. Company name (use a real tech company like Google, Microsoft, Amazon, etc.)
            2. Job title/role (specific to the requirements)
            3. Duration (realistic timeframe like "2022-2024" or "Jan 2023 - Present")
            4. 4-6 professional bullet points with:
               - Action verbs and quantifiable results
               - Technical skills mentioned in requirements
               - ATS-optimized language
               - Progressive responsibility
            
            IMPORTANT: 
            - Use REAL company names, not placeholders
            - Use REAL job titles, not generic ones
            - Use REAL timeframes, not placeholders
            - Make bullet points specific and detailed
            
            Return ONLY valid JSON with fields: company, role, duration, bullets (array of strings)
            
            Example format:
            {{
              "company": "Google",
              "role": "DevOps Engineer", 
              "duration": "2022-2024",
              "bullets": ["Deployed applications using Kubernetes", "Managed CI/CD pipelines"]
            }}
            """
        elif content_type == 'project':
            prompt = f"""
            Based on the following requirements, generate a project entry:
            
            Requirements: {requirements}
            {existing_context}
            
            Generate a project entry with:
            1. Project name
            2. Brief description
            3. 3-5 bullet points with:
               - Technical implementation details
               - Technologies used
               - Results/impact
               - Challenges overcome
            
            Return as JSON with fields: name, description, bullets (array of strings)
            """
        elif content_type == 'skill':
            prompt = f"""
            Based on the following requirements, generate a skills section:
            
            Requirements: {requirements}
            {existing_context}
            
            Generate a skills section with:
            1. Categorized skills (Technical, Tools, Languages, etc.)
            2. Relevant to the person's background
            3. Industry-standard terminology
            4. ATS-friendly format
            
            Return as JSON with fields: categories (object with category names as keys and skill arrays as values)
            """
        elif content_type == 'education':
            prompt = f"""
            Based on the following requirements, generate an education entry:
            
            Requirements: {requirements}
            {existing_context}
            
            Generate an education entry with:
            1. Institution name
            2. Degree/qualification
            3. Relevant coursework (if applicable)
            4. Graduation year
            5. Any honors or achievements
            
            Return as JSON with fields: institution, degree, year, coursework (array), honors (array)
            """
        elif content_type == 'bullet-improvement':
            current_bullet = context.get('bulletText', '')
            section_title = context.get('sectionTitle', 'Work Experience')
            company_name = context.get('companyName', '')
            job_title = context.get('jobTitle', '')
            
            prompt = f"""
            Improve the following bullet point for a resume:
            
            Current bullet point: "{current_bullet}"
            Section: {section_title}
            Company: {company_name}
            Role: {job_title}
            
            Requirements: {requirements}
            
            Please improve this bullet point by:
            - Adding specific metrics and quantifiable results
            - Using strong action verbs
            - Making it more achievement-focused
            - Keeping it concise but impactful
            - Maintaining professional tone
            
            Return as JSON with field: improvedBullet (string)
            """
        else:
            raise HTTPException(status_code=400, detail="Invalid content type")
        
        headers = {
            "Authorization": f"Bearer {openai_client['api_key']}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": openai_client['model'],
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 1000,
            "temperature": 0.7
        }
        
        response = openai_client['requests'].post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code}")
        
        result = response.json()
        content = result['choices'][0]['message']['content'].strip()
        
        logger.info("OpenAI response content: " + str(content))
        
        # Try to parse as JSON
        try:
            import json
            parsed_content = json.loads(content)
            logger.info("Parsed JSON content: " + str(parsed_content))
            return parsed_content
        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse JSON: {e}")
            logger.warning("Raw content: " + str(content))
            
            # Try to extract JSON from the content if it's wrapped in markdown
            import re
            json_match = re.search(r'```json\s*(\{.*?\})\s*```', content, re.DOTALL)
            if json_match:
                try:
                    parsed_content = json.loads(json_match.group(1))
                    logger.info("Extracted JSON from markdown: " + str(parsed_content))
                    return parsed_content
                except json.JSONDecodeError:
                    pass
            
            # If all else fails, create a structured response based on content type
            if content_type == 'job':
                return {
                    "company": "Generated Company",
                    "role": "Generated Role", 
                    "duration": "2023-2024",
                    "bullets": [content]
                }
            elif content_type == 'bullet-improvement':
                return {
                    "improvedBullet": content
                }
            else:
                return {"content": content}
            
    except Exception as e:
        logger.error(f"Content generation failed: {e}")
        error_message = "Failed to generate content: " + str(e)
        raise HTTPException(status_code=500, detail=error_message)

TEMPLATES = {
    "tech": {
        "name": "Tech Professional",
        "industry": "Technology",
        "styles": {
            "header_align": "left",
            "header_border": "3px solid #2563eb",
            "font": "Arial, sans-serif",
            "section_uppercase": False,
            "layout": "single"
        }
    },
    "healthcare": {
        "name": "Healthcare Pro",
        "industry": "Healthcare",
        "styles": {
            "header_align": "center",
            "header_border": "2px solid #059669",
            "font": "Georgia, serif",
            "section_uppercase": True,
            "layout": "single"
        }
    },
    "finance": {
        "name": "Finance Executive",
        "industry": "Finance",
        "styles": {
            "header_align": "center",
            "header_border": "3px solid #1e40af",
            "font": "Times New Roman, serif",
            "section_uppercase": True,
            "layout": "single"
        }
    },
    "creative": {
        "name": "Creative Portfolio",
        "industry": "Marketing & Design",
        "styles": {
            "header_align": "left",
            "header_border": "none",
            "font": "Verdana, sans-serif",
            "section_uppercase": False,
            "layout": "single"
        }
    },
    "academic": {
        "name": "Academic Scholar",
        "industry": "Education",
        "styles": {
            "header_align": "center",
            "header_border": "1px solid #000",
            "font": "Times New Roman, serif",
            "section_uppercase": False,
            "layout": "single"
        }
    },
    "legal": {
        "name": "Legal Professional",
        "industry": "Legal",
        "styles": {
            "header_align": "center",
            "header_border": "2px solid #1e293b",
            "font": "Georgia, serif",
            "section_uppercase": True,
            "layout": "single"
        }
    },
    "engineering": {
        "name": "Engineering Pro",
        "industry": "Engineering",
        "styles": {
            "header_align": "left",
            "header_border": "2px solid #ea580c",
            "font": "Arial, sans-serif",
            "section_uppercase": False,
            "layout": "single"
        }
    },
    "sales": {
        "name": "Sales Leader",
        "industry": "Sales",
        "styles": {
            "header_align": "left",
            "header_border": "3px solid #dc2626",
            "font": "Arial, sans-serif",
            "section_uppercase": False,
            "layout": "single"
        }
    },
    "consulting": {
        "name": "Consultant Elite",
        "industry": "Consulting",
        "styles": {
            "header_align": "center",
            "header_border": "2px solid #4338ca",
            "font": "Georgia, serif",
            "section_uppercase": True,
            "layout": "single"
        }
    },
    "hr": {
        "name": "HR Professional",
        "industry": "Human Resources",
        "styles": {
            "header_align": "left",
            "header_border": "2px solid #7c3aed",
            "font": "Arial, sans-serif",
            "section_uppercase": False,
            "layout": "single"
        }
    },
    "operations": {
        "name": "Operations Manager",
        "industry": "Operations",
        "styles": {
            "header_align": "left",
            "header_border": "2px solid #0891b2",
            "font": "Arial, sans-serif",
            "section_uppercase": False,
            "layout": "single"
        }
    },
    "customer": {
        "name": "Customer Success",
        "industry": "Customer Service",
        "styles": {
            "header_align": "center",
            "header_border": "2px solid #06b6d4",
            "font": "Arial, sans-serif",
            "section_uppercase": False,
            "layout": "single"
        }
    },
    "data": {
        "name": "Data Scientist",
        "industry": "Data & Analytics",
        "styles": {
            "header_align": "left",
            "header_border": "2px solid #8b5cf6",
            "font": "Courier New, monospace",
            "section_uppercase": False,
            "layout": "single"
        }
    },
    "product": {
        "name": "Product Manager",
        "industry": "Product",
        "styles": {
            "header_align": "left",
            "header_border": "3px solid #ec4899",
            "font": "Arial, sans-serif",
            "section_uppercase": False,
            "layout": "single"
        }
    },
    "executive": {
        "name": "Executive Leader",
        "industry": "Executive",
        "styles": {
            "header_align": "center",
            "header_border": "3px solid #1e293b",
            "font": "Georgia, serif",
            "section_uppercase": True,
            "layout": "single"
        }
    }
}

@app.get("/api/resume/templates")
async def get_templates():
    return {
        "templates": [
            {"id": tid, "name": t["name"], "industry": t.get("industry", "General")} 
            for tid, t in TEMPLATES.items()
        ]
    }

@app.post("/api/resume/parse-file")
async def parse_file(file: UploadFile = File(...)):
    """Industry-standard resume parsing with multiple extraction methods"""
    try:
        logger.info(f"Parsing file: {file.filename}, size: {file.size}")
        
        # File validation
        if not file.filename:
            return {"success": False, "error": "No filename provided"}
        
        if file.size > 10 * 1024 * 1024:  # 10MB limit
            return {"success": False, "error": "File too large. Maximum size is 10MB"}
        
        file_extension = file.filename.split('.')[-1].lower()
        if file_extension not in ['pdf', 'docx', 'doc', 'txt']:
            return {"success": False, "error": "Unsupported file type. Please upload PDF, DOCX, DOC, or TXT"}
        
        # Read file content
        file_content = await file.read()
        
        # Extract text using multiple methods for reliability
        text = ""
        extraction_methods = []
        
        if file_extension == 'pdf':
            text, methods = extract_pdf_text(file_content)
            extraction_methods.extend(methods)
        
        elif file_extension == 'docx':
            text, methods = extract_docx_text(file_content)
            extraction_methods.extend(methods)
        
        elif file_extension == 'doc':
            text, methods = extract_doc_text(file_content)
            extraction_methods.extend(methods)
        
        elif file_extension == 'txt':
            try:
                text = file_content.decode('utf-8')
                extraction_methods.append("UTF-8 decode")
            except:
                try:
                    text = file_content.decode('latin-1')
                    extraction_methods.append("Latin-1 decode")
                except:
                    return {"success": False, "error": "Could not decode text file"}
        
        # Clean and validate extracted text
        text = clean_extracted_text(text)
        
        if not text.strip():
            return {
                "success": False,
                "error": "No readable text found. The file might be:\n• Scanned image (try OCR)\n• Password protected\n• Corrupted\n• Empty"
            }
        
        if len(text.strip()) < 50:
            return {
                "success": False,
                "error": "Text too short. Please ensure the file contains a complete resume."
            }
        
        # Use regex parsing for now (AI is hanging)
        logger.info(f"Parsing {len(text)} characters with regex parser using methods: {', '.join(extraction_methods)}")
        parsed_data = parse_resume_with_regex(text)
        
        return {
            "success": True,
            "data": parsed_data,
            "raw_text": text[:500],  # First 500 chars for debugging
            "extraction_methods": extraction_methods,
            "message": f"Resume parsed successfully using {len(extraction_methods)} method(s) - {len(parsed_data.get('sections', []))} sections extracted"
        }
        
    except Exception as e:
        logger.error(f"File parsing error: {str(e)}")
        return {
            "success": False,
            "error": f"Upload failed: {str(e)}"
        }

def extract_pdf_text(file_content: bytes) -> tuple[str, list[str]]:
    """Extract text from PDF using multiple methods"""
    text = ""
    methods = []
    
    # Method 1: PyMuPDF (fitz) - Most reliable
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_content, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        methods.append("PyMuPDF")
        logger.info(f"PyMuPDF extracted {len(text)} characters")
    except Exception as e:
        logger.warning(f"PyMuPDF failed: {e}")
    
    # Method 2: pdfplumber - Good for complex layouts
    if len(text.strip()) < 100:
        try:
            import pdfplumber
            import io
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            methods.append("pdfplumber")
            logger.info(f"pdfplumber extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
    
    # Method 3: PyPDF2 - Fallback
    if len(text.strip()) < 100:
        try:
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            methods.append("PyPDF2")
            logger.info(f"PyPDF2 extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")
    
    return text, methods

def extract_docx_text(file_content: bytes) -> tuple[str, list[str]]:
    """Extract text from DOCX using multiple methods"""
    text = ""
    methods = []
    
    # Method 1: python-docx - Standard method
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(file_content))
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        methods.append("python-docx")
        logger.info(f"python-docx extracted {len(text)} characters")
    except Exception as e:
        logger.warning(f"python-docx failed: {e}")
    
    # Method 2: docx2txt - Alternative method
    if len(text.strip()) < 100:
        try:
            import docx2txt
            import io
            text = docx2txt.process(io.BytesIO(file_content))
            methods.append("docx2txt")
            logger.info(f"docx2txt extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"docx2txt failed: {e}")
    
    return text, methods

def extract_doc_text(file_content: bytes) -> tuple[str, list[str]]:
    """Extract text from DOC files"""
    text = ""
    methods = []
    
    # Method 1: antiword (if available)
    try:
        import subprocess
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            result = subprocess.run(['antiword', temp_file_path], 
                                  capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                text = result.stdout
                methods.append("antiword")
                logger.info(f"antiword extracted {len(text)} characters")
        finally:
            os.unlink(temp_file_path)
    except Exception as e:
        logger.warning(f"antiword failed: {e}")
    
    # Method 2: Try reading as text (fallback)
    if len(text.strip()) < 100:
        try:
            text = file_content.decode('utf-8', errors='ignore')
            methods.append("UTF-8 decode")
            logger.info(f"UTF-8 decode extracted {len(text)} characters")
        except Exception as e:
            logger.warning(f"UTF-8 decode failed: {e}")
    
    return text, methods

def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Max 2 newlines
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces to single
    
    # Remove common PDF artifacts
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII
    text = re.sub(r'\f', '\n', text)  # Form feeds to newlines
    
    # Clean up bullet points
    text = re.sub(r'[•·▪▫‣⁃]', '•', text)  # Normalize bullets
    
    return text.strip()

def parse_resume_with_regex(text: str) -> dict:
    """Simple, reliable resume parser that actually works"""
    logger.info("Using simple regex-based resume parsing")
    
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # Extract basic info
    name = ""
    title = ""
    email = ""
    phone = ""
    location = ""
    
    # Find name (first line that looks like a name)
    for line in lines[:3]:
        if len(line) > 2 and not any(char.isdigit() for char in line) and '@' not in line and '•' not in line:
            name = line
            break
    
    # Find title (line after name)
    for i, line in enumerate(lines[1:4]):
        if len(line) > 2 and not any(char.isdigit() for char in line) and '@' not in line and '•' not in line and line != name:
            title = line
            break
    
    # Find email
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    if email_match:
        email = email_match.group()
    
    # Find phone
    phone_match = re.search(r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})', text)
    if phone_match:
        phone = phone_match.group().strip()
    
    # Find location
    location_match = re.search(r'([A-Za-z\s]+),\s*([A-Z]{2})', text)
    if location_match:
        location = location_match.group().strip()
    
    # Extract sections
    sections = []
    current_section = None
    current_bullets = []
    
    section_keywords = ['work experience', 'professional experience', 'employment', 'experience', 'projects', 'education', 'skills', 'certifications']
    
    for line in lines:
        line_lower = line.lower()
        
        # Check if this is a section header
        is_section = any(keyword in line_lower for keyword in section_keywords)
        
        if is_section:
            # Save previous section
            if current_section and current_bullets:
                sections.append({
                    "title": current_section,
                    "bullets": [{"id": f"{len(sections)}-{i}", "text": bullet, "params": {}} for i, bullet in enumerate(current_bullets)],
                    "id": str(len(sections))
                })
            
            # Start new section
            current_section = line
            current_bullets = []
        
        elif current_section and line:
            # This is content for current section
            if line.startswith(('•', '-', '*')):
                bullet_text = line[1:].strip()
                if bullet_text:
                    current_bullets.append(f"• {bullet_text}")
            elif ' / ' in line and any(char.isdigit() for char in line):
                # Job entry
                current_bullets.append(f"**{line}**")
            else:
                current_bullets.append(line)
    
    # Save last section
    if current_section and current_bullets:
        sections.append({
            "title": current_section,
            "bullets": [{"id": f"{len(sections)}-{i}", "text": bullet, "params": {}} for i, bullet in enumerate(current_bullets)],
            "id": str(len(sections))
        })
    
    return {
        "name": name,
        "title": title,
        "email": email,
        "phone": phone,
        "location": location,
        "summary": "",
        "sections": sections,
        "detected_variables": {}
    }

@app.post("/api/resume/debug-extract")
async def debug_extract(file: UploadFile = File(...)):
    """Debug endpoint to see raw extracted text without AI parsing"""
    try:
        logger.info(f"Debug extracting file: {file.filename}")
        
        # File validation
        if not file.filename:
            return {"success": False, "error": "No filename provided"}
        
        file_extension = file.filename.split('.')[-1].lower()
        if file_extension not in ['pdf', 'docx', 'doc', 'txt']:
            return {"success": False, "error": "Unsupported file type"}
        
        # Read file content
        file_content = await file.read()
        
        # Extract text using multiple methods
        text = ""
        extraction_methods = []
        
        if file_extension == 'pdf':
            text, methods = extract_pdf_text(file_content)
            extraction_methods.extend(methods)
        elif file_extension == 'docx':
            text, methods = extract_docx_text(file_content)
            extraction_methods.extend(methods)
        elif file_extension == 'doc':
            text, methods = extract_doc_text(file_content)
            extraction_methods.extend(methods)
        elif file_extension == 'txt':
            try:
                text = file_content.decode('utf-8')
                extraction_methods.append("UTF-8 decode")
            except:
                text = file_content.decode('latin-1')
                extraction_methods.append("Latin-1 decode")
        
        # Clean text
        text = clean_extracted_text(text)
        
        return {
            "success": True,
            "filename": file.filename,
            "file_size": file.size,
            "extraction_methods": extraction_methods,
            "raw_text_length": len(text),
            "raw_text": text,
            "first_500_chars": text[:500],
            "last_500_chars": text[-500:] if len(text) > 500 else text
        }
        
    except Exception as e:
        logger.error(f"Debug extraction error: {str(e)}")
        return {"success": False, "error": str(e)}

@app.post("/api/resume/parse-text")
async def parse_text(payload: dict):
    try:
        text = payload.get("text", "")
        
        if not text.strip():
            return {"success": False, "error": "No text provided"}
        
        logger.info(f"Parsing text: {len(text)} characters with regex parser")
        parsed_data = parse_resume_with_regex(text)
        
        return {
            "success": True,
            "data": parsed_data,
            "message": "Resume parsed with AI - sections automatically organized!"
        }
    except Exception as e:
        logger.error(f"Text parsing error: {str(e)}")
        return {"success": False, "error": str(e)}

@app.post("/api/resume/preview")
async def preview_resume(payload: ResumePayload):
    assembled = f"{payload.name} — {payload.title}\n\n"
    if payload.summary:
        assembled += payload.summary + "\n\n"
    for s in payload.sections:
        assembled += s.title + "\n"
        for b in s.bullets:
            extra = []
            if b.metric: extra.append(b.metric)
            if b.tool: extra.append(b.tool)
            meta = f" ({', '.join(extra)})" if extra else ""
            assembled += f"• {b.text}{meta}\n"
        assembled += "\n"
    return {"variant": payload.variant or "default", "preview_text": assembled}

def apply_replacements(text: str, replacements: Dict[str, str]) -> str:
    for key, value in replacements.items():
        text = text.replace(key, value)
    return text

def format_bold_text(text: str) -> str:
    """Convert **text** to <strong>text</strong> for HTML"""
    return text.replace('**', '<strong>').replace('**', '</strong>')

def format_work_experience_bullets(bullets, replacements):
    """Format work experience bullets with proper company headers and tasks"""
    html_parts = []
    current_company = None
    
    for bullet in bullets:
        if not bullet.text.strip():
            # Empty separator - add spacing
            html_parts.append('<div class="job-separator"></div>')
            continue
            
        bullet_text = apply_replacements(bullet.text, replacements)
        
        # Check if this is a company header (starts with **)
        if bullet_text.startswith('**') and '**' in bullet_text[2:]:
            # Company header - extract company name and format
            company_text = bullet_text.replace('**', '').strip()
            html_parts.append(f'<div class="job-entry"><div class="company-name">{company_text}</div>')
            current_company = company_text
        else:
            # Regular task bullet - remove any existing bullet points
            task_text = bullet_text.replace('•', '').replace('*', '').strip()
            if task_text:
                html_parts.append(f'<li>{task_text}</li>')
    
    # Close the last job entry if needed
    if current_company:
        html_parts.append('</div>')
    
    return '\n'.join(html_parts)

def format_regular_bullets(bullets, replacements):
    """Format regular section bullets"""
    html_parts = []
    
    for bullet in bullets:
        if bullet.text.strip():
            bullet_text = apply_replacements(bullet.text, replacements)
            # Remove any existing bullet points and format
            clean_text = bullet_text.replace('•', '').replace('*', '').strip()
            if clean_text:
                html_parts.append(f'<li>{format_bold_text(clean_text)}</li>')
    
    return '\n'.join(html_parts)

@app.post("/api/resume/export/pdf")
async def export_pdf(payload: ExportPayload, user_email: str = None, db: Session = Depends(get_db)):
    try:
        from weasyprint import HTML
        from io import BytesIO
        
        replacements = payload.replacements or {}
        template_id = payload.template or "tech"
        template_style = TEMPLATES.get(template_id, TEMPLATES["tech"])
        
        logger.info(f"Exporting PDF with template: {template_id}")
        logger.info(f"Template style: {template_style}")
        logger.info(f"Two-column settings: left={payload.two_column_left}, right={payload.two_column_right}, width={payload.two_column_left_width}")
        
        font_family = template_style["styles"]["font"]
        header_align = template_style["styles"]["header_align"]
        header_border = template_style["styles"]["header_border"]
        section_uppercase = "text-transform: uppercase;" if template_style["styles"]["section_uppercase"] else ""
        layout = template_style["styles"]["layout"]
        
        # Build HTML content sections
        contact_info = apply_replacements(payload.email or '', replacements)
        if payload.phone:
            contact_info += ' • ' + apply_replacements(payload.phone, replacements)
        if payload.location:
            contact_info += ' • ' + apply_replacements(payload.location, replacements)
        
        # Add cover letter if provided
        cover_letter_html = ""
        if payload.cover_letter:
            cover_letter_content = payload.cover_letter.replace('\n', '<br>')
            cover_letter_html = f'''
            <div class="cover-letter-section">
                <h2>Cover Letter</h2>
                <div class="cover-letter-content">
                    {cover_letter_content}
                </div>
            </div>
            '''
        
        # Build sections HTML
        if layout == 'two-column':
            left_sections_html = ""
            right_sections_html = ""
            
            # Summary will be handled as part of section distribution
            # Don't add it separately here
            
            # Use localStorage configuration if provided, otherwise fallback to alternating
            left_section_ids = set(payload.two_column_left or [])
            right_section_ids = set(payload.two_column_right or [])
            
            # If no configuration provided, use alternating logic
            if not left_section_ids and not right_section_ids:
                for i, s in enumerate(payload.sections):
                    section_title = s.title.lower()
                    is_work_experience = 'experience' in section_title or 'work' in section_title or 'employment' in section_title
                    
                    if is_work_experience:
                        bullets_html = format_work_experience_bullets(s.bullets, replacements)
                        section_html = f'''
                        <div class="section">
                            <h2>{apply_replacements(s.title, replacements)}</h2>
                            {bullets_html}
                        </div>'''
                    else:
                        bullets_html = format_regular_bullets(s.bullets, replacements)
                        section_html = f'''
                        <div class="section">
                            <h2>{apply_replacements(s.title, replacements)}</h2>
                            <ul>
                                {bullets_html}
                            </ul>
                        </div>'''
                    
                    if i % 2 == 0:
                        left_sections_html += section_html
                    else:
                        right_sections_html += section_html
            else:
                # Use provided configuration
                for s in payload.sections:
                    section_title = s.title.lower()
                    is_work_experience = 'experience' in section_title or 'work' in section_title or 'employment' in section_title
                    
                    if is_work_experience:
                        bullets_html = format_work_experience_bullets(s.bullets, replacements)
                        section_html = f'''
                        <div class="section">
                            <h2>{apply_replacements(s.title, replacements)}</h2>
                            {bullets_html}
                        </div>'''
                    else:
                        bullets_html = format_regular_bullets(s.bullets, replacements)
                        section_html = f'''
                        <div class="section">
                            <h2>{apply_replacements(s.title, replacements)}</h2>
                            <ul>
                                {bullets_html}
                            </ul>
                        </div>'''
                    
                    if s.id and s.id in left_section_ids:
                        left_sections_html += section_html
                    elif s.id and s.id in right_section_ids:
                        right_sections_html += section_html
            
            # Calculate column widths with spacing
            left_width = payload.two_column_left_width or 50
            right_width = 100 - left_width
            
            # Adjust for spacing between columns (2% gap)
            gap = 2
            left_width_adjusted = left_width - (gap / 2)
            right_width_adjusted = right_width - (gap / 2)
            
            content_html = f'''
            {cover_letter_html}
            <div class="two-column">
                <div class="column" style="width: {left_width_adjusted}%; margin-right: {gap}%;">{left_sections_html}</div>
                <div class="column" style="width: {right_width_adjusted}%;">{right_sections_html}</div>
                <div class="clearfix"></div>
            </div>'''
        else:
            # Single column layout
            summary_html = f'<div class="summary">{apply_replacements(payload.summary, replacements)}</div>' if payload.summary else ''
            sections_html = ''
            for s in payload.sections:
                section_title = s.title.lower()
                is_work_experience = 'experience' in section_title or 'work' in section_title or 'employment' in section_title
                
                if is_work_experience:
                    bullets_html = format_work_experience_bullets(s.bullets, replacements)
                    sections_html += f'''
                    <div class="section">
                        <h2>{apply_replacements(s.title, replacements)}</h2>
                        {bullets_html}
                    </div>'''
                else:
                    bullets_html = format_regular_bullets(s.bullets, replacements)
                    sections_html += f'''
                    <div class="section">
                        <h2>{apply_replacements(s.title, replacements)}</h2>
                        <ul>
                            {bullets_html}
                        </ul>
                    </div>'''
            content_html = cover_letter_html + summary_html + sections_html
        
        html_content = f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        @page {{ size: A4; margin: 2cm; }}
        body {{ font-family: {font_family}; font-size: 11pt; line-height: 1.4; color: #333; }}
        .header {{ text-align: {header_align}; border-bottom: {header_border}; padding-bottom: 10px; margin-bottom: 15px; }}
        .header h1 {{ margin: 0; font-size: 24pt; font-weight: bold; }}
        .header .title {{ font-size: 14pt; margin: 5px 0; color: #555; }}
        .header .contact {{ font-size: 10pt; color: #666; margin-top: 5px; }}
        .summary {{ margin-bottom: 15px; font-size: 10pt; line-height: 1.5; }}
        .section {{ margin-bottom: 15px; }}
        .section h2 {{ font-size: 12pt; font-weight: bold; {section_uppercase}
                       border-bottom: 1px solid #333; padding-bottom: 3px; margin-bottom: 8px; }}
        .section ul {{ margin: 0; padding-left: 20px; list-style-type: disc; }}
        .section li {{ margin-bottom: 6px; font-size: 10pt; }}
        .job-entry {{ margin-bottom: 20px; }}
        .company-name {{ font-weight: bold; font-size: 1.1em; color: #333; margin-bottom: 5px; margin-left: 0; }}
        .job-separator {{ height: 10px; }}
        .two-column {{ width: 100%; }}
        .column {{ float: left; }}
        .clearfix {{ clear: both; }}
        .cover-letter-section {{ margin-bottom: 20px; page-break-after: always; }}
        .cover-letter-section h2 {{ font-size: 14pt; font-weight: bold; margin-bottom: 10px; border-bottom: 2px solid #333; padding-bottom: 5px; }}
        .cover-letter-content {{ font-size: 11pt; line-height: 1.6; text-align: justify; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{apply_replacements(payload.name, replacements)}</h1>
        <div class="title">{apply_replacements(payload.title, replacements)}</div>
        <div class="contact">{contact_info}</div>
    </div>
    {content_html}
</body>
</html>'''
        
        logger.info(f"Generated HTML length: {len(html_content)}")
        logger.info(f"Layout type: {layout}")
        
        pdf_bytes = HTML(string=html_content).write_pdf()
        logger.info(f"PDF generated successfully, size: {len(pdf_bytes)} bytes")
        
        # Track export analytics
        if user_email and db:
            try:
                user = db.query(User).filter(User.email == user_email).first()
                if user:
                    # Find or create resume record
                    resume = db.query(Resume).filter(
                        Resume.user_id == user.id,
                        Resume.name == payload.name
                    ).first()
                    
                    if not resume:
                        resume = Resume(
                            user_id=user.id,
                            name=payload.name,
                            title=payload.title,
                            email=payload.email,
                            phone=payload.phone,
                            location=payload.location,
                            summary=payload.summary,
                            template=template_id
                        )
                        db.add(resume)
                        db.commit()
                        db.refresh(resume)
                    
                    # Create export analytics record
                    export_analytics = ExportAnalytics(
                        user_id=user.id,
                        resume_id=resume.id,
                        export_format='pdf',
                        template_used=template_id,
                        file_size=len(pdf_bytes),
                        export_success=True
                    )
                    db.add(export_analytics)
                    db.commit()
                    
                    logger.info(f"Export analytics tracked for user {user_email}: PDF export")
            except Exception as e:
                logger.error(f"Failed to track export analytics: {e}")
        
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=resume.pdf"}
        )
    except Exception as e:
        logger.error(f"PDF export error: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/test/connection")
async def test_connection():
    """Test endpoint to verify API connectivity"""
    return {
        "status": "success",
        "message": "API is working",
        "timestamp": datetime.utcnow().isoformat(),
        "cors_origins": ALLOWED_ORIGINS
    }

@app.post("/api/resume/upload")
async def upload_resume(file: UploadFile = File(...)):
    try:
        from io import BytesIO
        
        contents = await file.read()
        filename = file.filename if file.filename else 'unknown'
        file_type = filename.split('.')[-1].lower() if '.' in filename else ''
        
        logger.info(f"Upload: filename={filename}, type={file_type}, content_type={file.content_type}, size={len(contents)}")
        
        magic_bytes = contents[:4]
        is_docx = magic_bytes == b'PK\x03\x04'
        is_pdf = contents[:5] == b'%PDF-'
        
        logger.info(f"File magic: is_docx={is_docx}, is_pdf={is_pdf}")
        
        text = ""
        actual_type = ""
        
        if is_docx or file_type in ['docx', 'doc'] or 'wordprocessingml' in str(file.content_type):
            logger.info("Processing as DOCX")
            actual_type = "DOCX"
            try:
                from docx import Document
                from docx.oxml.text.paragraph import CT_P
                from docx.oxml.table import CT_Tbl
                from docx.table import _Cell, Table
                from docx.text.paragraph import Paragraph
                
                docx_file = BytesIO(contents)
                doc = Document(docx_file)
                
                def extract_text_from_element(element):
                    elem_text = ""
                    if isinstance(element, Paragraph):
                        if element.text.strip():
                            elem_text += element.text + "\n"
                    elif isinstance(element, Table):
                        for row in element.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if para.text.strip():
                                        elem_text += para.text + "\n"
                    return elem_text
                
                for element in doc.element.body:
                    if isinstance(element, CT_P):
                        para = Paragraph(element, doc)
                        if para.text.strip():
                            text += para.text + "\n"
                    elif isinstance(element, CT_Tbl):
                        table = Table(element, doc)
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if para.text.strip():
                                        text += para.text + "\n"
                
                for section in doc.sections:
                    if section.header:
                        for para in section.header.paragraphs:
                            if para.text.strip():
                                text += para.text + "\n"
                    if section.footer:
                        for para in section.footer.paragraphs:
                            if para.text.strip():
                                text += para.text + "\n"
                
                logger.info(f"DOCX extracted {len(text)} characters from {len(doc.paragraphs)} paragraphs")
            except Exception as docx_error:
                logger.error(f"DOCX error: {str(docx_error)}")
                return {
                    "success": False,
                    "error": f"Could not read DOCX file: {str(docx_error)}"
                }
        
        elif is_pdf or file_type == 'pdf' or file.content_type == 'application/pdf':
            logger.info("Processing as PDF")
            actual_type = "PDF"
            try:
                import pdfplumber
                pdf_file = BytesIO(contents)
                with pdfplumber.open(pdf_file) as pdf:
                    total_pages = len(pdf.pages)
                    logger.info(f"PDF has {total_pages} pages")
                    
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num} ---\n"
                            text += page_text + "\n"
                            logger.info(f"Extracted {len(page_text)} characters from page {page_num}")
                
                logger.info(f"PDF extraction complete: {total_pages} pages, {len(text)} total characters")
            except Exception as pdf_error:
                error_msg = str(pdf_error)
                logger.error(f"PDF error: {error_msg}")
                if "Root object" in error_msg or "really a PDF" in error_msg:
                    return {
                        "success": False,
                        "error": "PDF file appears to be corrupted or invalid. Try:\n1. Re-saving the PDF from its original application\n2. Converting to DOCX first\n3. Or use manual entry"
                    }
                raise
        
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return {
                "success": False,
                "error": f"Unsupported file type: {file_type}. Please upload PDF or DOCX."
            }
        
        if not text.strip():
            logger.warning(f"No text extracted from {actual_type} file")
            return {
                "success": False,
                "error": f"Could not extract text from {actual_type} file. The file might be:\n• Empty\n• Using special formatting (text boxes)\n• Password protected\n\nPlease try exporting from Word as a new DOCX, or use manual entry."
            }
        
        logger.info("Using AI-powered parsing for better resume organization...")
        parsed_data = parse_resume_with_ai(text)
        
        return {
            "success": True,
            "data": parsed_data,
            "raw_text": text[:1000],
            "debug": {
                "file_type": file_type,
                "text_length": len(text),
                "has_content": bool(text.strip())
            }
        }
    except Exception as e:
        return {"success": False, "error": f"Upload failed: {str(e)}"}

def parse_resume_with_ai(text: str) -> Dict:
    """Use AI to intelligently parse and structure resume content"""
    if not openai_client:
        logger.warning("AI parsing not available, falling back to basic parsing")
        return parse_resume_text(text)
    
    try:
        prompt = f"""Parse this resume text and extract structured information. This resume may have been exported from a resume builder app, so pay special attention to preserving the exact structure and content.

CRITICAL: You MUST return valid JSON with ALL property names enclosed in double quotes. Do not use single quotes or unquoted property names.

Return a JSON object with this exact structure:

{{
  "name": "Full Name",
  "title": "Professional Title or Current Role",
  "email": "email@example.com",
  "phone": "+1-234-567-8900",
  "location": "City, State/Country",
  "summary": "Professional summary or objective (2-3 sentences)",
  "sections": [
    {{
      "title": "Work Experience",
      "bullets": [
        "**Company Name / Job Title / Start Date - End Date**",
        "• Achievement or task for this role",
        "• Another achievement with metrics",
        "• Key responsibility or accomplishment",
        "",
        "**Another Company / Job Title / Start Date - End Date**",
        "• Achievement at this company",
        "• Another task or responsibility"
      ]
    }},
    {{
      "title": "Education",
      "bullets": [
        "Degree, Institution, Year",
        "GPA or honors if applicable"
      ]
    }},
    {{
      "title": "Skills",
      "bullets": [
        "Skill category: specific skills",
        "Another category: more skills"
      ]
    }}
  ]
}}

CRITICAL PARSING RULES FOR APP-EXPORTED RESUMES:
1. PRESERVE ALL CONTENT - Don't skip, summarize, or paraphrase anything
2. Maintain exact bullet point structure and wording
3. For work experience sections, preserve company names, job titles, dates, and ALL achievements
4. Keep quantified metrics exactly as written (e.g., "30% reduction", "over 100 applications")
5. Preserve technical skills and tools mentioned (AWS, Kubernetes, Docker, etc.)
6. Maintain project descriptions with all details
7. Extract contact information from header (name, email, phone, location)
8. Handle multi-page content by combining everything
9. If you see patterns like "reduced deployment time by 30%" repeated, keep each instance
10. For DevOps/Engineering resumes, preserve all technical details and metrics
11. Don't paraphrase or shorten content - extract exactly as written
12. Group content into logical sections but preserve all bullet points

SPECIAL ATTENTION TO:
- Quantified achievements (percentages, numbers, metrics)
- Technical skills and tools (AWS, Kubernetes, Docker, Jenkins, etc.)
- Company names and job titles
- Project descriptions and outcomes
- All bullet points with their complete text
- Work experience entries with exact dates and roles

CRITICAL RULES FOR WORK EXPERIENCE:
1. Company/Job line MUST be: **Company Name / Job Title / Date Range** (with ** for bold)
2. Tasks/achievements MUST start with "• " (bullet point)
3. Separate different jobs with empty string ""
4. Extract complete date ranges (e.g., "Jan 2020 - Dec 2022" or "2020-2022")
5. Include ALL jobs from the resume

OTHER RULES:
6. Identify ALL sections (Education, Skills, Projects, Certifications, Awards, etc.)
7. Preserve important details (dates, technologies, metrics, achievements)
8. If something is missing (like phone or email), leave it as empty string
9. Professional summary should be concise (2-3 sentences max)
10. Return ONLY valid JSON, no markdown code blocks
11. This may be a MULTI-PAGE resume - extract ALL information from ALL pages
12. Include everything: all work experience, education, skills, projects, certifications

Resume Text (Full Content):
{text[:12000]}
"""

        # For very long resumes, use higher token limit
        resume_length = len(text)
        max_tokens_for_resume = min(4000, OPENAI_MAX_TOKENS) if resume_length > 8000 else OPENAI_MAX_TOKENS
        
        logger.info(f"Parsing resume: {resume_length} characters, using {max_tokens_for_resume} max tokens")

        response = openai_client['requests'].post(
            'https://api.openai.com/v1/chat/completions',
            headers={
                'Authorization': f"Bearer {openai_client['api_key']}",
                'Content-Type': 'application/json'
            },
            json={
                'model': openai_client['model'],
                'messages': [
                    {'role': 'system', 'content': 'You are a resume parsing expert specializing in app-exported resumes. Extract structured information from multi-page resumes accurately. Capture ALL sections and information without paraphrasing or summarizing. Preserve exact wording, metrics, and technical details. Always return valid JSON.'},
                    {'role': 'user', 'content': prompt}
                ],
                'temperature': 0.3,
                'max_tokens': max_tokens_for_resume
            },
            timeout=120
        )
        
        if response.status_code != 200:
            logger.error(f"OpenAI API error: {response.status_code} - {response.text}")
            return parse_resume_text(text)
        
        result = response.json()
        ai_response = result['choices'][0]['message']['content'].strip()
        
        # Clean up the response
        ai_response = re.sub(r'^```json\s*', '', ai_response)
        ai_response = re.sub(r'\s*```$', '', ai_response)
        
        # Try to fix common JSON formatting issues
        try:
            parsed_data = json.loads(ai_response)
        except json.JSONDecodeError as e:
            logger.warning(f"Initial JSON parse failed: {e}")
            # Try to fix common JSON issues
            ai_response = re.sub(r'(\w+):', r'"\1":', ai_response)  # Add quotes to unquoted keys
            ai_response = re.sub(r':\s*([^",\[\{][^,}\]]*?)([,}\]])', r': "\1"\2', ai_response)  # Add quotes to unquoted string values
            try:
                parsed_data = json.loads(ai_response)
            except json.JSONDecodeError as e2:
                logger.error(f"JSON cleanup failed: {e2}")
                raise e2
        
        for i, section in enumerate(parsed_data.get('sections', [])):
            section['id'] = str(i)
            for j, bullet in enumerate(section.get('bullets', [])):
                if isinstance(bullet, str):
                    if 'bullets' not in section:
                        section['bullets'] = []
                    section['bullets'][j] = {
                        'id': f"{i}-{j}",
                        'text': bullet,
                        'params': {}
                    }
        
        parsed_data.setdefault('detected_variables', {})
        
        logger.info(f"AI parsing successful: {len(parsed_data.get('sections', []))} sections extracted")
        return parsed_data
        
    except json.JSONDecodeError as e:
        logger.error(f"AI returned invalid JSON: {e}")
        logger.info("Falling back to basic text parsing")
        return parse_resume_text(text)
    except Exception as e:
        logger.error(f"AI parsing failed: {str(e)}")
        logger.info("Falling back to basic text parsing")
        return parse_resume_text(text)

def parse_resume_text(text: str) -> Dict:
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    
    data = {
        "name": "",
        "title": "",
        "email": "",
        "phone": "",
        "location": "",
        "summary": "",
        "sections": [],
        "detected_variables": {}
    }
    
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if email_match:
        data["email"] = email_match.group()
    
    phone_match = re.search(r'[\+\(]?[0-9][0-9 .\-\(\)]{8,}[0-9]', text)
    if phone_match:
        data["phone"] = phone_match.group()
    
    if lines:
        data["name"] = lines[0]
        if len(lines) > 1:
            data["title"] = lines[1]
    
    section_headers = ['experience', 'education', 'skills', 'projects', 'certifications', 'summary']
    current_section = None
    current_bullets = []
    summary_lines = []
    
    for i, line in enumerate(lines[2:], start=2):
        line_lower = line.lower()
        
        if any(header in line_lower for header in section_headers):
            if current_section:
                data["sections"].append({
                    "id": str(len(data["sections"])),
                    "title": current_section,
                    "bullets": [{"id": str(j), "text": b, "params": {}} for j, b in enumerate(current_bullets)]
                })
            current_section = line
            current_bullets = []
        elif current_section:
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                bullet_text = re.sub(r'^[•\-\*]\s*', '', line)
                current_bullets.append(bullet_text)
            elif len(line) > 20 and not any(header in line_lower for header in section_headers):
                current_bullets.append(line)
        elif 'summary' in line_lower:
            pass
        elif not current_section and len(line) > 30:
            summary_lines.append(line)
    
    if current_section:
        data["sections"].append({
            "id": str(len(data["sections"])),
            "title": current_section,
            "bullets": [{"id": str(j), "text": b, "params": {}} for j, b in enumerate(current_bullets)]
        })
    
    if summary_lines:
        data["summary"] = " ".join(summary_lines)
    
    data["detected_variables"] = detect_parameterization(text)
    
    return data

def detect_parameterization(text: str) -> Dict[str, str]:
    variables = {}
    
    companies = re.findall(r'\b(?:at|for|with)\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)\b', text)
    if companies:
        most_common = max(set(companies), key=companies.count)
        variables["{{company}}"] = most_common
    
    percentages = re.findall(r'(\d+)%', text)
    if percentages:
        variables["{{metric}}"] = percentages[0]
    
    tech_keywords = ['AWS', 'Azure', 'GCP', 'Kubernetes', 'Docker', 'Python', 'Java', 'React', 'Node.js']
    found_tech = [tech for tech in tech_keywords if tech in text]
    if found_tech:
        variables["{{tech}}"] = found_tech[0]
    
    return variables

@app.post("/api/resume/export/docx")
async def export_docx(payload: ExportPayload, user_email: str = None, db: Session = Depends(get_db)):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        
        replacements = payload.replacements or {}
        template_id = payload.template or "tech"
        template_style = TEMPLATES.get(template_id, TEMPLATES["tech"])
        
        logger.info(f"Exporting DOCX with template: {template_id}")
        logger.info(f"Template layout: {template_style['styles']['layout']}")
        logger.info(f"Two-column settings: left={payload.two_column_left}, right={payload.two_column_right}, width={payload.two_column_left_width}")
        
        layout = template_style["styles"]["layout"]
        
        doc = Document()
        
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(apply_replacements(payload.name, replacements))
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.name = 'Arial' if 'Arial' in template_style["styles"]["font"] else 'Georgia'
        if template_style["styles"]["header_align"] == "center":
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(apply_replacements(payload.title, replacements))
        title_run.font.size = Pt(12)
        title_run.font.name = 'Arial' if 'Arial' in template_style["styles"]["font"] else 'Georgia'
        if template_style["styles"]["header_align"] == "center":
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact_parts = []
        if payload.email:
            contact_parts.append(apply_replacements(payload.email, replacements))
        if payload.phone:
            contact_parts.append(apply_replacements(payload.phone, replacements))
        if payload.location:
            contact_parts.append(apply_replacements(payload.location, replacements))
        
        if contact_parts:
            contact_para = doc.add_paragraph(' • '.join(contact_parts))
            contact_para.runs[0].font.size = Pt(10)
            if template_style["styles"]["header_align"] == "center":
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Add cover letter if provided
        if payload.cover_letter:
            cover_letter_heading = doc.add_paragraph()
            cover_letter_heading_run = cover_letter_heading.add_run("Cover Letter")
            cover_letter_heading_run.font.size = Pt(14)
            cover_letter_heading_run.font.bold = True
            
            cover_letter_para = doc.add_paragraph(payload.cover_letter)
            cover_letter_para.runs[0].font.size = Pt(11)
            cover_letter_para.runs[0].font.name = 'Arial'
            
            doc.add_paragraph()  # Add spacing
        
        logger.info(f"Layout decision: layout='{layout}', is_two_column={layout == 'two-column'}")
        if layout == 'two-column':
            logger.info("Creating two-column layout")
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            # Create a table for two-column layout
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            
            # Remove table borders
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)
            tblPr.append(tblBorders)
            
            # Set column widths based on user preference
            left_width_percent = payload.two_column_left_width or 50
            right_width_percent = 100 - left_width_percent
            
            # Convert percentages to inches (assuming 8.5" total width)
            left_width_inches = (left_width_percent / 100) * 8.5
            right_width_inches = (right_width_percent / 100) * 8.5
            
            table.columns[0].width = Inches(left_width_inches)
            table.columns[1].width = Inches(right_width_inches)
            
            logger.info(f"Column widths: left={left_width_inches:.1f}\", right={right_width_inches:.1f}\"")
            
            # Left column
            left_cell = table.cell(0, 0)
            # Clear existing paragraphs in left cell
            for paragraph in left_cell.paragraphs:
                paragraph.clear()
            
            # Right column  
            right_cell = table.cell(0, 1)
            # Clear existing paragraphs in right cell
            for paragraph in right_cell.paragraphs:
                paragraph.clear()
            
            # Summary will be handled as part of section distribution
            # Don't add it separately here
            
            # Use localStorage configuration if provided, otherwise fallback to alternating
            left_section_ids = set(payload.two_column_left or [])
            right_section_ids = set(payload.two_column_right or [])
            
            logger.info(f"Section IDs from payload: left={left_section_ids}, right={right_section_ids}")
            logger.info(f"Available sections: {[s.id for s in payload.sections if s.id]}")  # Only log sections with IDs
            
            if not left_section_ids and not right_section_ids:
                # Use alternating logic if no configuration
                logger.info("Using alternating logic - no localStorage configuration provided")
                left_sections = [s for i, s in enumerate(payload.sections) if i % 2 == 0]
                right_sections = [s for i, s in enumerate(payload.sections) if i % 2 == 1]
            else:
                # Use provided configuration
                logger.info("Using localStorage configuration")
                left_sections = [s for s in payload.sections if s.id and s.id in left_section_ids]
                right_sections = [s for s in payload.sections if s.id and s.id in right_section_ids]
            
            logger.info(f"Final section distribution: left={[s.title for s in left_sections]}, right={[s.title for s in right_sections]}")
            logger.info(f"Left sections count: {len(left_sections)}, Right sections count: {len(right_sections)}")
            
            # Add sections to left column
            for s in left_sections:
                section_title = apply_replacements(s.title, replacements)
                if template_style["styles"]["section_uppercase"]:
                    section_title = section_title.upper()
                
                heading_para = left_cell.add_paragraph()
                heading_run = heading_para.add_run(section_title)
                heading_run.font.size = Pt(12)
                heading_run.font.bold = True
                heading_run.font.name = 'Arial' if 'Arial' in template_style["styles"]["font"] else 'Georgia'
                
                # Create a proper bullet list for this section
                if s.bullets:
                    # Add bullets as a proper list
                    for b in s.bullets:
                        if not b.text.strip():  # Skip empty bullets
                            continue
                        bullet_text = apply_replacements(b.text, replacements)
                        bullet_para = left_cell.add_paragraph()
                        bullet_para.style = 'List Bullet'
                        
                        # Handle bold text formatting
                        if '**' in bullet_text:
                            parts = bullet_text.split('**')
                            for i, part in enumerate(parts):
                                if i % 2 == 0:  # Regular text
                                    if part.strip():
                                        run = bullet_para.add_run(part)
                                        run.font.size = Pt(10)
                                else:  # Bold text
                                    run = bullet_para.add_run(part)
                                    run.font.size = Pt(10)
                                    run.font.bold = True
                        else:
                            run = bullet_para.add_run(bullet_text)
                            run.font.size = Pt(10)
                
                left_cell.add_paragraph()
            
            # Add sections to right column
            for s in right_sections:
                section_title = apply_replacements(s.title, replacements)
                if template_style["styles"]["section_uppercase"]:
                    section_title = section_title.upper()
                
                heading_para = right_cell.add_paragraph()
                heading_run = heading_para.add_run(section_title)
                heading_run.font.size = Pt(12)
                heading_run.font.bold = True
                heading_run.font.name = 'Arial' if 'Arial' in template_style["styles"]["font"] else 'Georgia'
                
                # Create a proper bullet list for this section
                if s.bullets:
                    # Add bullets as a proper list
                    for b in s.bullets:
                        if not b.text.strip():  # Skip empty bullets
                            continue
                        bullet_text = apply_replacements(b.text, replacements)
                        bullet_para = right_cell.add_paragraph()
                        bullet_para.style = 'List Bullet'
                        
                        # Handle bold text formatting
                        if '**' in bullet_text:
                            parts = bullet_text.split('**')
                            for i, part in enumerate(parts):
                                if i % 2 == 0:  # Regular text
                                    if part.strip():
                                        run = bullet_para.add_run(part)
                                        run.font.size = Pt(10)
                                else:  # Bold text
                                    run = bullet_para.add_run(part)
                                    run.font.size = Pt(10)
                                    run.font.bold = True
                        else:
                            run = bullet_para.add_run(bullet_text)
                            run.font.size = Pt(10)
                
                right_cell.add_paragraph()
        else:
            # Single column layout
            if payload.summary:
                summary_para = doc.add_paragraph(apply_replacements(payload.summary, replacements))
                summary_para.runs[0].font.size = Pt(10)
                doc.add_paragraph()
            
            for s in payload.sections:
                section_title = apply_replacements(s.title, replacements)
                if template_style["styles"]["section_uppercase"]:
                    section_title = section_title.upper()
                
                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run(section_title)
                heading_run.font.size = Pt(12)
                heading_run.font.bold = True
                heading_run.font.name = 'Arial' if 'Arial' in template_style["styles"]["font"] else 'Georgia'
                
                # Create a proper bullet list for this section
                if s.bullets:
                    # Add bullets as a proper list
                    for b in s.bullets:
                        if not b.text.strip():  # Skip empty bullets
                            continue
                        bullet_text = apply_replacements(b.text, replacements)
                        bullet_para = doc.add_paragraph()
                        bullet_para.style = 'List Bullet'
                        
                        # Handle bold text formatting
                        if '**' in bullet_text:
                            parts = bullet_text.split('**')
                            for i, part in enumerate(parts):
                                if i % 2 == 0:  # Regular text
                                    if part.strip():
                                        run = bullet_para.add_run(part)
                                        run.font.size = Pt(10)
                                else:  # Bold text
                                    run = bullet_para.add_run(part)
                                    run.font.size = Pt(10)
                                    run.font.bold = True
                        else:
                            run = bullet_para.add_run(bullet_text)
                            run.font.size = Pt(10)
                
                doc.add_paragraph()
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        docx_bytes = buffer.getvalue()
        logger.info(f"DOCX generated successfully, size: {len(docx_bytes)} bytes")
        
        # Track export analytics
        if user_email and db:
            try:
                user = db.query(User).filter(User.email == user_email).first()
                if user:
                    # Find or create resume record
                    resume = db.query(Resume).filter(
                        Resume.user_id == user.id,
                        Resume.name == payload.name
                    ).first()
                    
                    if not resume:
                        resume = Resume(
                            user_id=user.id,
                            name=payload.name,
                            title=payload.title,
                            email=payload.email,
                            phone=payload.phone,
                            location=payload.location,
                            summary=payload.summary,
                            template=template_id
                        )
                        db.add(resume)
                        db.commit()
                        db.refresh(resume)
                    
                    # Create export analytics record
                    export_analytics = ExportAnalytics(
                        user_id=user.id,
                        resume_id=resume.id,
                        export_format='docx',
                        template_used=template_id,
                        file_size=len(docx_bytes),
                        export_success=True
                    )
                    db.add(export_analytics)
                    db.commit()
                    
                    logger.info(f"Export analytics tracked for user {user_email}: DOCX export")
            except Exception as e:
                logger.error(f"Failed to track export analytics: {e}")
        
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=resume.docx"}
        )
    except Exception as e:
        logger.error(f"DOCX export error: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

# ============================================================
# REAL-TIME COLLABORATION (WebSocket)
# ============================================================

class Comment(BaseModel):
    id: str
    user_name: str
    text: str
    timestamp: str
    target_type: str
    target_id: str
    resolved: bool = False

class CollaborationRoom:
    def __init__(self):
        self.rooms: Dict[str, Dict[str, any]] = {}
        self.comments: Dict[str, List[Dict]] = {}
    
    def create_room(self, room_id: str):
        if room_id not in self.rooms:
            self.rooms[room_id] = {
                "connections": {},
                "resume_data": None,
                "last_update": datetime.now().isoformat()
            }
            self.comments[room_id] = []
            logger.info(f"Created collaboration room: {room_id}")
    
    def add_connection(self, room_id: str, user_id: str, websocket: WebSocket, user_name: str):
        self.create_room(room_id)
        
        # Remove any existing connection with the same user_id to prevent duplicates
        if user_id in self.rooms[room_id]["connections"]:
            logger.info(f"Replacing existing connection for user {user_name} ({user_id})")
        
        self.rooms[room_id]["connections"][user_id] = {
            "ws": websocket,
            "name": user_name,
            "joined_at": datetime.now().isoformat()
        }
        logger.info(f"User {user_name} ({user_id}) joined room {room_id}")
    
    def remove_connection(self, room_id: str, user_id: str):
        if room_id in self.rooms and user_id in self.rooms[room_id]["connections"]:
            del self.rooms[room_id]["connections"][user_id]
            logger.info(f"User {user_id} left room {room_id}")
            if not self.rooms[room_id]["connections"]:
                del self.rooms[room_id]
                if room_id in self.comments:
                    del self.comments[room_id]
                logger.info(f"Room {room_id} is now empty and removed")
    
    def get_active_users(self, room_id: str) -> List[Dict]:
        if room_id not in self.rooms:
            return []
        return [
            {"user_id": uid, "name": conn["name"], "joined_at": conn["joined_at"]}
            for uid, conn in self.rooms[room_id]["connections"].items()
        ]
    
    async def broadcast(self, room_id: str, message: dict, exclude_user: str = None):
        if room_id not in self.rooms:
            return
        
        dead_connections = []
        for user_id, conn in self.rooms[room_id]["connections"].items():
            if exclude_user and user_id == exclude_user:
                continue
            try:
                await conn["ws"].send_json(message)
            except Exception as e:
                logger.error(f"Failed to send to user {user_id}: {e}")
                dead_connections.append(user_id)
        
        for user_id in dead_connections:
            self.remove_connection(room_id, user_id)
    
    def add_comment(self, room_id: str, comment: Dict) -> Dict:
        self.create_room(room_id)
        comment_data = {
            **comment,
            "id": secrets.token_urlsafe(8),
            "timestamp": datetime.now().isoformat()
        }
        self.comments[room_id].append(comment_data)
        logger.info(f"Added comment in room {room_id}: {comment_data['id']}")
        return comment_data
    
    def get_comments(self, room_id: str, target_id: str = None) -> List[Dict]:
        if room_id not in self.comments:
            return []
        comments = self.comments[room_id]
        if target_id:
            return [c for c in comments if c.get("target_id") == target_id]
        return comments
    
    def resolve_comment(self, room_id: str, comment_id: str):
        if room_id in self.comments:
            for comment in self.comments[room_id]:
                if comment["id"] == comment_id:
                    comment["resolved"] = True
                    logger.info(f"Resolved comment {comment_id} in room {room_id}")
                    return True
        return False
    
    def delete_comment(self, room_id: str, comment_id: str):
        if room_id in self.comments:
            self.comments[room_id] = [c for c in self.comments[room_id] if c["id"] != comment_id]
            logger.info(f"Deleted comment {comment_id} from room {room_id}")
            return True
        return False

collab_manager = CollaborationRoom()

@app.get("/api/collab/room/create")
async def create_collab_room():
    room_id = secrets.token_urlsafe(8)
    collab_manager.create_room(room_id)
    return {"room_id": room_id, "url": f"/editor?room={room_id}"}

@app.get("/api/collab/room/{room_id}/users")
async def get_room_users(room_id: str):
    users = collab_manager.get_active_users(room_id)
    return {"room_id": room_id, "users": users, "count": len(users)}

@app.get("/api/collab/room/{room_id}/comments")
async def get_room_comments(room_id: str, target_id: str = None):
    comments = collab_manager.get_comments(room_id, target_id)
    return {"room_id": room_id, "comments": comments, "count": len(comments)}

@app.post("/api/collab/room/{room_id}/comments")
async def add_comment(room_id: str, comment: Comment):
    comment_data = collab_manager.add_comment(room_id, comment.dict())
    return {"success": True, "comment": comment_data}

@app.post("/api/collab/room/{room_id}/comments/{comment_id}/resolve")
async def resolve_comment_endpoint(room_id: str, comment_id: str):
    success = collab_manager.resolve_comment(room_id, comment_id)
    return {"success": success}

@app.delete("/api/collab/room/{room_id}/comments/{comment_id}")
async def delete_comment_endpoint(room_id: str, comment_id: str):
    success = collab_manager.delete_comment(room_id, comment_id)
    return {"success": success}

@app.post("/api/ai/ats_score")
async def get_ats_score(payload: ResumePayload):
    """Get ATS compatibility score and suggestions for resume"""
    try:
        logger.info("Processing ATS score request")
        
        # Check if ATS checker is available
        if not ats_checker:
            return {
                "success": False,
                "score": 0,
                "suggestions": ["ATS analysis is not available. Please install required dependencies."],
                "details": {},
                "error": "ATS checker not available"
            }
        
        # Convert ResumePayload to dict for ATSChecker
        resume_data = {
            'name': payload.name,
            'title': payload.title,
            'email': payload.email,
            'phone': payload.phone,
            'location': payload.location,
            'summary': payload.summary,
            'sections': payload.sections
        }
        
        # Get ATS score and analysis
        result = ats_checker.get_ats_score(resume_data)
        
        logger.info(f"ATS analysis completed. Score: {result.get('score', 0)}")
        
        return {
            "success": True,
            "score": result.get('score', 0),
            "suggestions": result.get('suggestions', []),
            "details": result.get('details', {}),
            "message": f"ATS compatibility score: {result.get('score', 0)}/100"
        }
        
    except Exception as e:
        logger.error(f"ATS score calculation error: {str(e)}")
        return {
            "success": False,
            "score": 0,
            "suggestions": ["Unable to analyze resume. Please check your content."],
            "details": {},
            "error": str(e)
        }

@app.post("/api/ai/enhanced_ats_score")
async def get_enhanced_ats_score(payload: EnhancedATSPayload):
    """Get enhanced ATS compatibility score with AI improvements"""
    try:
        logger.info("Processing enhanced ATS score request")
        
        # Check if enhanced ATS checker is available
        if not enhanced_ats_checker:
            return {
                "success": False,
                "score": 0,
                "suggestions": ["Enhanced ATS analysis is not available. Please install required dependencies."],
                "details": {},
                "ai_improvements": [],
                "error": "Enhanced ATS checker not available"
            }
        
        # Convert ResumePayload to dict for EnhancedATSChecker
        resume_data = {
            'name': payload.resume_data.name,
            'title': payload.resume_data.title,
            'email': payload.resume_data.email,
            'phone': payload.resume_data.phone,
            'location': payload.resume_data.location,
            'summary': payload.resume_data.summary,
            'sections': [
                {
                    'id': section.id,
                    'title': section.title,
                    'bullets': [
                        {
                            'id': bullet.id,
                            'text': bullet.text,
                            'params': bullet.params
                        }
                        for bullet in section.bullets
                    ]
                }
                for section in payload.resume_data.sections
            ]
        }
        
        # Get enhanced ATS score and analysis
        result = enhanced_ats_checker.get_enhanced_ats_score(resume_data, payload.job_description)
        
        logger.info(f"Enhanced ATS analysis completed. Score: {result.get('score', 0)}")
        
        return {
            "success": True,
            "score": result.get('score', 0),
            "suggestions": result.get('suggestions', []),
            "details": result.get('details', {}),
            "ai_improvements": result.get('ai_improvements', []),
            "message": f"Enhanced ATS compatibility score: {result.get('score', 0)}/100"
        }
        
    except Exception as e:
        logger.error(f"Enhanced ATS score calculation error: {str(e)}")
        return {
            "success": False,
            "score": 0,
            "suggestions": ["Unable to analyze resume. Please check your content."],
            "details": {},
            "ai_improvements": [],
            "error": str(e)
        }

@app.post("/api/ai/improvement_suggestions")
async def get_ai_improvement_suggestions(payload: AIImprovementPayload):
    """Get AI-powered improvement suggestions based on 10 strategies"""
    try:
        logger.info("Processing AI improvement suggestions request")
        
        # Check if AI improvement engine is available
        if not ai_improvement_engine:
            return {
                "success": False,
                "suggestions": ["AI improvement engine is not available. Please install required dependencies."],
                "error": "AI improvement engine not available"
            }
        
        # Convert ResumePayload to dict
        resume_data = {
            'name': payload.resume_data.name,
            'title': payload.resume_data.title,
            'email': payload.resume_data.email,
            'phone': payload.resume_data.phone,
            'location': payload.resume_data.location,
            'summary': payload.resume_data.summary,
            'sections': [
                {
                    'id': section.id,
                    'title': section.title,
                    'bullets': [
                        {
                            'id': bullet.id,
                            'text': bullet.text,
                            'params': bullet.params
                        }
                        for bullet in section.bullets
                    ]
                }
                for section in payload.resume_data.sections
            ]
        }
        
        # Get AI improvement suggestions
        result = ai_improvement_engine.get_improvement_suggestions(
            resume_data, 
            payload.job_description, 
            payload.target_role, 
            payload.industry
        )
        
        logger.info(f"AI improvement suggestions generated. Total: {result.get('total_improvements', 0)}")
        
        return result
        
    except Exception as e:
        logger.error(f"AI improvement suggestions error: {str(e)}")
        return {
            "success": False,
            "suggestions": ["Unable to generate improvement suggestions. Please check your content."],
            "error": str(e)
        }

@app.post("/api/ai/improve_ats_score")
async def improve_ats_score_bulk(payload: EnhancedATSPayload):
    """Apply multiple AI improvements to boost ATS score"""
    try:
        logger.info("Processing bulk ATS score improvement request")
        
        # Check if enhanced ATS checker and AI improvement engine are available
        if not enhanced_ats_checker or not ai_improvement_engine:
            return {
                "success": False,
                "improved_resume": None,
                "score_improvement": 0,
                "applied_improvements": [],
                "error": "Required services not available"
            }
        
        # Convert ResumePayload to dict
        resume_data = {
            'name': payload.resume_data.name,
            'title': payload.resume_data.title,
            'email': payload.resume_data.email,
            'phone': payload.resume_data.phone,
            'location': payload.resume_data.location,
            'summary': payload.resume_data.summary,
            'sections': [
                {
                    'id': section.id,
                    'title': section.title,
                    'bullets': [
                        {
                            'id': bullet.id,
                            'text': bullet.text,
                            'params': bullet.params
                        }
                        for bullet in section.bullets
                    ]
                }
                for section in payload.resume_data.sections
            ]
        }
        
        # Get current ATS score
        current_result = enhanced_ats_checker.get_enhanced_ats_score(resume_data, payload.job_description)
        current_score = current_result.get('score', 0)
        
        # Get AI improvements
        improvements = current_result.get('ai_improvements', [])
        
        # Sort improvements by impact score (highest first) and priority
        priority_order = {'high': 3, 'medium': 2, 'low': 1}
        improvements.sort(key=lambda x: (priority_order.get(x.get('priority', 'low'), 1), x.get('impact_score', 0)), reverse=True)
        
        # Apply top improvements (limit to 5 to avoid overwhelming changes)
        applied_improvements = []
        improved_resume = resume_data.copy()
        
        for improvement in improvements[:5]:
            try:
                # Map category to strategy
                strategy_mapping = {
                    'professional_summary': 'summary',
                    'quantified_achievements': 'achievements', 
                    'job_alignment': 'keywords',
                    'career_transition': 'experience',
                    'content_audit': 'content',
                    'modern_format': 'format',
                    'skills_enhancement': 'skills',
                    'leadership_emphasis': 'leadership',
                    'contact_optimization': 'contact',
                    'ats_compatibility': 'ats'
                }
                
                strategy = strategy_mapping.get(improvement.get('category', '').lower().replace(' ', '_'), 'content')
                
                # Generate improvement using AI
                if openai_client:
                    prompt = f"""Apply this specific ATS improvement to the resume:

Improvement: {improvement.get('title', '')}
Description: {improvement.get('description', '')}
Suggestion: {improvement.get('specific_suggestion', '')}

Current Resume Data:
{str(improved_resume)[:2000]}

Job Description Context:
{payload.job_description[:500] if payload.job_description else 'Not provided'}

Requirements:
- Apply the improvement naturally and professionally
- Maintain resume structure and formatting
- Ensure ATS compatibility
- Return the complete updated resume as JSON

Return ONLY the updated resume JSON, no explanations."""
                    
                    response = openai_client['requests'].post(
                        "https://api.openai.com/v1/chat/completions",
                        headers={
                            "Authorization": f"Bearer {openai_client['api_key']}",
                            "Content-Type": "application/json"
                        },
                        json={
                            "model": openai_client['model'],
                            "messages": [
                                {"role": "system", "content": "You are an expert resume writer specializing in ATS optimization. Apply improvements while maintaining professional quality."},
                                {"role": "user", "content": prompt}
                            ],
                            "max_tokens": 2000,
                            "temperature": 0.6
                        }
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        improved_content = result['choices'][0]['message']['content'].strip()
                        
                        # Try to parse the improved resume
                        try:
                            import json
                            updated_resume = json.loads(improved_content)
                            improved_resume = updated_resume
                            applied_improvements.append(improvement)
                            logger.info(f"Applied improvement: {improvement.get('title', '')}")
                        except json.JSONDecodeError:
                            logger.warning(f"Could not parse improved resume for: {improvement.get('title', '')}")
                            continue
                    else:
                        logger.warning(f"OpenAI API error for improvement: {improvement.get('title', '')}")
                        continue
                        
            except Exception as e:
                logger.error(f"Error applying improvement {improvement.get('title', '')}: {str(e)}")
                continue
        
        # Calculate new ATS score
        new_result = enhanced_ats_checker.get_enhanced_ats_score(improved_resume, payload.job_description)
        new_score = new_result.get('score', current_score)
        score_improvement = new_score - current_score
        
        logger.info(f"ATS score improved from {current_score} to {new_score} (+{score_improvement})")
        
        return {
            "success": True,
            "improved_resume": improved_resume,
            "original_score": current_score,
            "new_score": new_score,
            "score_improvement": score_improvement,
            "applied_improvements": applied_improvements,
            "remaining_improvements": len(improvements) - len(applied_improvements)
        }
        
    except Exception as e:
        logger.error(f"Bulk ATS improvement error: {str(e)}")
        return {
            "success": False,
            "improved_resume": None,
            "score_improvement": 0,
            "applied_improvements": [],
            "error": str(e)
        }

@app.post("/api/ai/apply_improvement")
async def apply_ai_improvement(payload: AIImprovementPayload):
    """Apply specific AI improvement to resume content"""
    try:
        logger.info(f"Processing AI improvement application for strategy: {payload.strategy}")
        
        # Check if AI improvement engine is available
        if not ai_improvement_engine:
            return {
                "success": False,
                "improved_content": "",
                "suggestions": ["AI improvement engine is not available."],
                "error": "AI improvement engine not available"
            }
        
        # Convert ResumePayload to dict
        resume_data = {
            'name': payload.resume_data.name,
            'title': payload.resume_data.title,
            'email': payload.resume_data.email,
            'phone': payload.resume_data.phone,
            'location': payload.resume_data.location,
            'summary': payload.resume_data.summary,
            'sections': [
                {
                    'id': section.id,
                    'title': section.title,
                    'bullets': [
                        {
                            'id': bullet.id,
                            'text': bullet.text,
                            'params': bullet.params
                        }
                        for bullet in section.bullets
                    ]
                }
                for section in payload.resume_data.sections
            ]
        }
        
        # Generate improvement prompt for specific strategy
        if payload.strategy:
            from ai_improvement_engine import ImprovementStrategy
            
            # Map category names to strategy enum values
            strategy_mapping = {
                'summary': 'professional_summary',
                'achievements': 'quantified_achievements',
                'keywords': 'job_alignment',
                'experience': 'career_transition',
                'content': 'content_audit',
                'format': 'modern_format',
                'skills': 'skills_enhancement',
                'leadership': 'leadership_emphasis',
                'contact': 'contact_optimization',
                'ats': 'ats_compatibility'
            }
            
            # Get the mapped strategy name
            mapped_strategy = strategy_mapping.get(payload.strategy, payload.strategy)
            logger.info(f"Strategy mapping: {payload.strategy} -> {mapped_strategy}")
            
            try:
                strategy = ImprovementStrategy(mapped_strategy)
                prompt = ai_improvement_engine.generate_improvement_prompt(
                    strategy, resume_data, payload.job_description, 
                    payload.target_role, payload.industry
                )
                
                # Use OpenAI to generate improved content
                if openai_client:
                    try:
                        response = openai_client['requests'].post(
                            "https://api.openai.com/v1/chat/completions",
                            headers={
                                "Authorization": f"Bearer {openai_client['api_key']}",
                                "Content-Type": "application/json"
                            },
                            json={
                                "model": openai_client['model'],
                                "messages": [
                                    {"role": "system", "content": "You are an expert resume writer. Provide specific, actionable improvements to resume content."},
                                    {"role": "user", "content": prompt}
                                ],
                                "max_tokens": OPENAI_MAX_TOKENS,
                                "temperature": 0.7
                            }
                        )
                        
                        if response.status_code == 200:
                            ai_response = response.json()
                            improved_content = ai_response['choices'][0]['message']['content']
                            
                            return {
                                "success": True,
                                "improved_content": improved_content,
                                "strategy": payload.strategy,
                                "prompt_used": prompt,
                                "message": f"AI improvement applied for {payload.strategy} strategy"
                            }
                        else:
                            logger.error(f"OpenAI API error: {response.status_code}")
                            return {
                                "success": False,
                                "improved_content": "",
                                "suggestions": ["AI improvement generation failed. Please try again."],
                                "error": f"OpenAI API error: {response.status_code}"
                            }
                    except Exception as e:
                        logger.error(f"OpenAI request error: {str(e)}")
                        return {
                            "success": False,
                            "improved_content": "",
                            "suggestions": ["AI improvement generation failed. Please try again."],
                            "error": str(e)
                        }
                else:
                    return {
                        "success": False,
                        "improved_content": "",
                        "suggestions": ["OpenAI client not available. Please configure API key."],
                        "error": "OpenAI client not available"
                    }
            except ValueError:
                return {
                    "success": False,
                    "improved_content": "",
                    "suggestions": [f"Invalid strategy: {payload.strategy}. Available strategies: {', '.join(strategy_mapping.keys())}"],
                    "error": f"Invalid strategy: {payload.strategy}. Mapped to: {mapped_strategy}"
                }
        else:
            return {
                "success": False,
                "improved_content": "",
                "suggestions": ["Please specify a strategy to apply."],
                "error": "No strategy specified"
            }
        
    except Exception as e:
        logger.error(f"AI improvement application error: {str(e)}")
        return {
            "success": False,
            "improved_content": "",
            "suggestions": ["Unable to apply AI improvement. Please try again."],
            "error": str(e)
        }

# Pydantic model for work experience generation
class WorkExperienceRequest(BaseModel):
    experienceDescription: str
    currentCompany: Optional[str] = None
    currentJobTitle: Optional[str] = None
    currentDateRange: Optional[str] = None

@app.post("/api/ai/generate-work-experience")
async def generate_work_experience(payload: WorkExperienceRequest):
    """Generate work experience content from user description"""
    try:
        logger.info("Processing work experience generation request")
        
        # Check if OpenAI is available
        if not openai_client:
            return {
                "success": False,
                "error": "OpenAI API not available",
                "companyName": payload.currentCompany or "Tech Company",
                "jobTitle": payload.currentJobTitle or "Software Engineer", 
                "dateRange": payload.currentDateRange or "2020-2023",
                "bullets": [
                    "Developed and maintained web applications using modern technologies",
                    "Collaborated with cross-functional teams to deliver high-quality software solutions",
                    "Implemented automated testing and CI/CD pipelines",
                    "Mentored junior developers and conducted code reviews"
                ]
            }
        
        # Create prompt for work experience generation
        prompt = f"""
        Based on the following work experience description, generate professional resume content:

        Experience Description: {payload.experienceDescription}

        Current Company: {payload.currentCompany or 'Not specified'}
        Current Job Title: {payload.currentJobTitle or 'Not specified'}
        Current Date Range: {payload.currentDateRange or 'Not specified'}

        Please generate:
        1. A professional company name (if not provided or generic)
        2. A professional job title (if not provided or generic)
        3. A realistic date range (if not provided)
        4. 4-6 professional bullet points that highlight achievements, responsibilities, and impact

        Return the response in this JSON format:
        {{
            "companyName": "Company Name",
            "jobTitle": "Job Title", 
            "dateRange": "Date Range",
            "bullets": [
                "Achievement-focused bullet point 1",
                "Responsibility-focused bullet point 2",
                "Impact-focused bullet point 3",
                "Skill-focused bullet point 4"
            ]
        }}

        Make the bullet points:
        - Start with action verbs (Led, Developed, Implemented, etc.)
        - Include quantifiable results where possible
        - Focus on achievements and impact, not just responsibilities
        - Be specific and professional
        - Use industry-standard terminology
        """

        headers = {
            "Authorization": f"Bearer {openai_client['api_key']}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": openai_client['model'],
            "messages": [
                {"role": "system", "content": "You are a professional resume writer and career coach. Generate high-quality, ATS-friendly resume content."},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 800,
            "temperature": 0.7
        }
        
        response = openai_client['requests'].post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code}")
        
        result = response.json()
        generated_content = result['choices'][0]['message']['content'].strip()
        
        # Parse the JSON response
        try:
            import json
            result = json.loads(generated_content)
            
            return {
                "success": True,
                "companyName": result.get("companyName", payload.currentCompany or "Tech Company"),
                "jobTitle": result.get("jobTitle", payload.currentJobTitle or "Software Engineer"),
                "dateRange": result.get("dateRange", payload.currentDateRange or "2020-2023"),
                "bullets": result.get("bullets", [
                    "Developed and maintained web applications using modern technologies",
                    "Collaborated with cross-functional teams to deliver high-quality software solutions",
                    "Implemented automated testing and CI/CD pipelines",
                    "Mentored junior developers and conducted code reviews"
                ])
            }
        except json.JSONDecodeError:
            # Fallback if JSON parsing fails
            return {
                "success": True,
                "companyName": payload.currentCompany or "Tech Company",
                "jobTitle": payload.currentJobTitle or "Software Engineer",
                "dateRange": payload.currentDateRange or "2020-2023",
                "bullets": [
                    "Developed and maintained web applications using modern technologies",
                    "Collaborated with cross-functional teams to deliver high-quality software solutions", 
                    "Implemented automated testing and CI/CD pipelines",
                    "Mentored junior developers and conducted code reviews"
                ]
            }
        
    except Exception as e:
        logger.error(f"Work experience generation error: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "companyName": payload.currentCompany or "Tech Company",
            "jobTitle": payload.currentJobTitle or "Software Engineer",
            "dateRange": payload.currentDateRange or "2020-2023", 
            "bullets": [
                "Developed and maintained web applications using modern technologies",
                "Collaborated with cross-functional teams to deliver high-quality software solutions",
                "Implemented automated testing and CI/CD pipelines",
                "Mentored junior developers and conducted code reviews"
            ]
        }

@app.websocket("/ws/collab/{room_id}")
async def websocket_collab(websocket: WebSocket, room_id: str):
    await websocket.accept()
    
    try:
        init_msg = await websocket.receive_json()
        user_id = init_msg.get("user_id", secrets.token_urlsafe(6))
        user_name = init_msg.get("user_name", "Anonymous")
        
        collab_manager.add_connection(room_id, user_id, websocket, user_name)
        
        active_users = collab_manager.get_active_users(room_id)
        await collab_manager.broadcast(room_id, {
            "type": "user_joined",
            "user_id": user_id,
            "user_name": user_name,
            "active_users": active_users
        })
        
        while True:
            data = await websocket.receive_json()
            message_type = data.get("type")
            
            if message_type == "resume_update":
                await collab_manager.broadcast(room_id, {
                    "type": "resume_update",
                    "user_id": user_id,
                    "user_name": user_name,
                    "data": data.get("data"),
                    "timestamp": datetime.now().isoformat()
                }, exclude_user=user_id)
            
            elif message_type == "cursor_position":
                await collab_manager.broadcast(room_id, {
                    "type": "cursor_position",
                    "user_id": user_id,
                    "user_name": user_name,
                    "position": data.get("position")
                }, exclude_user=user_id)
            
            elif message_type == "ping":
                await websocket.send_json({"type": "pong"})
            
            elif message_type == "add_comment":
                comment = collab_manager.add_comment(room_id, {
                    "user_name": user_name,
                    "text": data.get("text"),
                    "target_type": data.get("target_type"),
                    "target_id": data.get("target_id"),
                    "resolved": False
                })
                await collab_manager.broadcast(room_id, {
                    "type": "comment_added",
                    "comment": comment
                })
            
            elif message_type == "resolve_comment":
                comment_id = data.get("comment_id")
                collab_manager.resolve_comment(room_id, comment_id)
                await collab_manager.broadcast(room_id, {
                    "type": "comment_resolved",
                    "comment_id": comment_id
                })
            
            elif message_type == "delete_comment":
                comment_id = data.get("comment_id")
                collab_manager.delete_comment(room_id, comment_id)
                await collab_manager.broadcast(room_id, {
                    "type": "comment_deleted",
                    "comment_id": comment_id
                })
    
    except WebSocketDisconnect:
        logger.info(f"WebSocket disconnected for user {user_id}")
        collab_manager.remove_connection(room_id, user_id)
        active_users = collab_manager.get_active_users(room_id)
        await collab_manager.broadcast(room_id, {
            "type": "user_left",
            "user_id": user_id,
            "user_name": user_name,
            "active_users": active_users
        })
    except Exception as e:
        logger.error(f"WebSocket error for user {user_id}: {e}")
        collab_manager.remove_connection(room_id, user_id)
        active_users = collab_manager.get_active_users(room_id)
        await collab_manager.broadcast(room_id, {
            "type": "user_left",
            "user_id": user_id,
            "user_name": user_name,
            "active_users": active_users
        })

# Version Control Endpoints

@app.get("/api/resumes")
async def list_user_resumes(user_email: str = Query(..., description="User email for authentication"), db: Session = Depends(get_db)):
    """Get all resumes for a user"""
    try:
        logger.info(f"list_user_resumes: Request received for user {user_email}")
        
        if not user_email:
            raise HTTPException(status_code=400, detail="user_email is required")
        
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            logger.error(f"list_user_resumes: User not found for email {user_email}")
            raise HTTPException(status_code=404, detail="User not found")
        
        logger.info(f"list_user_resumes: User found with id {user.id}")
        
        resumes = db.query(Resume).filter(Resume.user_id == user.id).order_by(Resume.updated_at.desc()).all()
        logger.info(f"list_user_resumes: Found {len(resumes)} resumes for user {user.id}")
        
        result = []
        for resume in resumes:
            # Get latest version info
            latest_version = db.query(ResumeVersion).filter(
                ResumeVersion.resume_id == resume.id
            ).order_by(ResumeVersion.version_number.desc()).first()
            
            # Get match sessions for this resume with JD and version info
            match_sessions = db.query(MatchSession).filter(
                MatchSession.resume_id == resume.id
            ).order_by(MatchSession.created_at.desc()).limit(5).all()
            
            recent_matches_data = []
            for ms in match_sessions:
                # Get job description details
                jd = db.query(JobDescription).filter(JobDescription.id == ms.job_description_id).first()
                
                # Get resume version for this match (if available via relationship)
                # Try to find the version that was used for this match
                # We'll look for versions created around the same time or the latest version before the match
                match_version = None
                if latest_version:
                    # Check if there's a version created before or at the match time
                    version_query = db.query(ResumeVersion).filter(
                        ResumeVersion.resume_id == resume.id,
                        ResumeVersion.created_at <= ms.created_at
                    ).order_by(ResumeVersion.version_number.desc()).first()
                    if version_query:
                        match_version = version_query
                    else:
                        # Fallback to latest version
                        match_version = latest_version
                
                recent_matches_data.append({
                    'id': ms.id,
                    'job_description_id': ms.job_description_id,
                    'jd_title': jd.title if jd else 'Unknown Job',
                    'jd_company': jd.company if jd else None,
                    'score': ms.score,
                    'keyword_coverage': ms.keyword_coverage,
                    'resume_version_id': match_version.id if match_version else None,
                    'resume_version_number': match_version.version_number if match_version else None,
                    'created_at': ms.created_at.isoformat() if ms.created_at else None
                })
            
            result.append({
                'id': resume.id,
                'name': resume.name,
                'title': resume.title,
                'template': resume.template,
                'created_at': resume.created_at.isoformat() if resume.created_at else None,
                'updated_at': resume.updated_at.isoformat() if resume.updated_at else None,
                'latest_version_id': latest_version.id if latest_version else None,
                'latest_version_number': latest_version.version_number if latest_version else None,
                'version_count': db.query(ResumeVersion).filter(ResumeVersion.resume_id == resume.id).count(),
                'match_count': len(match_sessions),
                'recent_matches': recent_matches_data
            })
        
        logger.info(f"list_user_resumes: Returning {len(result)} resumes")
        return {
            "success": True,
            "resumes": result,
            "count": len(result)
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing resumes: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to list resumes")

@app.delete("/api/resumes/{resume_id}")
async def delete_resume(resume_id: int, user_email: str = Query(..., description="User email for authentication"), db: Session = Depends(get_db)):
    """Delete a resume and all its associated data"""
    try:
        logger.info(f"delete_resume: Attempting to delete resume {resume_id} for user {user_email}")
        
        if not user_email:
            raise HTTPException(status_code=400, detail="user_email is required")
        
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            logger.error(f"delete_resume: User not found for email {user_email}")
            raise HTTPException(status_code=404, detail="User not found")
        
        # Find the resume and verify ownership
        resume = db.query(Resume).filter(
            Resume.id == resume_id,
            Resume.user_id == user.id
        ).first()
        
        if not resume:
            logger.warning(f"delete_resume: Resume {resume_id} not found or doesn't belong to user {user_email}")
            raise HTTPException(status_code=404, detail="Resume not found or access denied")
        
        # Delete associated data (cascade should handle most, but we'll be explicit)
        # Delete match sessions
        match_sessions = db.query(MatchSession).filter(MatchSession.resume_id == resume_id).all()
        for ms in match_sessions:
            db.delete(ms)
        logger.info(f"delete_resume: Deleted {len(match_sessions)} match sessions")
        
        # Delete shared resumes
        shared_resumes = db.query(SharedResume).filter(SharedResume.resume_id == resume_id).all()
        for sr in shared_resumes:
            db.delete(sr)
        logger.info(f"delete_resume: Deleted {len(shared_resumes)} shared resume records")
        
        # Delete export analytics
        export_analytics = db.query(ExportAnalytics).filter(ExportAnalytics.resume_id == resume_id).all()
        for ea in export_analytics:
            db.delete(ea)
        logger.info(f"delete_resume: Deleted {len(export_analytics)} export analytics records")
        
        # Delete job matches
        job_matches = db.query(JobMatch).filter(JobMatch.resume_id == resume_id).all()
        for jm in job_matches:
            db.delete(jm)
        logger.info(f"delete_resume: Deleted {len(job_matches)} job matches")
        
        # Resume versions will be deleted by cascade
        # Delete the resume itself
        db.delete(resume)
        db.commit()
        
        logger.info(f"delete_resume: Successfully deleted resume {resume_id}")
        return {
            "success": True,
            "message": "Resume deleted successfully",
            "resume_id": resume_id
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"delete_resume: Error deleting resume {resume_id}: {e}", exc_info=True)
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to delete resume: {str(e)}")

@app.post("/api/resume/save")
async def save_resume(payload: SaveResumePayload, user_email: str = Query(..., description="User email for authentication"), db: Session = Depends(get_db)):
    """Save or update a resume with version control"""
    try:
        if not user_email:
            logger.error("save_resume: user_email is required")
            raise HTTPException(status_code=400, detail="user_email is required")
        
        logger.info(f"save_resume: Attempting to save resume for user {user_email}, resume name: {payload.name}")
        
        # Get user from database
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            logger.error(f"save_resume: User not found for email {user_email}")
            raise HTTPException(status_code=404, detail="User not found")
        
        logger.info(f"save_resume: User found with id {user.id}")
        
        # Validate payload
        if not payload.name or not payload.name.strip():
            logger.error(f"save_resume: Resume name is required")
            raise HTTPException(status_code=400, detail="Resume name is required")
        
        # Check if resume exists
        resume = db.query(Resume).filter(
            Resume.user_id == user.id,
            Resume.name == payload.name
        ).first()
        
        try:
            if resume:
                # Update existing resume
                logger.info(f"save_resume: Updating existing resume id {resume.id}")
                resume.title = payload.title
                resume.email = payload.email
                resume.phone = payload.phone
                resume.location = payload.location
                resume.summary = payload.summary
                resume.template = payload.template or "tech"
                resume.updated_at = datetime.utcnow()
                db.commit()
                db.refresh(resume)
            else:
                # Create new resume
                logger.info(f"save_resume: Creating new resume")
                resume = Resume(
                    user_id=user.id,
                    name=payload.name,
                    title=payload.title or "",
                    email=payload.email or "",
                    phone=payload.phone or "",
                    location=payload.location or "",
                    summary=payload.summary or "",
                    template=payload.template or "tech"
                )
                db.add(resume)
                db.commit()
                db.refresh(resume)
                logger.info(f"save_resume: Created resume with id {resume.id}")
            
            # Create version with safe section processing
            version_service = VersionControlService(db)
            
            # Safely process sections to handle missing attributes
            sections_data = []
            for s in payload.sections or []:
                try:
                    section_id = getattr(s, 'id', str(datetime.utcnow().timestamp()))
                    section_title = getattr(s, 'title', 'Untitled Section')
                    bullets_data = []
                    
                    for b in getattr(s, 'bullets', []):
                        try:
                            bullet_id = getattr(b, 'id', str(datetime.utcnow().timestamp()))
                            bullet_text = getattr(b, 'text', '')
                            bullet_params = getattr(b, 'params', {})
                            if not isinstance(bullet_params, dict):
                                bullet_params = {}
                            bullets_data.append({
                                "id": str(bullet_id),
                                "text": str(bullet_text),
                                "params": bullet_params
                            })
                        except Exception as bullet_error:
                            logger.warning(f"save_resume: Error processing bullet: {bullet_error}")
                            continue
                    
                    sections_data.append({
                        "id": str(section_id),
                        "title": str(section_title),
                        "bullets": bullets_data
                    })
                except Exception as section_error:
                    logger.warning(f"save_resume: Error processing section: {section_error}")
                    continue
            
            resume_data = {
                "personalInfo": {
                    "name": payload.name or "",
                    "email": payload.email or "",
                    "phone": payload.phone or "",
                    "location": payload.location or ""
                },
                "summary": payload.summary or "",
                "sections": sections_data
            }
            
            logger.info(f"save_resume: Creating version with {len(sections_data)} sections")
            version = version_service.create_version(
                user_id=user.id,
                resume_id=resume.id,
                resume_data=resume_data,
                change_summary="Manual save",
                is_auto_save=False
            )
            
            logger.info(f"save_resume: Successfully saved resume id {resume.id}, version id {version.id}")
            
            return {
                "success": True,
                "resume_id": resume.id,
                "version_id": version.id,
                "message": "Resume saved successfully"
            }
            
        except Exception as db_error:
            db.rollback()
            logger.error(f"save_resume: Database error: {str(db_error)}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"Database error: {str(db_error)}")
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"save_resume: Unexpected error saving resume for {user_email}: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to save resume: {str(e)}")

@app.post("/api/resume/version/create")
async def create_version(payload: CreateVersionPayload, user_email: str = Query(..., description="User email for authentication"), db: Session = Depends(get_db)):
    """Create a new version of a resume"""
    try:
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Verify resume belongs to user
        resume = db.query(Resume).filter(
            Resume.id == payload.resume_id,
            Resume.user_id == user.id
        ).first()
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        version_service = VersionControlService(db)
        version = version_service.create_version(
            user_id=user.id,
            resume_id=payload.resume_id,
            resume_data=payload.resume_data,
            change_summary=payload.change_summary,
            is_auto_save=payload.is_auto_save
        )
        
        return {
            "success": True,
            "version_id": version.id,
            "version_number": version.version_number,
            "message": "Version created successfully"
        }
        
    except Exception as e:
        logger.error(f"Error creating version: {e}")
        raise HTTPException(status_code=500, detail="Failed to create version")

@app.get("/api/resume/{resume_id}/versions")
async def get_resume_versions(resume_id: int, user_email: str, db: Session = Depends(get_db)):
    """Get all versions for a resume"""
    try:
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Verify resume belongs to user
        resume = db.query(Resume).filter(
            Resume.id == resume_id,
            Resume.user_id == user.id
        ).first()
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        version_service = VersionControlService(db)
        versions = version_service.get_resume_versions(resume_id, user.id)
        
        return {
            "success": True,
            "versions": [
                {
                    "id": v.id,
                    "version_number": v.version_number,
                    "change_summary": v.change_summary,
                    "is_auto_save": v.is_auto_save,
                    "created_at": v.created_at.isoformat()
                }
                for v in versions
            ]
        }
        
    except Exception as e:
        logger.error(f"Error getting versions: {e}")
        raise HTTPException(status_code=500, detail="Failed to get versions")

@app.post("/api/resume/version/rollback")
async def rollback_version(payload: RollbackVersionPayload, user_email: str, db: Session = Depends(get_db)):
    """Rollback to a specific version"""
    try:
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        version_service = VersionControlService(db)
        new_version = version_service.rollback_to_version(payload.version_id, user.id)
        
        if not new_version:
            raise HTTPException(status_code=404, detail="Version not found")
        
        return {
            "success": True,
            "new_version_id": new_version.id,
            "version_number": new_version.version_number,
            "message": "Rollback successful"
        }
        
    except Exception as e:
        logger.error(f"Error rolling back version: {e}")
        raise HTTPException(status_code=500, detail="Failed to rollback version")

@app.get("/api/resume/version/{version_id}")
async def get_version(version_id: int, user_email: str, db: Session = Depends(get_db)):
    """Get a specific version data"""
    try:
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        version_service = VersionControlService(db)
        version = version_service.get_version(version_id, user.id)
        
        if not version:
            raise HTTPException(status_code=404, detail="Version not found")
        
        return {
            "success": True,
            "version": {
                "id": version.id,
                "version_number": version.version_number,
                "resume_data": version.resume_data,
                "change_summary": version.change_summary,
                "is_auto_save": version.is_auto_save,
                "created_at": version.created_at.isoformat()
            }
        }
        
    except Exception as e:
        logger.error(f"Error getting version: {e}")
        raise HTTPException(status_code=500, detail="Failed to get version")

@app.post("/api/resume/version/compare")
async def compare_versions(payload: CompareVersionsPayload, user_email: str, db: Session = Depends(get_db)):
    """Compare two versions"""
    try:
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        version_service = VersionControlService(db)
        comparison = version_service.compare_versions(
            payload.version1_id, 
            payload.version2_id, 
            user.id
        )
        
        if not comparison:
            raise HTTPException(status_code=404, detail="Versions not found or incompatible")
        
        return {
            "success": True,
            "comparison": comparison
        }
        
    except Exception as e:
        logger.error(f"Error comparing versions: {e}")
        raise HTTPException(status_code=500, detail="Failed to compare versions")

@app.delete("/api/resume/version/{version_id}")
async def delete_version(version_id: int, user_email: str, db: Session = Depends(get_db)):
    """Delete a specific version"""
    try:
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        version_service = VersionControlService(db)
        success = version_service.delete_version(version_id, user.id)
        
        if not success:
            raise HTTPException(status_code=400, detail="Cannot delete version (only version or not found)")
        
        return {
            "success": True,
            "message": "Version deleted successfully"
        }
        
    except Exception as e:
        logger.error(f"Error deleting version: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete version")

# Export Analytics Endpoints

@app.get("/api/analytics/exports")
async def get_export_analytics(user_email: str, db: Session = Depends(get_db)):
    """Get export analytics for a user"""
    try:
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Get export analytics
        exports = db.query(ExportAnalytics).filter(
            ExportAnalytics.user_id == user.id
        ).order_by(ExportAnalytics.created_at.desc()).all()
        
        # Get summary statistics
        total_exports = len(exports)
        pdf_exports = len([e for e in exports if e.export_format == 'pdf'])
        docx_exports = len([e for e in exports if e.export_format == 'docx'])
        
        # Get template usage
        template_usage = {}
        for export in exports:
            template = export.template_used or 'unknown'
            template_usage[template] = template_usage.get(template, 0) + 1
        
        # Get recent exports (last 30 days)
        from datetime import datetime, timedelta
        thirty_days_ago = datetime.utcnow() - timedelta(days=30)
        recent_exports = [e for e in exports if e.created_at >= thirty_days_ago]
        
        return {
            "success": True,
            "analytics": {
                "total_exports": total_exports,
                "pdf_exports": pdf_exports,
                "docx_exports": docx_exports,
                "recent_exports": len(recent_exports),
                "template_usage": template_usage,
                "exports": [
                    {
                        "id": e.id,
                        "format": e.export_format,
                        "template": e.template_used,
                        "file_size": e.file_size,
                        "success": e.export_success,
                        "created_at": e.created_at.isoformat(),
                        "resume_name": e.resume.name if e.resume else "Unknown"
                    }
                    for e in exports
                ]
            }
        }
        
    except Exception as e:
        logger.error(f"Error getting export analytics: {e}")
        raise HTTPException(status_code=500, detail="Failed to get export analytics")

# Shared Resume Endpoints

@app.post("/api/resume/share")
async def create_shared_resume(resume_id: int, user_email: str, password: str = None, expires_days: int = None, db: Session = Depends(get_db)):
    """Create a shareable link for a resume"""
    try:
        logger.info(f"Creating shared resume for resume_id={resume_id}, user_email={user_email}")
        
        # Try case-insensitive lookup
        user = db.query(User).filter(User.email.ilike(user_email)).first()
        if not user:
            logger.error(f"User not found for email: {user_email}")
            raise HTTPException(status_code=404, detail="User not found")
        
        logger.info(f"Found user: {user.id}")
        
        # Verify resume belongs to user
        resume = db.query(Resume).filter(
            Resume.id == resume_id,
            Resume.user_id == user.id
        ).first()
        if not resume:
            logger.error(f"Resume not found for resume_id={resume_id}, user_id={user.id}")
            raise HTTPException(status_code=404, detail="Resume not found")
        
        logger.info(f"Found resume: {resume.id}")
        
        # Generate unique share token
        share_token = secrets.token_urlsafe(32)
        
        # Calculate expiration date
        expires_at = None
        if expires_days:
            from datetime import datetime, timedelta
            expires_at = datetime.utcnow() + timedelta(days=expires_days)
        
        # Create shared resume record
        logger.info(f"Creating SharedResume record with resume_id={resume_id}, user_id={user.id}")
        shared_resume = SharedResume(
            resume_id=resume_id,
            user_id=user.id,
            share_token=share_token,
            password_protected=bool(password),
            password_hash=password,  # In production, hash this password
            expires_at=expires_at
        )
        
        logger.info(f"Adding shared resume to database")
        db.add(shared_resume)
        logger.info(f"Committing to database")
        db.commit()
        logger.info(f"Refreshing shared resume")
        db.refresh(shared_resume)
        logger.info(f"Shared resume created with ID: {shared_resume.id}")
        
        # Generate shareable URL
        base_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        if not base_url or base_url == "http://localhost:3000":
            logger.warning("FRONTEND_URL not set or using localhost - this will break in production!")
        share_url = f"{base_url}/shared/{share_token}"
        
        return {
            "success": True,
            "share_token": share_token,
            "share_url": share_url,
            "expires_at": expires_at.isoformat() if expires_at else None,
            "password_protected": bool(password)
        }
        
    except HTTPException:
        # Bubble up known errors like 404 User/Resume not found
        raise
    except Exception as e:
        logger.error(f"Error creating shared resume: {e}")
        raise HTTPException(status_code=500, detail="Failed to create shared resume")

@app.get("/api/resume/shared/{share_token}")
async def get_shared_resume(share_token: str, password: str = None, db: Session = Depends(get_db)):
    """Get a shared resume by token"""
    try:
        shared_resume = db.query(SharedResume).filter(
            SharedResume.share_token == share_token,
            SharedResume.is_active == True
        ).first()
        
        if not shared_resume:
            raise HTTPException(status_code=404, detail="Shared resume not found")
        
        # Check if expired
        if shared_resume.expires_at and shared_resume.expires_at < datetime.utcnow():
            raise HTTPException(status_code=410, detail="Shared resume has expired")
        
        # Check password if required
        if shared_resume.password_protected:
            if not password or password != shared_resume.password_hash:
                raise HTTPException(status_code=401, detail="Password required")
        
        # Get resume data
        resume = shared_resume.resume
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        # Get latest version
        version_service = VersionControlService(db)
        latest_version = version_service.get_latest_version(resume.id, resume.user_id)
        
        resume_data = latest_version.resume_data if latest_version else {
            "personalInfo": {
                "name": resume.name,
                "email": resume.email,
                "phone": resume.phone,
                "location": resume.location
            },
            "summary": resume.summary,
            "sections": []
        }
        
        return {
            "success": True,
            "resume": {
                "id": resume.id,
                "name": resume.name,
                "title": resume.title,
                "template": resume.template,
                "created_at": resume.created_at.isoformat(),
                "updated_at": resume.updated_at.isoformat()
            },
            "resume_data": resume_data,
            "shared_info": {
                "created_at": shared_resume.created_at.isoformat(),
                "expires_at": shared_resume.expires_at.isoformat() if shared_resume.expires_at else None
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting shared resume: {e}")
        raise HTTPException(status_code=500, detail="Failed to get shared resume")

@app.post("/api/resume/shared/{share_token}/view")
async def track_resume_view(share_token: str, request: Request, db: Session = Depends(get_db)):
    """Track a view of a shared resume"""
    try:
        shared_resume = db.query(SharedResume).filter(
            SharedResume.share_token == share_token,
            SharedResume.is_active == True
        ).first()
        
        if not shared_resume:
            raise HTTPException(status_code=404, detail="Shared resume not found")
        
        # Get client info
        client_ip = request.client.host
        user_agent = request.headers.get("user-agent", "")
        referrer = request.headers.get("referer", "")
        
        # Create view record
        view = ResumeView(
            shared_resume_id=shared_resume.id,
            viewer_ip=client_ip,
            viewer_user_agent=user_agent,
            referrer=referrer
        )
        
        db.add(view)
        db.commit()
        
        return {"success": True, "message": "View tracked"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error tracking view: {e}")
        raise HTTPException(status_code=500, detail="Failed to track view")

@app.get("/api/resume/shared/{share_token}/analytics")
async def get_shared_resume_analytics(share_token: str, user_email: str, db: Session = Depends(get_db)):
    """Get analytics for a shared resume"""
    try:
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        shared_resume = db.query(SharedResume).filter(
            SharedResume.share_token == share_token,
            SharedResume.user_id == user.id
        ).first()
        
        if not shared_resume:
            raise HTTPException(status_code=404, detail="Shared resume not found")
        
        # Get views
        views = db.query(ResumeView).filter(
            ResumeView.shared_resume_id == shared_resume.id
        ).order_by(ResumeView.created_at.desc()).all()
        
        # Calculate analytics
        total_views = len(views)
        unique_ips = len(set(v.viewer_ip for v in views))
        
        # Group by date
        from collections import defaultdict
        views_by_date = defaultdict(int)
        for view in views:
            date_key = view.created_at.date().isoformat()
            views_by_date[date_key] += 1
        
        return {
            "success": True,
            "analytics": {
                "total_views": total_views,
                "unique_visitors": unique_ips,
                "views_by_date": dict(views_by_date),
                "recent_views": [
                    {
                        "ip": v.viewer_ip,
                        "user_agent": v.viewer_user_agent,
                        "referrer": v.referrer,
                        "created_at": v.created_at.isoformat()
                    }
                    for v in views[:10]  # Last 10 views
                ]
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting shared resume analytics: {e}")
        raise HTTPException(status_code=500, detail="Failed to get analytics")

@app.delete("/api/resume/shared/{share_token}")
async def deactivate_shared_resume(share_token: str, user_email: str, db: Session = Depends(get_db)):
    """Deactivate a shared resume link"""
    try:
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        shared_resume = db.query(SharedResume).filter(
            SharedResume.share_token == share_token,
            SharedResume.user_id == user.id
        ).first()
        
        if not shared_resume:
            raise HTTPException(status_code=404, detail="Shared resume not found")
        
        shared_resume.is_active = False
        db.commit()
        
        return {"success": True, "message": "Shared resume deactivated"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deactivating shared resume: {e}")
        raise HTTPException(status_code=500, detail="Failed to deactivate shared resume")

# Shared Resume Comments Endpoints

class SharedResumeCommentPayload(BaseModel):
    commenter_name: str
    commenter_email: Optional[str] = None
    text: str
    target_type: str  # 'resume', 'section', 'bullet'
    target_id: str    # ID of the specific element

@app.get("/api/resume/shared/{share_token}/comments")
async def get_shared_resume_comments(share_token: str, db: Session = Depends(get_db)):
    """Get comments for a shared resume"""
    try:
        shared_resume = db.query(SharedResume).filter(
            SharedResume.share_token == share_token,
            SharedResume.is_active == True
        ).first()
        
        if not shared_resume:
            raise HTTPException(status_code=404, detail="Shared resume not found")
        
        comments = db.query(SharedResumeComment).filter(
            SharedResumeComment.shared_resume_id == shared_resume.id
        ).order_by(SharedResumeComment.created_at.desc()).all()
        
        return {
            "success": True,
            "comments": [
                {
                    "id": comment.id,
                    "commenter_name": comment.commenter_name,
                    "commenter_email": comment.commenter_email,
                    "text": comment.text,
                    "target_type": comment.target_type,
                    "target_id": comment.target_id,
                    "resolved": comment.resolved,
                    "created_at": comment.created_at.isoformat()
                }
                for comment in comments
            ]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting shared resume comments: {e}")
        raise HTTPException(status_code=500, detail="Failed to get comments")

@app.post("/api/resume/shared/{share_token}/comments")
async def add_shared_resume_comment(share_token: str, payload: SharedResumeCommentPayload, db: Session = Depends(get_db)):
    """Add a comment to a shared resume"""
    try:
        shared_resume = db.query(SharedResume).filter(
            SharedResume.share_token == share_token,
            SharedResume.is_active == True
        ).first()
        
        if not shared_resume:
            raise HTTPException(status_code=404, detail="Shared resume not found")
        
        comment = SharedResumeComment(
            shared_resume_id=shared_resume.id,
            commenter_name=payload.commenter_name,
            commenter_email=payload.commenter_email,
            text=payload.text,
            target_type=payload.target_type,
            target_id=payload.target_id
        )
        
        db.add(comment)
        db.commit()
        db.refresh(comment)
        
        logger.info(f"Added comment to shared resume {share_token}: {comment.id}")
        
        return {
            "success": True,
            "comment": {
                "id": comment.id,
                "commenter_name": comment.commenter_name,
                "commenter_email": comment.commenter_email,
                "text": comment.text,
                "target_type": comment.target_type,
                "target_id": comment.target_id,
                "resolved": comment.resolved,
                "created_at": comment.created_at.isoformat()
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error adding shared resume comment: {e}")
        raise HTTPException(status_code=500, detail="Failed to add comment")

@app.post("/api/resume/shared/{share_token}/comments/{comment_id}/resolve")
async def resolve_shared_resume_comment(share_token: str, comment_id: int, db: Session = Depends(get_db)):
    """Resolve a comment on a shared resume"""
    try:
        shared_resume = db.query(SharedResume).filter(
            SharedResume.share_token == share_token,
            SharedResume.is_active == True
        ).first()
        
        if not shared_resume:
            raise HTTPException(status_code=404, detail="Shared resume not found")
        
        comment = db.query(SharedResumeComment).filter(
            SharedResumeComment.id == comment_id,
            SharedResumeComment.shared_resume_id == shared_resume.id
        ).first()
        
        if not comment:
            raise HTTPException(status_code=404, detail="Comment not found")
        
        comment.resolved = True
        db.commit()
        
        logger.info(f"Resolved comment {comment_id} on shared resume {share_token}")
        
        return {"success": True, "message": "Comment resolved"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error resolving shared resume comment: {e}")
        raise HTTPException(status_code=500, detail="Failed to resolve comment")

@app.delete("/api/resume/shared/{share_token}/comments/{comment_id}")
async def delete_shared_resume_comment(share_token: str, comment_id: int, db: Session = Depends(get_db)):
    """Delete a comment from a shared resume"""
    try:
        shared_resume = db.query(SharedResume).filter(
            SharedResume.share_token == share_token,
            SharedResume.is_active == True
        ).first()
        
        if not shared_resume:
            raise HTTPException(status_code=404, detail="Shared resume not found")
        
        comment = db.query(SharedResumeComment).filter(
            SharedResumeComment.id == comment_id,
            SharedResumeComment.shared_resume_id == shared_resume.id
        ).first()
        
        if not comment:
            raise HTTPException(status_code=404, detail="Comment not found")
        
        db.delete(comment)
        db.commit()
        
        logger.info(f"Deleted comment {comment_id} from shared resume {share_token}")
        
        return {"success": True, "message": "Comment deleted"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting shared resume comment: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete comment")

# Job Match Analytics Endpoints

@app.get("/api/analytics/job-matches")
async def get_job_match_analytics(user_email: str, db: Session = Depends(get_db)):
    """Get job match analytics for a user"""
    try:
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Get job matches
        job_matches = db.query(JobMatch).filter(
            JobMatch.user_id == user.id
        ).order_by(JobMatch.created_at.desc()).all()
        
        # Calculate analytics
        total_matches = len(job_matches)
        if total_matches == 0:
            return {
                "success": True,
                "analytics": {
                    "total_matches": 0,
                    "average_score": 0,
                    "score_trend": [],
                    "top_missing_keywords": [],
                    "improvement_areas": [],
                    "matches": []
                }
            }
        
        # Calculate average score
        average_score = sum(match.match_score for match in job_matches) / total_matches
        
        # Calculate score trend (last 10 matches)
        recent_matches = job_matches[:10]
        score_trend = [
            {
                "date": match.created_at.date().isoformat(),
                "score": match.match_score,
                "resume_name": match.resume.name if match.resume else "Unknown"
            }
            for match in reversed(recent_matches)
        ]
        
        # Get top missing keywords
        all_missing_keywords = []
        for match in job_matches:
            if match.missing_keywords:
                all_missing_keywords.extend(match.missing_keywords)
        
        from collections import Counter
        keyword_counts = Counter(all_missing_keywords)
        top_missing_keywords = [
            {"keyword": keyword, "count": count}
            for keyword, count in keyword_counts.most_common(10)
        ]
        
        # Get improvement areas
        improvement_areas = []
        for match in job_matches:
            if match.improvement_suggestions:
                for suggestion in match.improvement_suggestions:
                    if isinstance(suggestion, dict) and 'category' in suggestion:
                        improvement_areas.append(suggestion['category'])
        
        improvement_counts = Counter(improvement_areas)
        top_improvement_areas = [
            {"area": area, "count": count}
            for area, count in improvement_counts.most_common(5)
        ]
        
        # Get recent matches with details
        recent_matches_details = [
            {
                "id": match.id,
                "resume_name": match.resume.name if match.resume else "Unknown",
                "match_score": match.match_score,
                "keyword_matches": match.keyword_matches or [],
                "missing_keywords": match.missing_keywords or [],
                "created_at": match.created_at.isoformat(),
                "job_description_preview": match.job_description[:200] + "..." if len(match.job_description) > 200 else match.job_description
            }
            for match in job_matches[:20]  # Last 20 matches
        ]
        
        return {
            "success": True,
            "analytics": {
                "total_matches": total_matches,
                "average_score": round(average_score, 1),
                "score_trend": score_trend,
                "top_missing_keywords": top_missing_keywords,
                "improvement_areas": top_improvement_areas,
                "matches": recent_matches_details
            }
        }
        
    except Exception as e:
        logger.error(f"Error getting job match analytics: {e}")
        raise HTTPException(status_code=500, detail="Failed to get job match analytics")

@app.get("/api/analytics/job-matches/{match_id}")
async def get_job_match_details(match_id: int, user_email: str, db: Session = Depends(get_db)):
    """Get detailed information about a specific job match"""
    try:
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        job_match = db.query(JobMatch).filter(
            JobMatch.id == match_id,
            JobMatch.user_id == user.id
        ).first()
        
        if not job_match:
            raise HTTPException(status_code=404, detail="Job match not found")
        
        return {
            "success": True,
            "match": {
                "id": job_match.id,
                "resume_name": job_match.resume.name if job_match.resume else "Unknown",
                "match_score": job_match.match_score,
                "keyword_matches": job_match.keyword_matches or [],
                "missing_keywords": job_match.missing_keywords or [],
                "improvement_suggestions": job_match.improvement_suggestions or [],
                "job_description": job_match.job_description,
                "created_at": job_match.created_at.isoformat()
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting job match details: {e}")
        raise HTTPException(status_code=500, detail="Failed to get job match details")

# ==========================
# Job Descriptions & Matches
# ==========================
class JobDescriptionCreate(BaseModel):
    id: Optional[int] = None
    title: str
    company: Optional[str] = None
    source: Optional[str] = None
    url: Optional[str] = None
    easy_apply_url: Optional[str] = None
    location: Optional[str] = None  # Company headquarters location
    work_type: Optional[str] = None  # Remote, Hybrid, Onsite
    job_type: Optional[str] = None  # Full Time, Contractor, Part-time, Internship
    content: str
    user_email: Optional[str] = None

class MatchCreate(BaseModel):
    resumeId: Optional[int] = None
    jobDescriptionId: int
    user_email: Optional[str] = None
    resume_name: Optional[str] = None
    resume_title: Optional[str] = None
    resume_snapshot: Optional[Dict[str, Any]] = None
    resume_version_id: Optional[int] = None
    ats_score: Optional[int] = None  # ATS score from match analysis
    jd_metadata: Optional[Dict[str, Any]] = None  # JD metadata (easy_apply_url, work_type, job_type, company)

def _classify_priority_keywords(extracted: Dict[str, Any]) -> Dict[str, List[str]]:
    keyword_freq: Dict[str, int] = extracted.get('keyword_frequency', {}) if isinstance(extracted, dict) else {}
    technical = set(extracted.get('technical_keywords', [])) if isinstance(extracted, dict) else set()
    high_priority: List[str] = []
    regular: List[str] = []
    for kw, cnt in keyword_freq.items():
        if cnt >= 3 or kw in technical:
            high_priority.append(kw)
        else:
            regular.append(kw)
    return { 'high_priority': list(sorted(set(high_priority))), 'regular': list(sorted(set(regular))) }

def _resume_to_text(resume_data: Dict[str, Any]) -> str:
    parts: List[str] = []
    for key in ['name','title','summary','email','phone','location']:
        val = resume_data.get(key)
        if isinstance(val, str):
            parts.append(val)
    for section in resume_data.get('sections', []) or []:
        parts.append(section.get('title',''))
        for b in section.get('bullets', []) or []:
            txt = b.get('text') if isinstance(b, dict) else None
            if txt:
                parts.append(txt)
    return '\n'.join([p for p in parts if p])

def _compute_match_breakdown(jd_text: str, resume_text: str, extracted_jd: Dict[str, Any]) -> Dict[str, Any]:
    similarity = keyword_extractor.calculate_similarity(jd_text, resume_text)
    hp = set(_classify_priority_keywords(extracted_jd)['high_priority'])
    matched = set(similarity['matching_keywords'])
    missing = set(similarity['missing_keywords'])

    total = len(hp) * 2 + max(1, similarity.get('total_job_keywords', 0) - len(hp))
    score = 0
    for kw in matched:
        score += 2 if kw in hp else 1
    critical_misses = len([k for k in missing if k in hp])
    score -= critical_misses
    score_pct = max(0, min(100, int(round((score / max(1, total)) * 100))))

    return {
        'score': score_pct,
        'matched_keywords': sorted(list(matched)),
        'missing_keywords': sorted(list(missing)),
        'priority_keywords': sorted(list(hp)),
        'keyword_coverage': round((len(matched) / max(1, similarity.get('total_job_keywords', 0))) * 100, 2),
        'similarity': similarity
    }

@app.post('/job-descriptions')
def create_or_update_job_description(payload: JobDescriptionCreate, db: Session = Depends(get_db)):
    try:
        user = None
        if payload.user_email:
            user = db.query(User).filter(User.email == payload.user_email).first()
        jd = None
        if payload.id:
            jd = db.query(JobDescription).filter(JobDescription.id == payload.id).first()
            if not jd:
                raise HTTPException(status_code=404, detail='Job description not found')
        else:
            jd = JobDescription()

        jd.user_id = user.id if user else None
        jd.title = payload.title
        jd.company = payload.company
        jd.source = payload.source
        jd.url = payload.url
        jd.easy_apply_url = payload.easy_apply_url
        jd.location = payload.location
        jd.work_type = payload.work_type
        jd.job_type = payload.job_type
        jd.content = payload.content

        extracted = keyword_extractor.extract_keywords(payload.content)
        pri = _classify_priority_keywords(extracted)
        jd.extracted_keywords = extracted
        jd.priority_keywords = pri.get('high_priority', [])
        jd.soft_skills = extracted.get('soft_skills', [])
        jd.high_frequency_keywords = extracted.get('high_frequency_keywords', [])
        jd.ats_insights = extracted.get('ats_keywords', {})

        db.add(jd)
        db.commit()
        db.refresh(jd)
        return {'id': jd.id, 'message': 'saved', 'extracted': extracted, 'priority_keywords': jd.priority_keywords}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception('Failed to save job description')
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.get('/job-descriptions')
def list_job_descriptions(user_email: Optional[str] = None, db: Session = Depends(get_db)):
    try:
        q = db.query(JobDescription)
        if user_email:
            user = db.query(User).filter(User.email == user_email).first()
            if user:
                q = q.filter(JobDescription.user_id == user.id)
            else:
                # User not found, but still return jobs that might not have user_id set
                # This handles cases where jobs were saved without user_email
                q = q.filter((JobDescription.user_id == None) | (JobDescription.user_id == 0))
        else:
            # No user_email provided, return all jobs (for admin or when user not logged in)
            pass
        items = q.order_by(JobDescription.created_at.desc()).limit(100).all()
        result = []
        for it in items:
            try:
                # Get latest match session for this JD
                latest_match = None
                try:
                    latest_match = db.query(MatchSession).filter(
                        MatchSession.job_description_id == it.id
                    ).order_by(MatchSession.created_at.desc()).first()
                except Exception as e:
                    logger.warning(f'Failed to get match session for JD {it.id}: {e}')
                    latest_match = None
                
                # Get resume info if available
                resume_name = None
                if latest_match and latest_match.resume_id:
                    try:
                        resume = db.query(Resume).filter(Resume.id == latest_match.resume_id).first()
                        if resume:
                            resume_name = resume.name
                    except Exception as e:
                        logger.warning(f'Failed to get resume for match {latest_match.id}: {e}')
                
                # Handle priority_keywords - ensure it's serializable
                priority_kw = it.priority_keywords
                if priority_kw is not None:
                    if isinstance(priority_kw, str):
                        try:
                            priority_kw = json.loads(priority_kw)
                        except:
                            priority_kw = []
                    elif not isinstance(priority_kw, list):
                        priority_kw = []
                
                item_data = {
                    'id': it.id,
                    'title': it.title or '',
                    'company': it.company or '',
                    'source': it.source or '',
                    'url': it.url or '',
                    'easy_apply_url': it.easy_apply_url or '',
                    'location': it.location or '',
                    'work_type': it.work_type or '',
                    'job_type': it.job_type or '',
                    'created_at': it.created_at.isoformat() if it.created_at else None,
                    'priority_keywords': priority_kw or [],
                    'soft_skills': it.soft_skills or [],
                    'high_frequency_keywords': it.high_frequency_keywords or [],
                    'ats_insights': it.ats_insights or {},
                }
                
                if latest_match:
                    try:
                        # Get resume version ID if available
                        resume_version_id = None
                        if latest_match.resume_id:
                            try:
                                latest_version = db.query(ResumeVersion).filter(
                                    ResumeVersion.resume_id == latest_match.resume_id
                                ).order_by(ResumeVersion.version_number.desc()).first()
                                if latest_version:
                                    resume_version_id = latest_version.id
                            except Exception as e:
                                logger.warning(f'Failed to get resume version for resume {latest_match.resume_id}: {e}')
                        
                        item_data['last_match'] = {
                            'id': latest_match.id,
                            'score': latest_match.score or 0,
                            'resume_id': latest_match.resume_id,
                            'resume_name': resume_name,
                            'resume_version_id': resume_version_id,
                            'keyword_coverage': latest_match.keyword_coverage,
                            'matched_keywords': latest_match.matched_keywords or [],
                            'missing_keywords': latest_match.missing_keywords or [],
                            'excess_keywords': latest_match.excess_keywords or [],
                            'created_at': latest_match.created_at.isoformat() if latest_match.created_at else None,
                        }
                    except Exception as e:
                        logger.warning(f'Failed to build last_match for JD {it.id}: {e}')
                        item_data['last_match'] = None
                else:
                    item_data['last_match'] = None
                
                result.append(item_data)
            except Exception as e:
                logger.error(f'Error processing JD {it.id}: {e}')
                # Still include the JD with minimal data
                result.append({
                    'id': it.id,
                    'title': it.title or '',
                    'company': it.company or '',
                    'source': it.source or '',
                    'url': it.url or '',
                    'easy_apply_url': it.easy_apply_url or '',
                    'location': it.location or '',
                    'work_type': it.work_type or '',
                    'job_type': it.job_type or '',
                    'created_at': it.created_at.isoformat() if it.created_at else None,
                    'priority_keywords': [],
                    'soft_skills': [],
                    'high_frequency_keywords': [],
                    'ats_insights': {},
                    'last_match': None,
                })
        
        return result
    except Exception as e:
        logger.exception('Failed to list job descriptions')
        raise HTTPException(status_code=500, detail=str(e))

@app.get('/job-descriptions/{jd_id}')
def get_job_description(jd_id: int, db: Session = Depends(get_db)):
    jd = db.query(JobDescription).filter(JobDescription.id == jd_id).first()
    if not jd:
        raise HTTPException(status_code=404, detail='Job description not found')
    
    # Handle JSON fields safely
    extracted_kw = jd.extracted_keywords
    if isinstance(extracted_kw, str):
        try:
            extracted_kw = json.loads(extracted_kw)
        except:
            extracted_kw = {}
    elif extracted_kw is None:
        extracted_kw = {}
    
    priority_kw = jd.priority_keywords
    if isinstance(priority_kw, str):
        try:
            priority_kw = json.loads(priority_kw)
        except:
            priority_kw = []
    elif priority_kw is None:
        priority_kw = []
    
    return {
        'id': jd.id,
        'title': jd.title or '',
        'company': jd.company or '',
        'source': jd.source or '',
        'url': jd.url or '',
        'easy_apply_url': jd.easy_apply_url or '',
        'location': jd.location or '',
        'work_type': jd.work_type or '',
        'job_type': jd.job_type or '',
        'content': jd.content or '',
        'extracted_keywords': extracted_kw,
        'priority_keywords': priority_kw,
        'soft_skills': jd.soft_skills or [],
        'high_frequency_keywords': jd.high_frequency_keywords or [],
        'ats_insights': jd.ats_insights or {},
        'created_at': jd.created_at.isoformat() if jd.created_at else None,
    }

@app.post('/matches')
def create_match(payload: MatchCreate, db: Session = Depends(get_db)):
    """Create a match session between resume and job description"""
    try:
        logger.info(f"create_match: Starting match creation for JD {payload.jobDescriptionId}, resumeId: {payload.resumeId}")
        
        # Find or create resume if needed
        resume = None
        if payload.resumeId:
            resume = db.query(Resume).filter(Resume.id == payload.resumeId).first()
            if not resume:
                logger.warning(f"create_match: Resume {payload.resumeId} not found, will try to create if user_email provided")
        
        # If resume doesn't exist but we have user_email and resume data, create it
        if not resume and payload.user_email and payload.resume_name:
            try:
                user = db.query(User).filter(User.email == payload.user_email).first()
                if user:
                    logger.info(f"create_match: Creating new resume for user {user.email}")
                    # Check if resume with this name already exists
                    existing_resume = db.query(Resume).filter(
                        Resume.user_id == user.id,
                        Resume.name == payload.resume_name
                    ).first()
                    
                    if existing_resume:
                        resume = existing_resume
                        logger.info(f"create_match: Found existing resume {resume.id} with same name")
                    else:
                        # Create new resume from snapshot or basic info
                        resume = Resume(
                            user_id=user.id,
                            name=payload.resume_name,
                            title=payload.resume_title or '',
                            email='',
                            phone='',
                            location='',
                            summary='',
                            template='tech'
                        )
                        db.add(resume)
                        db.commit()
                        db.refresh(resume)
                        logger.info(f"create_match: Created new resume {resume.id}")
                        
                        # Create version if snapshot provided
                        if payload.resume_snapshot:
                            try:
                                version_service = VersionControlService(db)
                                resume_data = {
                                    "personalInfo": payload.resume_snapshot.get('personalInfo', {
                                        "name": payload.resume_name,
                                        "title": payload.resume_title or '',
                                        "email": "",
                                        "phone": "",
                                        "location": ""
                                    }),
                                    "summary": payload.resume_snapshot.get('summary', ''),
                                    "sections": payload.resume_snapshot.get('sections', [])
                                }
                                version = version_service.create_version(
                                    user_id=user.id,
                                    resume_id=resume.id,
                                    resume_data=resume_data,
                                    change_summary="Auto-created from match session",
                                    is_auto_save=False
                                )
                                logger.info(f"create_match: Created version {version.id} for resume {resume.id}")
                            except Exception as e:
                                logger.warning(f"create_match: Failed to create version: {e}")
                else:
                    logger.warning(f"create_match: User {payload.user_email} not found")
            except Exception as e:
                logger.error(f"create_match: Error creating resume: {e}")
        
        if not resume:
            raise HTTPException(status_code=404, detail='Resume not found and could not be created. Provide resumeId or user_email with resume_name.')
        
        # Get job description
        jd = db.query(JobDescription).filter(JobDescription.id == payload.jobDescriptionId).first()
        if not jd:
            raise HTTPException(status_code=404, detail='Job description not found')
        
        # Update JD with metadata if provided
        if payload.jd_metadata:
            metadata = payload.jd_metadata
            if metadata.get('easy_apply_url') and not jd.easy_apply_url:
                jd.easy_apply_url = metadata.get('easy_apply_url')
            if metadata.get('work_type') and not jd.work_type:
                jd.work_type = metadata.get('work_type')
            if metadata.get('job_type') and not jd.job_type:
                jd.job_type = metadata.get('job_type')
            if metadata.get('company') and not jd.company:
                jd.company = metadata.get('company')
            db.commit()
            db.refresh(jd)

        # Get resume text from version or resume data
        resume_text = ''
        if payload.resume_version_id:
            rv = db.query(ResumeVersion).filter(ResumeVersion.id == payload.resume_version_id).first()
            if rv:
                resume_text = _resume_to_text(rv.resume_data)
        
        if not resume_text:
            rv = db.query(ResumeVersion).filter(ResumeVersion.resume_id == resume.id).order_by(ResumeVersion.version_number.desc()).first()
            resume_text = _resume_to_text(rv.resume_data) if rv else '\n'.join([resume.name or '', resume.title or '', resume.summary or ''])

        breakdown = _compute_match_breakdown(jd.content, resume_text, jd.extracted_keywords or {})

        # Use provided ATS score if available, otherwise use computed score
        final_score = payload.ats_score if payload.ats_score is not None else int(breakdown.get('score', 0))
        logger.info(f"create_match: Using ATS score: {final_score} (provided: {payload.ats_score}, computed: {breakdown.get('score', 0)})")

        # Ensure JSON fields are Python lists/dicts, not JSON strings
        matched_kw = breakdown.get('matched_keywords', [])
        missing_kw = breakdown.get('missing_keywords', [])
        
        logger.info(f"create_match: Raw breakdown data types - matched_keywords: {type(matched_kw)}, missing_keywords: {type(missing_kw)}")
        
        # Convert to list if it's already a string (defensive)
        if isinstance(matched_kw, str):
            logger.warning(f"create_match: matched_keywords is string, parsing: {matched_kw[:50]}")
            try:
                matched_kw = json.loads(matched_kw)
            except Exception as e:
                logger.error(f"create_match: Failed to parse matched_keywords string: {e}")
                matched_kw = []
        if isinstance(missing_kw, str):
            logger.warning(f"create_match: missing_keywords is string, parsing: {missing_kw[:50]}")
            try:
                missing_kw = json.loads(missing_kw)
            except Exception as e:
                logger.error(f"create_match: Failed to parse missing_keywords string: {e}")
                missing_kw = []
        
        # Ensure they're lists (handle sets, tuples, etc.)
        if not isinstance(matched_kw, list):
            matched_kw = list(matched_kw) if matched_kw else []
        if not isinstance(missing_kw, list):
            missing_kw = list(missing_kw) if missing_kw else []
        
        logger.info(f"create_match: Final data types - matched_keywords: {type(matched_kw)} (len={len(matched_kw)}), missing_keywords: {type(missing_kw)} (len={len(missing_kw)})")

        # Get user_id from resume
        if not resume.user_id:
            raise HTTPException(status_code=400, detail="Resume must belong to a user to create match session")

        ms = MatchSession(
            user_id=resume.user_id,
            resume_id=resume.id,
            job_description_id=jd.id,
            score=final_score,  # Use ATS score
            keyword_coverage=float(breakdown.get('keyword_coverage', 0.0)) if breakdown.get('keyword_coverage') is not None else 0.0,
            matched_keywords=matched_kw,
            missing_keywords=missing_kw,
            excess_keywords=[],  # Empty list for now
        )
        db.add(ms)
        db.commit()
        db.refresh(ms)

        logger.info(f"create_match: Successfully created match session {ms.id} with score {ms.score}")
        return { 'id': ms.id, **breakdown }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception('Failed to create match')
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.get('/matches/{match_id}')
def get_match(match_id: int, db: Session = Depends(get_db)):
    ms = db.query(MatchSession).filter(MatchSession.id == match_id).first()
    if not ms:
        raise HTTPException(status_code=404, detail='Match not found')
    return {
        'id': ms.id,
        'resume_id': ms.resume_id,
        'job_description_id': ms.job_description_id,
        'score': ms.score,
        'keyword_coverage': ms.keyword_coverage,
        'matched_keywords': ms.matched_keywords,
        'missing_keywords': ms.missing_keywords,
        'excess_keywords': ms.excess_keywords,
        'created_at': ms.created_at.isoformat(),
    }

# Compatibility aliases under /api/* for existing clients
@app.post('/api/job-descriptions')
def create_or_update_job_description_api(
    payload: JobDescriptionCreate,
    authorization: Optional[str] = Header(None, alias="Authorization"),
    db: Session = Depends(get_db)
):
    """Create or update job description, supporting both token-based (extension) and email-based auth"""
    try:
        logger.info(f"create_or_update_job_description_api: Received request with title={payload.title}, company={payload.company}")
        
        # Validate required fields
        if not payload.title or not payload.title.strip():
            raise HTTPException(status_code=400, detail="Title is required")
        if not payload.content or not payload.content.strip():
            raise HTTPException(status_code=400, detail="Content is required")
        
        # Extract user_email from payload if provided
        user_email = payload.user_email
        
        # If Authorization header is provided, try to extract user info
        # Note: Currently tokens are not validated/stored, so we allow saving without auth
        # This allows extension to work even if token validation isn't fully implemented
        if authorization and authorization.startswith('Bearer '):
            token = authorization.replace('Bearer ', '').strip()
            logger.info(f"create_or_update_job_description_api: Received token (not validated yet)")
            # TODO: Implement token validation to extract user_email from token
            # For now, allow saving without user (extension users can save anonymously)
        
        # If no user_email in payload and no token validation, allow saving without user
        if not user_email:
            logger.info("create_or_update_job_description_api: Saving job description without user (anonymous)")
        
        result = create_or_update_job_description(payload, db)
        logger.info(f"create_or_update_job_description_api: Successfully saved job description with id={result.get('id')}")
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"create_or_update_job_description_api: Error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to save job description: {str(e)}")

@app.get('/api/job-descriptions')
def list_job_descriptions_api(user_email: Optional[str] = None, db: Session = Depends(get_db)):
    try:
        # If no user_email provided, return jobs with no user_id (for extension without login)
        if not user_email:
            # Return jobs that were saved without user_id (from extension)
            items = db.query(JobDescription).filter(
                JobDescription.user_id.is_(None)
            ).order_by(JobDescription.created_at.desc()).limit(100).all()
            
            result = []
            for it in items:
                try:
                    priority_kw = it.priority_keywords
                    if priority_kw is not None:
                        if isinstance(priority_kw, str):
                            try:
                                priority_kw = json.loads(priority_kw)
                            except:
                                priority_kw = []
                        elif not isinstance(priority_kw, list):
                            priority_kw = []
                    
                    result.append({
                        'id': it.id,
                        'title': it.title or '',
                        'company': it.company or '',
                        'source': it.source or '',
                        'url': it.url or '',
                        'created_at': it.created_at.isoformat() if it.created_at else None,
                        'priority_keywords': priority_kw or [],
                        'last_match': None,
                    })
                except Exception as e:
                    logger.warning(f'Error processing JD {it.id}: {e}')
                    continue
            
            return result
        
        return list_job_descriptions(user_email, db)
    except Exception as e:
        logger.exception(f'Failed to list job descriptions via API (user_email: {user_email})')
        return []

@app.delete('/api/job-descriptions/{jd_id}')
def delete_job_description_api(jd_id: int, user_email: Optional[str] = Query(None), db: Session = Depends(get_db)):
    """Delete a job description"""
    try:
        logger.info(f"delete_job_description_api: Attempting to delete JD {jd_id}")
        
        # Find the job description
        jd = db.query(JobDescription).filter(JobDescription.id == jd_id).first()
        if not jd:
            logger.warning(f"delete_job_description_api: JD {jd_id} not found")
            raise HTTPException(status_code=404, detail='Job description not found')
        
        # If user_email is provided, verify ownership (optional for now, allow deletion)
        if user_email:
            user = db.query(User).filter(User.email == user_email).first()
            if user and jd.user_id and jd.user_id != user.id:
                logger.warning(f"delete_job_description_api: User {user_email} doesn't own JD {jd_id}")
                raise HTTPException(status_code=403, detail='Not authorized to delete this job description')
        
        # Delete associated match sessions first (if cascade doesn't handle it)
        match_sessions = db.query(MatchSession).filter(MatchSession.job_description_id == jd_id).all()
        for ms in match_sessions:
            db.delete(ms)
        
        # Delete the job description
        db.delete(jd)
        db.commit()
        
        logger.info(f"delete_job_description_api: Successfully deleted JD {jd_id}")
        return {
            "success": True,
            "message": "Job description deleted successfully",
            "id": jd_id
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"delete_job_description_api: Error deleting JD {jd_id}: {e}", exc_info=True)
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to delete job description: {str(e)}")

@app.get('/api/job-descriptions/{jd_id}')
def get_job_description_api(jd_id: int, db: Session = Depends(get_db)):
    return get_job_description(jd_id, db)

@app.post('/api/matches')
def create_match_api(payload: MatchCreate, db: Session = Depends(get_db)):
    return create_match(payload, db)

@app.get('/api/matches/{match_id}')
def get_match_api(match_id: int, db: Session = Depends(get_db)):
    return get_match(match_id, db)

@app.get('/api/job-descriptions/{jd_id}/matches')
def get_job_description_matches(jd_id: int, db: Session = Depends(get_db)):
    """Get all match sessions for a specific job description"""
    try:
        matches = db.query(MatchSession).filter(
            MatchSession.job_description_id == jd_id
        ).order_by(MatchSession.created_at.desc()).all()
        
        result = []
        for match in matches:
            resume_name = None
            resume_version_id = None
            
            if match.resume_id:
                resume = db.query(Resume).filter(Resume.id == match.resume_id).first()
                if resume:
                    resume_name = resume.name
                    # Get latest version ID if available
                    try:
                        latest_version = db.query(ResumeVersion).filter(
                            ResumeVersion.resume_id == resume.id
                        ).order_by(ResumeVersion.version_number.desc()).first()
                        if latest_version:
                            resume_version_id = latest_version.id
                    except Exception as e:
                        logger.warning(f'Failed to get resume version for resume {resume.id}: {e}')
            
            result.append({
                'id': match.id,
                'score': match.score,
                'resume_id': match.resume_id,
                'resume_name': resume_name,
                'resume_version_id': resume_version_id,
                'keyword_coverage': match.keyword_coverage,
                'matched_keywords': match.matched_keywords or [],
                'missing_keywords': match.missing_keywords or [],
                'excess_keywords': match.excess_keywords or [],
                'created_at': match.created_at.isoformat() if match.created_at else None,
            })
        
        return result
    except Exception as e:
        logger.exception('Failed to get job description matches')
        raise HTTPException(status_code=500, detail=str(e))

